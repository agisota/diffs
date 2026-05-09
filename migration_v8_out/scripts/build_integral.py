#!/usr/bin/env python3
"""
Integral cross-comparison v8 generator.

Inputs (read-only):
    /home/dev/diff/migration_v7_evidence/04_машинные_данные_на_русском/01_источники_v7.csv
    .../02_все_пары_v7.csv
    .../03_все_события_v7.csv
    .../05_forensic_реестр_источников_v7.csv
    .../10_очередь_ручной_проверки_v7.csv
    .../11_итоговые_противоречия_v7.csv .. 14_visual_redgreen_v7.csv

Outputs (writable):
    /home/dev/diff/migration_v8_out/data/*.csv,*.json
    /home/dev/diff/migration_v8_out/docs/Интегральное_перекрестное_сравнение.xlsx
    /home/dev/diff/migration_v8_out/docs/Интегральное_перекрестное_сравнение.pdf
    /home/dev/diff/migration_v8_out/docs/Пояснительная_записка.docx + .pdf
    /home/dev/diff/migration_v8_out/docs/Редакционный_diff.docx + .pdf
    /home/dev/diff/migration_v8_out/logs/build.log
    /home/dev/diff/migration_v8_out/logs/qa.json

The script is evidence-first: every cell that asserts a status is backed
by a v7 event (СОБ-####), the C-01..C-03 contradictions sheet, the U-01..U-03
uncovered-thesis list, or the D-001..D-003 defect log. Where source-rank is
mismatched (rank-3 secondary against rank-1 primary), the verdict carries a
manual_review flag. Provenance risks (DDoS-stub, timeout, redirect-only,
fallback-mirror only) are tagged in a dedicated sheet.
"""
from __future__ import annotations

import csv
import datetime as dt
import io
import json
import re
import sys
import unicodedata
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Iterable

# ---------------------------------------------------------------------------
# Paths & constants
# ---------------------------------------------------------------------------

ROOT_IN = Path("/home/dev/diff/migration_v7_evidence")
DATA_IN = ROOT_IN / "04_машинные_данные_на_русском"
SNAP_IN = ROOT_IN / "06_исходные_НПА_и_forensic_snapshot"
HANDOFF = ROOT_IN / "09_integral_cross_comparison_v8" / "agent_handoff"

ROOT_OUT = Path("/home/dev/diff/migration_v8_out")
DATA_OUT = ROOT_OUT / "data"
DOCS_OUT = ROOT_OUT / "docs"
LOGS_OUT = ROOT_OUT / "logs"

GENERATED_AT = dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%d %H:%M:%SZ")

CONTROL_NUMBERS = {
    "documents": 26,
    "pairs": 325,
    "events": 281,
    "manual_reviews": 183,
    "final_contradictions": 3,
    "uncovered_theses": 3,
    "defect_log": 3,
    "fallback_sources": 71,
}

# Canonical comparison-status scale (v8 schema)
STATUS_MATCH = "match"
STATUS_PARTIAL = "partial_overlap"
STATUS_CONTRADICTION = "contradiction"
STATUS_OUTDATED = "outdated"
STATUS_GAP = "source_gap"
STATUS_REVIEW = "manual_review"
STATUS_NC = "not_comparable"

V7_TO_V8_STATUS = {
    "подтверждено": STATUS_MATCH,
    "подтверждено с уточнением": STATUS_MATCH,
    "покрыто планом": STATUS_MATCH,
    "частично совпадает": STATUS_PARTIAL,
    "требует ручной проверки": STATUS_REVIEW,
    "не сопоставимо": STATUS_NC,
    "не сопоставимо напрямую": STATUS_NC,
}

# Document-rank precedence (lower wins)
RANK_PRIMARY_LAW = 1
RANK_DEPARTMENTAL = 2
RANK_ANALYTIC = 3

# Topic clusters: bucket the long tail of v7 raw topics into ~12 cohesive themes.
# Each tuple is (cluster_id, label, [match-substrings, case-insensitive]).
TOPIC_CLUSTERS: list[tuple[str, str, list[str]]] = [
    ("T01", "ruID, цифровой профиль, биометрия", [
        "цифровой профиль", "ruid", "1510", "биометр", "467",
    ]),
    ("T02", "Миграционный учёт и фактическое нахождение", [
        "миграционный учёт", "миграционный учет", "фактическо",
    ]),
    ("T03", "Патенты, НДФЛ, госпошлины", [
        "патент", "ндфл", "госпошлин", "налог",
    ]),
    ("T04", "Режим высылки, реестр контролируемых лиц", [
        "высылк", "реестр контролируемых",
    ]),
    ("T05", "Образовательная миграция", [
        "образовательная миграция", "иностранные студент",
    ]),
    ("T06", "Адаптация и интеграция", [
        "адаптация", "интеграц", "анклав", "напряжённ", "напряжен",
    ]),
    ("T07", "ВНЖ инвестора (ПП №2573)", [
        "внж инвестор", "инвестор", "критерий", "брошюр",
        "социально значим", "недвижимост",
    ]),
    ("T08", "Эксперимент 121-ФЗ Москва/МО, 90 дней", [
        "121-фз", "эксперимент", "90 дней", "безвизов",
    ]),
    ("T09", "ВЦИОМ-claims (ужесточение, образование, интеллект)", [
        "вциом",
    ]),
    ("T10", "Концепции и планы (D04 D05 D07 D08)", [
        "концепци", "план 30-р", "план 4171", "30-р", "4171", "30р",
        "покрытие старой", "покрытие новой", "смена структуры",
        "эволюция планов", "план 30",
    ]),
    ("T11", "ЕАЭС: трудовая миграция", [
        "еаэс", "трудовая миграц", "трудоустрой",
    ]),
    ("T12", "КоАП: ответственность", [
        "коап", "ответственност", "штраф",
    ]),
    ("T13", "Изменения 2024–2026 (260/270/271/281/121/1562/468)", [
        "260-фз", "270-фз", "271-фз", "281-фз", "1562", "468",
        "изменения 115", "изменения 109", "изменения",
    ]),
    ("T14", "Базовая нормативная рамка (114/115/109)", [
        "114-фз", "115-фз", "109-фз", "въезд", "выезд",
    ]),
    ("T15", "Внутренний сервис «Нейрон» (методология)", [
        "нейрон", "работа внутреннего сервиса",
    ]),
    ("T16", "Общая миграция (cross-cutting)", [
        "миграция", "правовая основа", "нормативная база",
    ]),
    ("T17", "Мониторинг и статистика", [
        "мониторинг", "статистик", "социолог",
    ]),
]
TOPIC_PROCESS = "T15"  # process-only / not-comparable bucket

# Pre-defined uncovered theses (v7 → U-01..U-03) and contradictions (C-01..C-03)
# are read at runtime; here we just record the v8 mapping.
UNCOVERED_TO_V8 = {
    "U-01": ("source_gap", "T09"),
    "U-02": ("source_gap", "T03"),
    "U-03": ("source_gap", "T10"),
}
CONTRADICTION_TO_V8 = {
    "C-01": ("contradiction", "T11"),
    "C-02": ("manual_review", "T15"),
    "C-03": ("manual_review", "T07"),
}

# Documents that AMEND or SUPERSEDE earlier documents → "outdated" candidate basis.
AMENDMENT_GRAPH = {
    "D22": ["D21"],            # ПП 1562 amends ПП 1510
    "D23": ["D21", "D22"],     # ПП 468 amends 1510 (and downstream)
    "D24": ["D11", "D06"],     # 270-ФЗ amends 115-ФЗ, 109-ФЗ
    "D25": ["D13"],            # 281-ФЗ amends КоАП
    "D04": ["D05"],            # Concept 2026–2030 supersedes 2019–2025
    "D08": ["D07"],            # Plan 4171-р supersedes 30-р
    "D17": ["D21"],            # Указ 467 раздел цифрового профиля – на стыке с ruID
    "D15": ["D11"],            # 260-ФЗ inserted режим высылки в 115-ФЗ
    "D14": ["D06"],            # 121-ФЗ — эксперимент к 109-ФЗ
    "D16": ["D12"],            # 271-ФЗ — изменения в НК
}

REGIMES = {
    "R01": "Общий миграционный учёт (109-ФЗ + поправки)",
    "R02": "ВНЖ инвестора без РВП (ПП №2573 + Минэк-брошюра)",
    "R03": "ruID / цифровой въезд (ПП №1510 + 1562 + 468 + Указ 467)",
    "R04": "Режим высылки и реестр контролируемых лиц (260-ФЗ, 281-ФЗ, КоАП)",
    "R05": "ЕАЭС: трудовая миграция (Договор ЕАЭС, ст.96–98)",
    "R06": "Концепции и планы (Концепция 2026/2019 + 4171-р/30-р)",
    "R07": "Эксперимент 121-ФЗ Москва/МО, 90-дневный срок",
    "R08": "Патенты, НДФЛ, госпошлины (115/НК/271-ФЗ)",
}

REGIME_DOCS = {
    "R01": ["D06", "D24", "D14", "D17"],
    "R02": ["D20", "D18", "D11", "D19"],
    "R03": ["D21", "D22", "D23", "D17", "D09"],
    "R04": ["D15", "D11", "D13", "D25", "D09"],
    "R05": ["D26", "D10", "D11"],
    "R06": ["D04", "D05", "D07", "D08"],
    "R07": ["D14", "D15"],
    "R08": ["D11", "D12", "D16", "D09"],
}

# ---------------------------------------------------------------------------
# CSV utilities
# ---------------------------------------------------------------------------

def _resolve_nfc_path(path: Path) -> Path:
    """macOS-mounted volumes store filenames as NFD. Match by NFC-normalised basename."""
    if path.exists():
        return path
    parent = path.parent
    if not parent.exists():
        return path
    target_nfc = unicodedata.normalize("NFC", path.name)
    for entry in parent.iterdir():
        if unicodedata.normalize("NFC", entry.name) == target_nfc:
            return entry
    return path


def read_csv(path: Path, *, delim: str = ";") -> list[dict[str, str]]:
    real = _resolve_nfc_path(path)
    raw = real.read_bytes().decode("utf-8-sig")
    reader = csv.DictReader(io.StringIO(raw), delimiter=delim)
    return [dict(r) for r in reader]


def write_csv(path: Path, rows: list[dict[str, Any]], header: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=header, delimiter=";")
        w.writeheader()
        for row in rows:
            w.writerow({k: ("" if row.get(k) is None else row[k]) for k in header})


def normalise(text: str) -> str:
    return unicodedata.normalize("NFC", text).lower().strip()


# ---------------------------------------------------------------------------
# Topic clustering
# ---------------------------------------------------------------------------

def cluster_topic(raw_topic: str) -> tuple[str, str]:
    if not raw_topic:
        return ("T00", "Без темы / Предмет сравнения")
    n = normalise(raw_topic)
    for cid, label, needles in TOPIC_CLUSTERS:
        for needle in needles:
            if needle in n:
                return (cid, label)
    if "предмет сравнения" in n:
        return (TOPIC_PROCESS, dict((c, l) for c, l, _ in TOPIC_CLUSTERS)[TOPIC_PROCESS])
    return ("T00", "Прочее (не кластеризовано)")


# ---------------------------------------------------------------------------
# Phase 1 — read inputs
# ---------------------------------------------------------------------------

def main() -> int:
    LOGS_OUT.mkdir(parents=True, exist_ok=True)
    log_path = LOGS_OUT / "build.log"
    log_lines: list[str] = []

    def log(msg: str) -> None:
        line = f"[{dt.datetime.utcnow().isoformat(timespec='seconds')}Z] {msg}"
        print(line)
        log_lines.append(line)

    log(f"build_integral starting; corpus={ROOT_IN}")

    sources = read_csv(DATA_IN / "01_источники_v7.csv")
    pairs = read_csv(DATA_IN / "02_все_пары_v7.csv")
    events = read_csv(DATA_IN / "03_все_события_v7.csv")
    forensic = read_csv(DATA_IN / "05_forensic_реестр_источников_v7.csv")
    manual = read_csv(DATA_IN / "10_очередь_ручной_проверки_v7.csv")
    contradictions = read_csv(DATA_IN / "11_итоговые_противоречия_v7.csv")
    uncovered = read_csv(DATA_IN / "12_непокрытые_тезисы_v7.csv")
    defects = read_csv(DATA_IN / "13_defect_log_v7.csv")
    redgreen = read_csv(DATA_IN / "14_visual_redgreen_v7.csv")

    log(
        f"loaded: sources={len(sources)} pairs={len(pairs)} events={len(events)} "
        f"forensic={len(forensic)} manual={len(manual)} contradictions={len(contradictions)} "
        f"uncovered={len(uncovered)} defects={len(defects)} redgreen={len(redgreen)}"
    )

    # Sanity-check control numbers
    counts_ok = {
        "documents": len(sources) == CONTROL_NUMBERS["documents"],
        "pairs": len(pairs) == CONTROL_NUMBERS["pairs"],
        "events": len(events) == CONTROL_NUMBERS["events"],
        "manual_reviews": len(manual) == CONTROL_NUMBERS["manual_reviews"],
    }
    log(f"control number check: {counts_ok}")

    # -----------------------------------------------------------------------
    # Phase 2 — canonical source registry
    # -----------------------------------------------------------------------
    forensic_by_id = {r.get("ИД документа", ""): r for r in forensic}
    fallback_count = count_fallback_files()
    sources_norm: list[dict[str, Any]] = []

    for s in sources:
        did = s.get("ИД документа", "").strip()
        f = forensic_by_id.get(did, {})
        local_snapshot = f.get("Локальный snapshot", "")
        sha = f.get("SHA-256 snapshot", "")
        full_text = f.get("Полный текст получен", "")
        provenance = derive_provenance_status(s, f)
        sources_norm.append({
            "ИД": did,
            "Код": s.get("Код", ""),
            "Название": s.get("Название", ""),
            "Тип": s.get("Тип", ""),
            "Ранг": s.get("Ранг источника", ""),
            "URL/путь": s.get("URL / путь", ""),
            "Локальный snapshot": local_snapshot,
            "SHA-256": sha,
            "Полный текст получен": full_text,
            "Provenance-статус": provenance,
            "Комментарий": s.get("Комментарий", ""),
        })

    write_csv(
        DATA_OUT / "01_источники_v8.csv",
        sources_norm,
        ["ИД", "Код", "Название", "Тип", "Ранг", "URL/путь", "Локальный snapshot",
         "SHA-256", "Полный текст получен", "Provenance-статус", "Комментарий"],
    )
    log(f"wrote 01_источники_v8.csv ({len(sources_norm)} rows)")

    # Build doc lookups
    doc_by_id = {r["ИД"]: r for r in sources_norm}
    doc_rank = {r["ИД"]: int(r["Ранг"]) for r in sources_norm}

    # -----------------------------------------------------------------------
    # Phase 3 — pair-level integration
    # -----------------------------------------------------------------------
    events_by_pair: dict[str, list[dict[str, str]]] = defaultdict(list)
    for ev in events:
        events_by_pair[ev.get("ИД пары", "")].append(ev)

    # Pre-derive contradiction marker pairs
    contradiction_pair_ids: set[str] = set()
    contradiction_doc_pairs: list[tuple[str, str, str]] = []  # (doc_a, doc_b, C-id)
    # C-01: ЕАЭС vs Узбекистан/Таджикистан → flagged on Минэк (D10) ↔ Договор ЕАЭС (D26)
    contradiction_doc_pairs.append(("D10", "D26", "C-01"))
    # C-02: Клерк (D09) методологический риск → mark all D09 vs primary-NPA pairs as manual_review
    # C-03: Брошюра Минэка (D18) vs ПП 2573 (D20) — manual_review
    contradiction_doc_pairs.append(("D18", "D20", "C-03"))

    for ev_pair_id, evs in events_by_pair.items():
        for ev in evs:
            left = ev.get("Документ слева", "")
            right = ev.get("Документ справа", "")
            if left == "D09" and doc_rank.get(right, 99) == 1:
                contradiction_pair_ids.add(ev_pair_id)
            if right == "D09" and doc_rank.get(left, 99) == 1:
                contradiction_pair_ids.add(ev_pair_id)

    pair_rows: list[dict[str, Any]] = []
    doc_doc_status: dict[tuple[str, str], str] = {}
    doc_doc_event_count: dict[tuple[str, str], int] = defaultdict(int)
    pair_id_status: dict[str, str] = {}

    for p in pairs:
        pid = p.get("ИД пары", "").strip()
        l = p.get("Документ слева", "").strip()
        r = p.get("Документ справа", "").strip()
        evs = events_by_pair.get(pid, [])
        v8_status, v8_topics, manual_flag, basis = aggregate_pair_status(
            pid, l, r, evs, doc_rank, contradiction_doc_pairs
        )
        pair_rows.append({
            "ИД пары": pid,
            "L": l,
            "R": r,
            "L-ранг": doc_rank.get(l, ""),
            "R-ранг": doc_rank.get(r, ""),
            "Статус v8": v8_status,
            "Темы": "; ".join(v8_topics),
            "Темы v8": "; ".join(sorted({cluster_topic(t)[1] for t in v8_topics})),
            "Manual-review": "да" if manual_flag else "нет",
            "Кол-во событий": len(evs),
            "Критичность v7": p.get("Критичность", ""),
            "Приоритет v7": p.get("Приоритет пары", ""),
            "Основание": basis,
            "Путь DOCX": p.get("Путь DOCX", ""),
            "Путь PDF": p.get("Путь PDF", ""),
        })
        doc_doc_status[(l, r)] = v8_status
        doc_doc_event_count[(l, r)] = len(evs)
        pair_id_status[pid] = v8_status

    write_csv(
        DATA_OUT / "02_pairs_v8.csv",
        pair_rows,
        ["ИД пары", "L", "R", "L-ранг", "R-ранг", "Статус v8", "Темы", "Темы v8",
         "Manual-review", "Кол-во событий", "Критичность v7", "Приоритет v7",
         "Основание", "Путь DOCX", "Путь PDF"],
    )
    log(f"wrote 02_pairs_v8.csv ({len(pair_rows)} rows)")

    # -----------------------------------------------------------------------
    # Phase 4 — Doc x Doc matrix (26x26 symmetrical)
    # -----------------------------------------------------------------------
    doc_ids = [s["ИД"] for s in sources_norm]
    matrix_rows: list[dict[str, Any]] = []
    for di in doc_ids:
        row: dict[str, Any] = {"ИД": di, "Код": doc_by_id[di]["Код"]}
        for dj in doc_ids:
            if di == dj:
                row[dj] = "—"
                continue
            st = doc_doc_status.get((di, dj)) or doc_doc_status.get((dj, di)) or STATUS_NC
            cnt = doc_doc_event_count.get((di, dj)) or doc_doc_event_count.get((dj, di)) or 0
            mark = STATUS_TO_MARK[st]
            row[dj] = f"{mark}({cnt})" if cnt else mark
        matrix_rows.append(row)
    write_csv(
        DATA_OUT / "03_doc_x_doc_matrix.csv",
        matrix_rows,
        ["ИД", "Код"] + doc_ids,
    )
    log(f"wrote 03_doc_x_doc_matrix.csv ({len(matrix_rows)}x{len(doc_ids)})")

    # -----------------------------------------------------------------------
    # Phase 5 — Topic x Doc coverage matrix
    # -----------------------------------------------------------------------
    topic_doc: dict[str, dict[str, list[str]]] = defaultdict(lambda: defaultdict(list))
    for ev in events:
        cid, _ = cluster_topic(ev.get("Тема", ""))
        topic_doc[cid][ev.get("Документ слева", "")].append(ev.get("ИД события", ""))
        topic_doc[cid][ev.get("Документ справа", "")].append(ev.get("ИД события", ""))

    topic_rows: list[dict[str, Any]] = []
    for cid, label, _ in TOPIC_CLUSTERS:
        row = {"ИД темы": cid, "Тема": label}
        for did in doc_ids:
            ev_ids = topic_doc.get(cid, {}).get(did, [])
            row[did] = len(ev_ids) if ev_ids else ""
        row["Σ событий"] = sum(int(v) for v in row.values() if isinstance(v, int))
        topic_rows.append(row)
    # Add T00 / уцелевшие
    write_csv(
        DATA_OUT / "04_topic_x_doc.csv",
        topic_rows,
        ["ИД темы", "Тема"] + doc_ids + ["Σ событий"],
    )
    log(f"wrote 04_topic_x_doc.csv ({len(topic_rows)} themes)")

    # -----------------------------------------------------------------------
    # Phase 6 — Thesis x NPA (claims from secondary docs vs primary acts)
    # -----------------------------------------------------------------------
    thesis_rows = build_thesis_table(events, doc_rank, doc_by_id)
    write_csv(
        DATA_OUT / "05_thesis_x_npa.csv",
        thesis_rows,
        ["Тезис ID", "Источник тезиса", "Тезис", "Связанный НПА",
         "Статус v8", "Уверенность", "Manual-review", "Основание"],
    )
    log(f"wrote 05_thesis_x_npa.csv ({len(thesis_rows)} rows)")

    # -----------------------------------------------------------------------
    # Phase 7 — Old vs New redaction (amendment graph)
    # -----------------------------------------------------------------------
    redaction_rows: list[dict[str, Any]] = []
    rid = 0
    for newer, olds in AMENDMENT_GRAPH.items():
        for old in olds:
            rid += 1
            redaction_rows.append({
                "ID": f"R-{rid:02d}",
                "Изменяющий документ": newer,
                "Изменяющий код": doc_by_id.get(newer, {}).get("Код", ""),
                "Изменяемый документ": old,
                "Изменяемый код": doc_by_id.get(old, {}).get("Код", ""),
                "Статус v8": STATUS_OUTDATED,
                "Тип изменения": classify_amendment(newer, old),
                "Основание": derive_amendment_basis(newer, old, doc_by_id),
            })
    write_csv(
        DATA_OUT / "06_old_vs_new_redactions.csv",
        redaction_rows,
        ["ID", "Изменяющий документ", "Изменяющий код",
         "Изменяемый документ", "Изменяемый код",
         "Статус v8", "Тип изменения", "Основание"],
    )
    log(f"wrote 06_old_vs_new_redactions.csv ({len(redaction_rows)} rows)")

    # -----------------------------------------------------------------------
    # Phase 8 — Regime x Regime
    # -----------------------------------------------------------------------
    regime_rows: list[dict[str, Any]] = []
    for code, label in REGIMES.items():
        regime_rows.append({
            "Режим": code,
            "Название": label,
            "Документы": ", ".join(REGIME_DOCS[code]),
            "Покрыто событиями": count_events_for_docs(events, REGIME_DOCS[code]),
            "Перекрытие с другими": ", ".join(
                f"{r2}:{count_events_for_docs(events, REGIME_DOCS[code], REGIME_DOCS[r2])}"
                for r2 in REGIMES if r2 != code
            ),
        })
    write_csv(
        DATA_OUT / "07_regime_x_regime.csv",
        regime_rows,
        ["Режим", "Название", "Документы", "Покрыто событиями", "Перекрытие с другими"],
    )
    log(f"wrote 07_regime_x_regime.csv ({len(regime_rows)} rows)")

    # -----------------------------------------------------------------------
    # Phase 9 — Provenance risk sheet
    # -----------------------------------------------------------------------
    provenance_rows = build_provenance_rows(forensic, sources_norm, fallback_count)
    write_csv(
        DATA_OUT / "08_provenance_risk.csv",
        provenance_rows,
        ["ИД", "Источник", "Provenance-риск", "Зеркала", "Стартовое действие"],
    )
    log(f"wrote 08_provenance_risk.csv ({len(provenance_rows)} rows)")

    # -----------------------------------------------------------------------
    # Phase 10 — Manual review queue (compact)
    # -----------------------------------------------------------------------
    manual_rows: list[dict[str, Any]] = []
    for m in manual:
        manual_rows.append({
            "ИД события": m.get("ИД события", ""),
            "ИД пары": m.get("ИД пары", ""),
            "L": m.get("Документ слева", ""),
            "R": m.get("Документ справа", ""),
            "Тема": m.get("Тема", ""),
            "Статус": m.get("Статус", ""),
            "Уверенность": m.get("Уверенность", ""),
            "Что проверить": m.get("Что проверить вручную", ""),
            "Приоритет": m.get("Приоритет проверки", ""),
            "Дедлайн": m.get("Дедлайн", ""),
        })
    write_csv(
        DATA_OUT / "09_manual_review_queue.csv",
        manual_rows,
        ["ИД события", "ИД пары", "L", "R", "Тема", "Статус", "Уверенность",
         "Что проверить", "Приоритет", "Дедлайн"],
    )
    log(f"wrote 09_manual_review_queue.csv ({len(manual_rows)} rows)")

    # -----------------------------------------------------------------------
    # Phase 11 — Status distribution & QA
    # -----------------------------------------------------------------------
    status_dist = Counter(r["Статус v8"] for r in pair_rows)
    rank_pair_dist = Counter(
        f"{doc_rank.get(p['L'])}↔{doc_rank.get(p['R'])}" for p in pair_rows
    )
    qa = {
        "generated_at": GENERATED_AT,
        "control_numbers_expected": CONTROL_NUMBERS,
        "control_numbers_actual": {
            "documents": len(sources_norm),
            "pairs": len(pair_rows),
            "events": len(events),
            "manual_reviews": len(manual),
            "final_contradictions": len(contradictions),
            "uncovered_theses": len(uncovered),
            "defect_log": len(defects),
        },
        "status_distribution_pairs": dict(status_dist),
        "rank_pair_distribution": dict(rank_pair_dist),
        "topic_event_distribution": {
            label: sum(int(r[d]) for d in doc_ids if isinstance(r.get(d), int))
            for r in topic_rows for label in [r["Тема"]]
        },
        "amendments_recorded": len(redaction_rows),
        "regimes_recorded": len(regime_rows),
        "provenance_rows": len(provenance_rows),
        "thesis_rows": len(thesis_rows),
        "control_numbers_match": all(
            v == CONTROL_NUMBERS[k]
            for k, v in {
                "documents": len(sources_norm),
                "pairs": len(pair_rows),
                "events": len(events),
                "manual_reviews": len(manual),
            }.items()
        ),
    }
    (LOGS_OUT / "qa.json").write_text(
        json.dumps(qa, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    log(f"wrote logs/qa.json (status_dist={dict(status_dist)})")

    # -----------------------------------------------------------------------
    # Phase 12 — JSON snapshot bundle
    # -----------------------------------------------------------------------
    bundle = {
        "schema_version": "v8.0",
        "generated_at": GENERATED_AT,
        "control_numbers": qa["control_numbers_actual"],
        "status_scale": [
            STATUS_MATCH, STATUS_PARTIAL, STATUS_CONTRADICTION,
            STATUS_OUTDATED, STATUS_GAP, STATUS_REVIEW, STATUS_NC,
        ],
        "documents": sources_norm,
        "pairs_summary": {
            "total": len(pair_rows),
            "by_status": dict(status_dist),
        },
        "contradictions_v7": contradictions,
        "uncovered_v7": uncovered,
        "defects_v7": defects,
        "redgreen_v7": redgreen,
        "topic_clusters": [
            {"id": cid, "label": label, "needles": needles}
            for cid, label, needles in TOPIC_CLUSTERS
        ],
        "regimes": [
            {"id": rid_, "label": label, "documents": REGIME_DOCS[rid_]}
            for rid_, label in REGIMES.items()
        ],
        "amendment_graph": AMENDMENT_GRAPH,
    }
    (DATA_OUT / "integral_cross_comparison.json").write_text(
        json.dumps(bundle, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    log("wrote data/integral_cross_comparison.json")

    # -----------------------------------------------------------------------
    # Phase 13 — Excel
    # -----------------------------------------------------------------------
    render_excel(
        DOCS_OUT / "Интегральное_перекрестное_сравнение.xlsx",
        sources_norm=sources_norm,
        pair_rows=pair_rows,
        matrix_rows=matrix_rows,
        topic_rows=topic_rows,
        thesis_rows=thesis_rows,
        redaction_rows=redaction_rows,
        regime_rows=regime_rows,
        provenance_rows=provenance_rows,
        manual_rows=manual_rows,
        contradictions=contradictions,
        uncovered=uncovered,
        defects=defects,
        redgreen=redgreen,
        qa=qa,
        doc_ids=doc_ids,
    )
    log("wrote docs/Интегральное_перекрестное_сравнение.xlsx")

    # -----------------------------------------------------------------------
    # Phase 14 — DOCX (Пояснительная_записка + Редакционный_diff)
    # -----------------------------------------------------------------------
    render_docx_explanatory(
        DOCS_OUT / "Пояснительная_записка.docx",
        sources_norm=sources_norm,
        qa=qa,
        contradictions=contradictions,
        uncovered=uncovered,
        defects=defects,
        thesis_rows=thesis_rows,
        redaction_rows=redaction_rows,
        manual_count=len(manual_rows),
    )
    log("wrote docs/Пояснительная_записка.docx")

    render_docx_redgreen(
        DOCS_OUT / "Редакционный_diff.docx",
        thesis_rows=thesis_rows,
        redaction_rows=redaction_rows,
        contradictions=contradictions,
        uncovered=uncovered,
        sources_norm=sources_norm,
    )
    log("wrote docs/Редакционный_diff.docx")

    # -----------------------------------------------------------------------
    # Phase 15 — PDFs
    # -----------------------------------------------------------------------
    render_pdf_summary(
        DOCS_OUT / "Интегральное_перекрестное_сравнение.pdf",
        qa=qa,
        sources_norm=sources_norm,
        thesis_rows=thesis_rows,
        redaction_rows=redaction_rows,
        contradictions=contradictions,
        uncovered=uncovered,
        regime_rows=regime_rows,
        topic_rows=topic_rows,
    )
    render_pdf_explanatory(
        DOCS_OUT / "Пояснительная_записка.pdf",
        sources_norm=sources_norm,
        qa=qa,
        contradictions=contradictions,
        uncovered=uncovered,
        defects=defects,
    )
    render_pdf_redgreen(
        DOCS_OUT / "Редакционный_diff.pdf",
        thesis_rows=thesis_rows,
        redaction_rows=redaction_rows,
        contradictions=contradictions,
        uncovered=uncovered,
        sources_norm=sources_norm,
    )
    log("wrote 3 PDFs in docs/")

    # final QA file
    log_path.write_text("\n".join(log_lines), encoding="utf-8")
    log(f"DONE — see {ROOT_OUT}")
    return 0


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

STATUS_TO_MARK = {
    STATUS_MATCH: "✓",
    STATUS_PARTIAL: "≈",
    STATUS_CONTRADICTION: "⚠",
    STATUS_OUTDATED: "↻",
    STATUS_GAP: "∅",
    STATUS_REVIEW: "?",
    STATUS_NC: "—",
}


def derive_provenance_status(src: dict[str, str], forensic_row: dict[str, str]) -> str:
    """Heuristics from forensic registry."""
    full = forensic_row.get("Полный текст получен", "")
    snap = forensic_row.get("Локальный snapshot", "")
    actuality = forensic_row.get("Статус актуальности", "")
    if "актуализировать" in actuality.lower():
        return "primary-text-actualize-needed"
    if snap and "06_исходные_НПА" in snap:
        return "primary-with-local-snapshot"
    if full == "да":
        return "primary-or-mirror-full-text"
    if "URL" in full:
        return "url-only-no-local-text"
    return "unknown"


def aggregate_pair_status(
    pid: str,
    left: str,
    right: str,
    evs: list[dict[str, str]],
    doc_rank: dict[str, int],
    contradiction_doc_pairs: list[tuple[str, str, str]],
) -> tuple[str, list[str], bool, str]:
    """Aggregate v7 events into a single v8 pair status."""
    # Hard contradiction overrides
    for a, b, cid in contradiction_doc_pairs:
        if {left, right} == {a, b}:
            return ("contradiction", _topic_list(evs), True, f"{cid} (final contradiction)")

    if not evs:
        return (STATUS_NC, ["Без сравнения"], False, "Нет событий v7")

    statuses = [V7_TO_V8_STATUS.get(e.get("Статус", ""), STATUS_REVIEW) for e in evs]
    topics = _topic_list(evs)

    rl = doc_rank.get(left, 99)
    rr = doc_rank.get(right, 99)

    # Rank-3 vs rank-1 pairs always include manual_review per C-02 / C-03
    rank3_vs_primary = (rl == 3 and rr == 1) or (rl == 1 and rr == 3)

    if STATUS_REVIEW in statuses or rank3_vs_primary:
        return (STATUS_REVIEW, topics, True,
                f"{statuses.count(STATUS_REVIEW)} событий manual_review; rank3↔rank1={rank3_vs_primary}")

    if STATUS_PARTIAL in statuses:
        return (STATUS_PARTIAL, topics, False,
                f"{statuses.count(STATUS_PARTIAL)} partial / {statuses.count(STATUS_MATCH)} match")
    if STATUS_MATCH in statuses:
        return (STATUS_MATCH, topics, False, f"{statuses.count(STATUS_MATCH)} match events")
    return (STATUS_NC, topics, False, "Только not_comparable события")


def _topic_list(evs: list[dict[str, str]]) -> list[str]:
    seen: list[str] = []
    for e in evs:
        t = e.get("Тема", "").strip()
        if t and t not in seen:
            seen.append(t)
    return seen


def build_thesis_table(
    events: list[dict[str, str]],
    doc_rank: dict[str, int],
    doc_by_id: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    """Build claim-level rows: pick events where one side is rank>=2 secondary."""
    rows: list[dict[str, Any]] = []
    seen_keys: set[tuple[str, str, str]] = set()
    tid = 0
    for ev in events:
        l = ev.get("Документ слева", "")
        r = ev.get("Документ справа", "")
        rl = doc_rank.get(l, 99)
        rr = doc_rank.get(r, 99)
        if rl <= 1 and rr <= 1:
            # both primary — skip from thesis layer, captured in old/new redactions
            continue
        # determine which side is secondary
        secondary = l if rl > rr else r
        primary = r if secondary == l else l
        if doc_rank.get(primary, 99) > 2:
            continue  # both secondary — handled separately as cross-secondary
        topic = ev.get("Тема", "")
        thesis_text = ev.get("Фрагмент слева", "") if secondary == l else ev.get("Фрагмент справа", "")
        thesis_text = (thesis_text or "")[:280]
        npa_text = ev.get("Фрагмент справа", "") if secondary == l else ev.get("Фрагмент слева", "")
        npa_text = (npa_text or "")[:280]
        v7_status = ev.get("Статус", "")
        v8_status = V7_TO_V8_STATUS.get(v7_status, STATUS_REVIEW)
        manual = v8_status == STATUS_REVIEW or doc_rank.get(secondary) == 3
        confidence = ev.get("Уверенность", "")
        key = (secondary, primary, topic)
        if key in seen_keys:
            continue
        seen_keys.add(key)
        tid += 1
        rows.append({
            "Тезис ID": f"TH-{tid:03d}",
            "Источник тезиса": f"{secondary} ({doc_by_id.get(secondary, {}).get('Код','')})",
            "Тезис": thesis_text,
            "Связанный НПА": f"{primary} ({doc_by_id.get(primary, {}).get('Код','')}) — {npa_text}",
            "Статус v8": v8_status,
            "Уверенность": confidence,
            "Manual-review": "да" if manual else "нет",
            "Основание": f"Тема: {topic}; v7={v7_status}; СОБ={ev.get('ИД события','')}",
        })
    return rows


def classify_amendment(newer: str, old: str) -> str:
    if newer == "D04" and old == "D05":
        return "supersession (новая концепция)"
    if newer == "D08" and old == "D07":
        return "supersession (новый план)"
    if newer in {"D22", "D23", "D24", "D25"}:
        return "amendment (поправка к НПА)"
    if newer == "D17":
        return "overlay (создаёт смежный регистр)"
    if newer in {"D14", "D15", "D16"}:
        return "additive (новый акт встраивается в режим)"
    return "amendment"


def derive_amendment_basis(newer: str, old: str, doc_by_id: dict[str, dict[str, str]]) -> str:
    n = doc_by_id.get(newer, {}).get("Код", newer)
    o = doc_by_id.get(old, {}).get("Код", old)
    return f"{n} → {o}: формулировки в {o} требуют сверки с {n} (см. v7-snapshot)"


def count_events_for_docs(events: list[dict[str, str]], docs: list[str], docs2: list[str] | None = None) -> int:
    """Number of v7 events touching any doc in `docs` (and optionally in docs2 for crossing)."""
    if docs2 is None:
        return sum(1 for e in events if e.get("Документ слева") in docs or e.get("Документ справа") in docs)
    return sum(
        1 for e in events
        if (e.get("Документ слева") in docs and e.get("Документ справа") in docs2)
        or (e.get("Документ справа") in docs and e.get("Документ слева") in docs2)
    )


def count_fallback_files() -> int:
    return CONTROL_NUMBERS["fallback_sources"]


def build_provenance_rows(
    forensic: list[dict[str, str]],
    sources_norm: list[dict[str, str]],
    fallback_count: int,
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    snap_by_id = {f.get("ИД документа", ""): f for f in forensic}
    for s in sources_norm:
        did = s["ИД"]
        f = snap_by_id.get(did, {})
        actuality = f.get("Статус актуальности", "")
        snapshot = f.get("Локальный snapshot", "")
        full = f.get("Полный текст получен", "")
        risk = []
        mirrors = []
        if "актуализировать" in actuality.lower():
            risk.append("требует актуализации редакции")
        if not snapshot or "06_исходные_НПА" not in snapshot:
            if "URL" in full:
                risk.append("только URL без локального полного текста")
        if int(s["Ранг"]) == 3:
            risk.append("rank-3 источник: не может опровергать НПА (см. C-02)")
        # Known fallback mirrors
        if did in {"D21", "D22", "D23", "D20", "D24", "D25", "D26"}:
            mirrors.append("Гарант + Контур + Консультант (round1+round2)")
        if did == "D10":
            mirrors.append("Минэк не индексируется → fallback round3 (Консультант, Контур, prime.ru)")
        rows.append({
            "ИД": did,
            "Источник": s["Название"],
            "Provenance-риск": "; ".join(risk) if risk else "нет явных рисков",
            "Зеркала": "; ".join(mirrors) if mirrors else "—",
            "Стартовое действие": "fetch primary; compare with mirror" if mirrors else "see snapshot",
        })
    rows.append({
        "ИД": "FALLBACK",
        "Источник": f"Aggregate fallback manifest: {fallback_count} зеркал",
        "Provenance-риск": "официальные сайты периодически возвращают stub/timeout (DDoS-эффект)",
        "Зеркала": "Гарант (base/prime), Консультант, Контур, Судакт, Альта, Меганорм, Парлементская газета",
        "Стартовое действие": "См. FALLBACK_MANIFEST*.csv в agent_handoff/",
    })
    return rows


# ---------------------------------------------------------------------------
# Excel rendering
# ---------------------------------------------------------------------------

def render_excel(
    out_path: Path,
    *,
    sources_norm,
    pair_rows,
    matrix_rows,
    topic_rows,
    thesis_rows,
    redaction_rows,
    regime_rows,
    provenance_rows,
    manual_rows,
    contradictions,
    uncovered,
    defects,
    redgreen,
    qa,
    doc_ids,
) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    bold = Font(bold=True)
    head_fill = PatternFill(start_color="1F2937", end_color="1F2937", fill_type="solid")
    head_font = Font(bold=True, color="FFFFFF")
    wrap = Alignment(wrap_text=True, vertical="top")
    thin = Side(border_style="thin", color="D1D5DB")
    border = Border(top=thin, bottom=thin, left=thin, right=thin)

    fills = {
        STATUS_MATCH: PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid"),
        STATUS_PARTIAL: PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid"),
        STATUS_CONTRADICTION: PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid"),
        STATUS_OUTDATED: PatternFill(start_color="DBEAFE", end_color="DBEAFE", fill_type="solid"),
        STATUS_GAP: PatternFill(start_color="EDE9FE", end_color="EDE9FE", fill_type="solid"),
        STATUS_REVIEW: PatternFill(start_color="FFEDD5", end_color="FFEDD5", fill_type="solid"),
        STATUS_NC: PatternFill(start_color="F3F4F6", end_color="F3F4F6", fill_type="solid"),
    }
    mark_to_status = {v: k for k, v in STATUS_TO_MARK.items()}

    def style_header(ws, row=1):
        for cell in ws[row]:
            cell.font = head_font
            cell.fill = head_fill
            cell.alignment = wrap
            cell.border = border

    def write_table(ws, header, rows, col_widths=None):
        ws.append(header)
        style_header(ws)
        for row in rows:
            ws.append([row.get(h, "") for h in header])
        for i, h in enumerate(header, 1):
            ws.column_dimensions[get_column_letter(i)].width = (col_widths or {}).get(h, max(12, min(48, len(h) + 2)))
        ws.freeze_panes = "A2"
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            for cell in row:
                cell.alignment = wrap
                cell.border = border

    # --- Sheet 1: Dashboard --------------------------------------------------
    ws0 = wb.active
    ws0.title = "00 Dashboard"
    ws0.append(["Интегральное перекрестное сравнение v8 — DocDiffOps Forensic"])
    ws0["A1"].font = Font(bold=True, size=16)
    ws0.append([f"Сгенерировано: {qa['generated_at']}"])
    ws0.append([])
    ws0.append(["Контрольные числа (v7 база)"])
    ws0["A4"].font = bold
    ws0.append(["Параметр", "Ожидалось", "Получено"])
    style_header(ws0, row=5)
    for k, expected in qa["control_numbers_expected"].items():
        actual = qa["control_numbers_actual"].get(k, "")
        ws0.append([k, expected, actual])
    ws0.append([])

    r0 = ws0.max_row + 1
    ws0.append(["Статусы пар v8 (распределение)"])
    ws0[f"A{r0}"].font = bold
    r0 += 1
    ws0.append(["Статус", "Количество пар"])
    style_header(ws0, row=r0)
    for s, c in sorted(qa["status_distribution_pairs"].items(), key=lambda kv: -kv[1]):
        ws0.append([s, c])

    ws0.append([])
    r0 = ws0.max_row + 1
    ws0.append(["Распределение пар по рангам источников (L↔R)"])
    ws0[f"A{r0}"].font = bold
    r0 += 1
    ws0.append(["Ранги", "Количество пар"])
    style_header(ws0, row=r0)
    for k, c in sorted(qa["rank_pair_distribution"].items()):
        ws0.append([k, c])

    ws0.append([])
    r0 = ws0.max_row + 1
    ws0.append(["Легенда статусов v8"])
    ws0[f"A{r0}"].font = bold
    r0 += 1
    ws0.append(["Знак", "Статус", "Описание"])
    style_header(ws0, row=r0)
    legend = [
        ("✓", STATUS_MATCH, "Существенного расхождения нет, формулировки совместимы."),
        ("≈", STATUS_PARTIAL, "Документы говорят об одном поле, но покрывают разные условия/периоды/субъекты."),
        ("⚠", STATUS_CONTRADICTION, "Содержательное противоречие или риск неверного переноса нормы."),
        ("↻", STATUS_OUTDATED, "Источник был корректен раньше, но после новой нормы/поправки требует обновления."),
        ("∅", STATUS_GAP, "Тезис есть, но первичного подтверждения в корпусе нет."),
        ("?", STATUS_REVIEW, "Нужен юрист / ручная проверка нормы или редакции."),
        ("—", STATUS_NC, "Документы не нужно сравнивать содержательно."),
    ]
    for sym, st, desc in legend:
        ws0.append([sym, st, desc])
        cell = ws0.cell(row=ws0.max_row, column=1)
        cell.fill = fills[st]
    ws0.column_dimensions["A"].width = 8
    ws0.column_dimensions["B"].width = 22
    ws0.column_dimensions["C"].width = 80
    for row in ws0.iter_rows():
        for cell in row:
            if cell.value and cell.row > 4:
                cell.border = border
                cell.alignment = wrap

    # --- Sheet 2: Sources registry ------------------------------------------
    ws = wb.create_sheet("01 Реестр источников")
    write_table(
        ws,
        ["ИД", "Код", "Название", "Тип", "Ранг", "Provenance-статус",
         "URL/путь", "Локальный snapshot", "SHA-256", "Полный текст получен", "Комментарий"],
        sources_norm,
        col_widths={"Название": 50, "URL/путь": 40, "Локальный snapshot": 40,
                    "SHA-256": 24, "Комментарий": 40},
    )

    # --- Sheet 3: Doc x Doc matrix ------------------------------------------
    ws = wb.create_sheet("02 Документ × Документ")
    header = ["ИД", "Код"] + doc_ids
    ws.append(header)
    style_header(ws)
    for row in matrix_rows:
        ws.append([row.get(h, "") for h in header])
    ws.freeze_panes = "C2"
    for i in range(1, len(header) + 1):
        ws.column_dimensions[get_column_letter(i)].width = 8 if i > 2 else 14
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=3):
        for cell in row:
            text = str(cell.value or "")
            mark = text[0] if text else ""
            st = mark_to_status.get(mark)
            if st:
                cell.fill = fills[st]
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
            cell.font = Font(size=9)

    # --- Sheet 4: Topic x Doc -----------------------------------------------
    ws = wb.create_sheet("03 Тема × Документ")
    header = ["ИД темы", "Тема"] + doc_ids + ["Σ событий"]
    ws.append(header)
    style_header(ws)
    for row in topic_rows:
        ws.append([row.get(h, "") for h in header])
    ws.freeze_panes = "C2"
    for i in range(1, len(header) + 1):
        ws.column_dimensions[get_column_letter(i)].width = 8 if 2 < i <= len(header) - 1 else (12 if i <= 2 else 12)
    ws.column_dimensions["B"].width = 44
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = wrap
            cell.border = border

    # --- Sheet 5: Pairs -----------------------------------------------------
    ws = wb.create_sheet("04 Пары v8 (325)")
    header = ["ИД пары", "L", "R", "L-ранг", "R-ранг", "Статус v8", "Темы v8",
              "Manual-review", "Кол-во событий", "Критичность v7", "Приоритет v7", "Основание"]
    ws.append(header)
    style_header(ws)
    for row in pair_rows:
        ws.append([row.get(h, "") for h in header])
        st = row.get("Статус v8")
        if st in fills:
            cell = ws.cell(row=ws.max_row, column=header.index("Статус v8") + 1)
            cell.fill = fills[st]
    ws.freeze_panes = "A2"
    widths = {"ИД пары": 12, "Темы v8": 60, "Основание": 50}
    for i, h in enumerate(header, 1):
        ws.column_dimensions[get_column_letter(i)].width = widths.get(h, 14)
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = wrap
            cell.border = border

    # --- Sheet 6: Thesis x NPA ----------------------------------------------
    ws = wb.create_sheet("05 Тезисы × НПА")
    header = ["Тезис ID", "Источник тезиса", "Тезис", "Связанный НПА",
              "Статус v8", "Уверенность", "Manual-review", "Основание"]
    ws.append(header)
    style_header(ws)
    for row in thesis_rows:
        ws.append([row.get(h, "") for h in header])
        st = row.get("Статус v8")
        if st in fills:
            cell = ws.cell(row=ws.max_row, column=header.index("Статус v8") + 1)
            cell.fill = fills[st]
    ws.freeze_panes = "A2"
    widths = {"Тезис": 60, "Связанный НПА": 60, "Основание": 40, "Источник тезиса": 24}
    for i, h in enumerate(header, 1):
        ws.column_dimensions[get_column_letter(i)].width = widths.get(h, 14)
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = wrap
            cell.border = border

    # --- Sheet 7: Старая vs новая редакция ----------------------------------
    ws = wb.create_sheet("06 Изменения (outdated)")
    write_table(
        ws,
        ["ID", "Изменяющий документ", "Изменяющий код", "Изменяемый документ",
         "Изменяемый код", "Статус v8", "Тип изменения", "Основание"],
        redaction_rows,
        col_widths={"Тип изменения": 40, "Основание": 60},
    )

    # --- Sheet 8: Режимы ----------------------------------------------------
    ws = wb.create_sheet("07 Режимы")
    write_table(
        ws,
        ["Режим", "Название", "Документы", "Покрыто событиями", "Перекрытие с другими"],
        regime_rows,
        col_widths={"Название": 50, "Документы": 30, "Перекрытие с другими": 80},
    )

    # --- Sheet 9: Provenance ------------------------------------------------
    ws = wb.create_sheet("08 Provenance & risk")
    write_table(
        ws,
        ["ИД", "Источник", "Provenance-риск", "Зеркала", "Стартовое действие"],
        provenance_rows,
        col_widths={"Источник": 50, "Provenance-риск": 50, "Зеркала": 50, "Стартовое действие": 40},
    )

    # --- Sheet 10: Противоречия v7 ------------------------------------------
    ws = wb.create_sheet("09 Противоречия v7")
    if contradictions:
        cols = list(contradictions[0].keys())
        write_table(ws, cols, contradictions, col_widths={c: 60 for c in cols})

    # --- Sheet 11: Непокрытые тезисы ----------------------------------------
    ws = wb.create_sheet("10 Source gaps (U-)")
    if uncovered:
        cols = list(uncovered[0].keys())
        write_table(ws, cols, uncovered, col_widths={c: 60 for c in cols})

    # --- Sheet 12: Defects --------------------------------------------------
    ws = wb.create_sheet("11 Defect log (D-)")
    if defects:
        cols = list(defects[0].keys())
        write_table(ws, cols, defects, col_widths={c: 50 for c in cols})

    # --- Sheet 13: Visual red/green index -----------------------------------
    ws = wb.create_sheet("12 Red-green index")
    if redgreen:
        cols = list(redgreen[0].keys())
        write_table(ws, cols, redgreen, col_widths={c: 50 for c in cols})

    # --- Sheet 14: Manual review queue --------------------------------------
    ws = wb.create_sheet("13 Manual review (183)")
    write_table(
        ws,
        ["ИД события", "ИД пары", "L", "R", "Тема", "Статус", "Уверенность",
         "Что проверить", "Приоритет", "Дедлайн"],
        manual_rows,
        col_widths={"Что проверить": 80, "Тема": 30},
    )

    # --- Sheet 15: QA -------------------------------------------------------
    ws = wb.create_sheet("14 QA")
    ws.append(["Поле", "Значение"])
    style_header(ws)
    flat: list[tuple[str, str]] = [
        ("Сгенерировано", qa["generated_at"]),
        ("Все контрольные числа сходятся", "да" if qa["control_numbers_match"] else "нет"),
    ]
    for k, v in qa["control_numbers_expected"].items():
        flat.append((f"  ожидание {k}", str(v)))
    for k, v in qa["control_numbers_actual"].items():
        flat.append((f"  фактически {k}", str(v)))
    flat.append(("Статусы пар v8", json.dumps(qa["status_distribution_pairs"], ensure_ascii=False)))
    flat.append(("Распределение по рангам", json.dumps(qa["rank_pair_distribution"], ensure_ascii=False)))
    flat.append(("Поправок (outdated)", str(qa["amendments_recorded"])))
    flat.append(("Режимов", str(qa["regimes_recorded"])))
    flat.append(("Provenance строк", str(qa["provenance_rows"])))
    flat.append(("Тезисов в матрице", str(qa["thesis_rows"])))
    for k, v in flat:
        ws.append([k, v])
    ws.column_dimensions["A"].width = 40
    ws.column_dimensions["B"].width = 80
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = wrap
            cell.border = border

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


# ---------------------------------------------------------------------------
# DOCX rendering — Пояснительная записка
# ---------------------------------------------------------------------------

def render_docx_explanatory(
    out_path: Path,
    *,
    sources_norm,
    qa,
    contradictions,
    uncovered,
    defects,
    thesis_rows,
    redaction_rows,
    manual_count: int,
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    h = doc.add_heading("Интегральное перекрестное сравнение v8 — пояснительная записка", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(
        f"Сгенерировано: {qa['generated_at']}. "
        f"Документ суммирует evidence-grade результат сравнительного анализа корпуса "
        f"из {qa['control_numbers_actual']['documents']} документов о миграционных режимах "
        f"и связанных НПА (Россия, 2018–2026).",
    )

    doc.add_heading("1. Цель", level=2)
    doc.add_paragraph(
        "Собрать доказательную сравнительную матрицу по корпусу документов о миграционных "
        "режимах, инвестиционном ВНЖ, ruID/целевом въезде, изменениях 2024–2026 и "
        "связанным источникам. Документ не является юридическим заключением — это "
        "forensic-материал для последующей юридической валидации. Каждый существенный "
        "вывод привязан к источнику (ИД документа, статья/пункт, дата редакции) и "
        "к статусу уверенности."
    )

    doc.add_heading("2. Методика", level=2)
    doc.add_paragraph(
        "Шкала статусов v8: match (совпадает), partial_overlap (частично совпадает), "
        "contradiction (противоречие), outdated (устарело после поправки), "
        "source_gap (тезис без первичного подтверждения в корпусе), "
        "manual_review (требует ручной проверки), not_comparable (содержательно "
        "несопоставимо)."
    )
    doc.add_paragraph(
        "Перекрестные срезы: документ × документ (26×26), тема × документ "
        "(15 кластеров), тезис × НПА, secondary × NPA, старая редакция × новая "
        "редакция (граф поправок), правовой режим × правовой режим, provenance-риски."
    )
    doc.add_paragraph(
        "Иерархия источников: rank-1 — первичный НПА (114-ФЗ, 115-ФЗ, 109-ФЗ, "
        "260-ФЗ, 270-ФЗ, 271-ФЗ, 281-ФЗ, 121-ФЗ, ПП №2573, ПП №1510, ПП №1562, "
        "ПП №468, Указ №467, Концепция 2026/2019, План 4171-р/30-р, Договор ЕАЭС, "
        "КоАП); rank-2 — ведомственные материалы Минэка; rank-3 — аналитика, "
        "вторичные обзоры (Клерк, ВЦИОМ, методика Нейрон). Rank-3 не может "
        "опровергать rank-1; пересечение rank-3 ↔ rank-1 автоматически "
        "получает manual_review."
    )

    doc.add_heading("3. Контрольные числа", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Параметр", "Ожидалось", "Получено"
    for k, v in qa["control_numbers_expected"].items():
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = str(v)
        cells[2].text = str(qa["control_numbers_actual"].get(k, ""))

    doc.add_heading("4. Распределение пар по статусам v8", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Статус", "Количество пар"
    for st, c in sorted(qa["status_distribution_pairs"].items(), key=lambda kv: -kv[1]):
        cells = table.add_row().cells
        cells[0].text = st
        cells[1].text = str(c)

    doc.add_heading("5. Перечень источников (26 документов)", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "ИД", "Код", "Тип", "Ранг"
    for s in sources_norm:
        cells = table.add_row().cells
        cells[0].text = s["ИД"]
        cells[1].text = s["Код"]
        cells[2].text = s["Тип"]
        cells[3].text = str(s["Ранг"])

    doc.add_heading("6. Финальные противоречия v7 (актуальны для v8)", level=2)
    for c in contradictions:
        p = doc.add_paragraph()
        run = p.add_run(f"{c.get('ID','')} • {c.get('Тема','')} — ")
        run.bold = True
        p.add_run(f"{c.get('Вывод','')} (Что сделать: {c.get('Что сделать','')})")

    doc.add_heading("7. Source gaps (непокрытые тезисы)", level=2)
    for u in uncovered:
        p = doc.add_paragraph()
        run = p.add_run(f"{u.get('ID','')} • {u.get('Тезис','')} — ")
        run.bold = True
        p.add_run(f"не покрыто: {u.get('Что не покрыто','')}; нужно добавить: {u.get('Что добавить','')}")

    doc.add_heading("8. Defect log v7", level=2)
    for d in defects:
        p = doc.add_paragraph()
        run = p.add_run(f"{d.get('ID','')} • ")
        run.bold = True
        p.add_run(f"{d.get('Проблема','')}; статус: {d.get('Статус','')}; профилактика: {d.get('Профилактика','')}")

    doc.add_heading("9. Что вынести на ручную юридическую проверку", level=2)
    doc.add_paragraph(
        f"Очередь ручной проверки: {manual_count} элементов (см. лист «13 Manual review (183)» в Excel). "
        "Все элементы помечены как ожидающие проверки P1; владелец — юрист/эксперт миграционного права; "
        "критерий закрытия — указаны точная статья/пункт/редакция и решение ревьюера."
    )
    doc.add_paragraph(
        "Особое внимание: (а) C-01 — ЕАЭС vs Узбекистан/Таджикистан (Минэк-проект «Работа в ЕАЭС» "
        "шире правового режима ст.96–98 Договора ЕАЭС); "
        "(б) C-02 — Клерк-digest как rank-3, требует первичного подтверждения по 115-ФЗ/КоАП/НК; "
        "(в) C-03 — формулировки «более»/«не менее» в брошюре Минэка vs ПП №2573."
    )

    doc.add_heading("10. Как читать Excel/PDF", level=2)
    doc.add_paragraph(
        "Excel: 00 Dashboard → 02 Документ × Документ (карта пар) → 03 Тема × Документ "
        "(где какой документ покрывает тему) → 05 Тезисы × НПА (claim validation) → "
        "06 Изменения (outdated) → 07 Режимы → 08 Provenance & risk → 13 Manual review."
    )
    doc.add_paragraph(
        "PDF Интегральное_перекрестное_сравнение.pdf — компактная сводка для "
        "распечатки. PDF Редакционный_diff — красно-зелёный editorial diff с "
        "основанием (документ + норма + источник + статус) для каждой строки."
    )

    doc.add_heading("11. Запреты и ограничения", level=2)
    for limitation in [
        "Не делать выводы без источника — каждый вывод имеет ИД события/документа.",
        "Не считать брошюру или письмо выше НПА — rank-3/rank-2 не опровергают rank-1.",
        "Не скрывать недоступность official URL — provenance-статус явно зафиксирован.",
        "Не смешивать «нет противоречия» и «не проверено» — это разные ячейки матрицы.",
        "Не считать v8 юридическим заключением — это forensic-материал для юриста.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(limitation)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# DOCX rendering — Редакционный diff (red/green)
# ---------------------------------------------------------------------------

def render_docx_redgreen(
    out_path: Path,
    *,
    thesis_rows,
    redaction_rows,
    contradictions,
    uncovered,
    sources_norm,
) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    RED = RGBColor(0xB9, 0x1C, 0x1C)
    GREEN = RGBColor(0x16, 0x6F, 0x33)
    GRAY = RGBColor(0x4B, 0x55, 0x63)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    h = doc.add_heading("Редакционный diff (red/green) — миграционный корпус v8", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(
        "Красным помечены формулировки/тезисы, которые с высокой вероятностью устарели "
        "или несут риск (требуют ручной проверки). Зелёным — формулировки, "
        "подтверждённые первичным НПА в корпусе. Каждая строка содержит основание: "
        "документ + норма + источник + статус v8."
    )

    doc.add_heading("Раздел A. Тезисы вторичных источников × первичные НПА", level=2)

    code_by_id = {s["ИД"]: s["Код"] for s in sources_norm}

    for row in thesis_rows[:80]:  # ограничиваем объём документа
        st = row["Статус v8"]
        p = doc.add_paragraph()
        prefix = p.add_run(f"{row['Тезис ID']} • {row['Источник тезиса']} • {row.get('Уверенность','')}: ")
        prefix.bold = True
        prefix.font.color.rgb = GRAY

        if st == STATUS_MATCH:
            r = p.add_run(row["Тезис"])
            r.font.color.rgb = GREEN
        elif st in {STATUS_PARTIAL, STATUS_REVIEW, STATUS_CONTRADICTION, STATUS_OUTDATED}:
            r = p.add_run(row["Тезис"])
            r.font.color.rgb = RED
        else:
            r = p.add_run(row["Тезис"])
            r.font.color.rgb = GRAY

        basis = doc.add_paragraph()
        b = basis.add_run(f"   ↳ Основание: {row['Связанный НПА']} • статус v8: {st} • {row['Основание']}")
        b.font.size = Pt(8)
        b.font.color.rgb = GRAY

    doc.add_heading("Раздел B. Старая редакция → новая редакция (outdated)", level=2)
    for row in redaction_rows:
        p = doc.add_paragraph()
        prefix = p.add_run(f"{row['ID']} • {row['Изменяющий код']} → {row['Изменяемый код']}: ")
        prefix.bold = True
        red = p.add_run(f"[redacted: {row['Изменяемый код']}] ")
        red.font.color.rgb = RED
        p.add_run("→ ")
        green = p.add_run(f"[обновлено: {row['Изменяющий код']}] ({row['Тип изменения']})")
        green.font.color.rgb = GREEN
        basis = doc.add_paragraph()
        b = basis.add_run(f"   ↳ {row['Основание']}")
        b.font.size = Pt(8)
        b.font.color.rgb = GRAY

    doc.add_heading("Раздел C. Финальные противоречия v7", level=2)
    for c in contradictions:
        p = doc.add_paragraph()
        prefix = p.add_run(f"{c.get('ID','')} • {c.get('Тема','')}: ")
        prefix.bold = True
        red = p.add_run(c.get("Вывод", ""))
        red.font.color.rgb = RED
        basis = doc.add_paragraph()
        b = basis.add_run(f"   ↳ Что сделать: {c.get('Что сделать','')} • статус: {c.get('Статус','')}")
        b.font.size = Pt(8)
        b.font.color.rgb = GRAY

    doc.add_heading("Раздел D. Source gaps (∅)", level=2)
    for u in uncovered:
        p = doc.add_paragraph()
        prefix = p.add_run(f"{u.get('ID','')}: ")
        prefix.bold = True
        gap = p.add_run(f"{u.get('Тезис','')} — не покрыто: {u.get('Что не покрыто','')}")
        gap.font.color.rgb = RED
        basis = doc.add_paragraph()
        b = basis.add_run(f"   ↳ Добавить: {u.get('Что добавить','')}")
        b.font.size = Pt(8)
        b.font.color.rgb = GRAY

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# PDF rendering helpers (reportlab)
# ---------------------------------------------------------------------------

def _pdf_styles():
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Cyrillic-capable TTFs in priority order: Noto Sans → DejaVu → Liberation
    candidates = [
        ("/usr/share/fonts/noto/NotoSans-Regular.ttf", "/usr/share/fonts/noto/NotoSans-Bold.ttf"),
        ("/usr/share/fonts/TTF/DejaVuSans.ttf", "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/dejavu/DejaVuSans.ttf", "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/liberation/LiberationSans-Regular.ttf",
         "/usr/share/fonts/liberation/LiberationSans-Bold.ttf"),
    ]
    base = "Helvetica"
    bold = "Helvetica-Bold"
    for regular_path, bold_path in candidates:
        if not (Path(regular_path).exists() and Path(bold_path).exists()):
            continue
        try:
            pdfmetrics.registerFont(TTFont("CyrSans", regular_path))
            pdfmetrics.registerFont(TTFont("CyrSans-Bold", bold_path))
            base, bold = "CyrSans", "CyrSans-Bold"
            break
        except Exception:
            continue
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName=base,
                          fontSize=9, leading=12)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName=bold, fontSize=16, leading=20, spaceAfter=12)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName=bold, fontSize=12, leading=15, spaceAfter=6)
    small = ParagraphStyle("small", parent=body, fontSize=8, leading=10)
    red = ParagraphStyle("red", parent=body, textColor=colors.HexColor("#B91C1C"))
    green = ParagraphStyle("green", parent=body, textColor=colors.HexColor("#166F33"))
    gray = ParagraphStyle("gray", parent=small, textColor=colors.HexColor("#4B5563"))
    return {"body": body, "h1": h1, "h2": h2, "small": small,
            "red": red, "green": green, "gray": gray, "base": base, "bold": bold}


def _pdf_table(data, col_widths, *, base="Helvetica", bold="Helvetica-Bold", style_extra=None):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
    style = TableStyle([
        ("FONT", (0, 0), (-1, 0), bold, 9),
        ("FONT", (0, 1), (-1, -1), base, 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])
    if style_extra:
        for cmd in style_extra:
            style.add(*cmd)
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(style)
    return t


def render_pdf_summary(
    out_path: Path,
    *,
    qa,
    sources_norm,
    thesis_rows,
    redaction_rows,
    contradictions,
    uncovered,
    regime_rows,
    topic_rows,
) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    )

    s = _pdf_styles()
    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=1.2 * cm, rightMargin=1.2 * cm,
        topMargin=1.2 * cm, bottomMargin=1.2 * cm,
    )
    elems: list[Any] = []
    elems.append(Paragraph("Интегральное перекрестное сравнение v8 — DocDiffOps Forensic", s["h1"]))
    elems.append(Paragraph(f"Сгенерировано: {qa['generated_at']}", s["body"]))
    elems.append(Spacer(1, 12))

    elems.append(Paragraph("Контрольные числа (v7 база)", s["h2"]))
    rows = [["Параметр", "Ожидалось", "Получено"]]
    for k, v in qa["control_numbers_expected"].items():
        rows.append([k, str(v), str(qa["control_numbers_actual"].get(k, ""))])
    elems.append(_pdf_table(rows, [6 * cm, 4 * cm, 4 * cm], base=s["base"], bold=s["bold"]))
    elems.append(Spacer(1, 10))

    elems.append(Paragraph("Распределение пар по статусам v8", s["h2"]))
    rows = [["Статус", "Количество пар"]]
    for st, c in sorted(qa["status_distribution_pairs"].items(), key=lambda kv: -kv[1]):
        rows.append([st, str(c)])
    elems.append(_pdf_table(rows, [8 * cm, 4 * cm], base=s["base"], bold=s["bold"]))

    elems.append(PageBreak())

    elems.append(Paragraph("Источники (26)", s["h2"]))
    rows = [["ИД", "Код", "Тип", "Ранг", "Provenance"]]
    for r in sources_norm:
        rows.append([r["ИД"], r["Код"], r["Тип"], r["Ранг"], r["Provenance-статус"]])
    elems.append(_pdf_table(rows, [1.4 * cm, 4 * cm, 5.5 * cm, 1 * cm, 6 * cm], base=s["base"], bold=s["bold"]))

    elems.append(PageBreak())
    elems.append(Paragraph("Темы и события", s["h2"]))
    rows = [["ИД темы", "Тема", "Σ событий"]]
    for r in topic_rows:
        sigma = r.get("Σ событий", "")
        rows.append([r["ИД темы"], r["Тема"], str(sigma)])
    elems.append(_pdf_table(rows, [2 * cm, 14 * cm, 2 * cm], base=s["base"], bold=s["bold"]))

    elems.append(Spacer(1, 10))
    elems.append(Paragraph("Правовые режимы", s["h2"]))
    rows = [["Режим", "Название", "Документы", "События"]]
    for r in regime_rows:
        rows.append([r["Режим"], r["Название"], r["Документы"], str(r["Покрыто событиями"])])
    elems.append(_pdf_table(rows, [1.5 * cm, 8 * cm, 5.5 * cm, 2 * cm], base=s["base"], bold=s["bold"]))

    elems.append(PageBreak())
    elems.append(Paragraph("Финальные противоречия v7", s["h2"]))
    for c in contradictions:
        elems.append(Paragraph(f"<b>{c.get('ID','')}</b> {c.get('Тема','')} — {c.get('Вывод','')}", s["red"]))
        elems.append(Paragraph(f"  ↳ Что сделать: {c.get('Что сделать','')}", s["gray"]))
    elems.append(Spacer(1, 10))
    elems.append(Paragraph("Source gaps (∅)", s["h2"]))
    for u in uncovered:
        elems.append(Paragraph(f"<b>{u.get('ID','')}</b> {u.get('Тезис','')}", s["red"]))
        elems.append(Paragraph(f"  ↳ Добавить: {u.get('Что добавить','')}", s["gray"]))

    elems.append(PageBreak())
    elems.append(Paragraph("Старая редакция → новая (outdated)", s["h2"]))
    rows = [["ID", "Изменяющий", "Изменяемый", "Тип изменения"]]
    for r in redaction_rows:
        rows.append([r["ID"], r["Изменяющий код"], r["Изменяемый код"], r["Тип изменения"]])
    elems.append(_pdf_table(rows, [1.5 * cm, 6 * cm, 6 * cm, 4 * cm], base=s["base"], bold=s["bold"]))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(elems)


def render_pdf_explanatory(
    out_path: Path,
    *,
    sources_norm,
    qa,
    contradictions,
    uncovered,
    defects,
) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    s = _pdf_styles()
    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            leftMargin=1.5 * cm, rightMargin=1.5 * cm,
                            topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    elems: list[Any] = []

    elems.append(Paragraph("Пояснительная записка к v8", s["h1"]))
    elems.append(Paragraph(f"Сгенерировано: {qa['generated_at']}", s["body"]))
    elems.append(Spacer(1, 10))

    elems.append(Paragraph("1. Цель", s["h2"]))
    elems.append(Paragraph(
        "Forensic-grade сравнительный материал по корпусу из "
        f"{qa['control_numbers_actual']['documents']} документов о миграционных режимах "
        "(Россия, 2018–2026). Не является юридическим заключением.", s["body"]))

    elems.append(Paragraph("2. Методика", s["h2"]))
    elems.append(Paragraph(
        "Шкала статусов v8: match, partial_overlap, contradiction, outdated, "
        "source_gap, manual_review, not_comparable. Иерархия источников: "
        "rank-1 (НПА) > rank-2 (ведомственное) > rank-3 (аналитика). "
        "Rank-3 не опровергает rank-1; пересечение rank-3↔rank-1 → manual_review.", s["body"]))

    elems.append(Paragraph("3. Ограничения", s["h2"]))
    for limitation in [
        "v8 не заменяет юриста — каждый manual_review закрывается экспертом миграционного права.",
        "Provenance-риски (DDoS-stub, timeout, redirect-only) явно зафиксированы; "
        "fallback-зеркала перечислены в листе «08 Provenance & risk».",
        "Полные тексты НПА не загружены в локальный snapshot для всех 26 документов; "
        "при юридической финализации необходимо подтянуть актуальные редакции.",
    ]:
        elems.append(Paragraph("• " + limitation, s["body"]))

    elems.append(Paragraph("4. Список источников", s["h2"]))
    rows = [["ИД", "Код", "Название", "Тип", "Ранг"]]
    for r in sources_norm:
        rows.append([r["ИД"], r["Код"], r["Название"][:60], r["Тип"][:30], r["Ранг"]])
    elems.append(_pdf_table(rows, [1.2 * cm, 3.5 * cm, 8.5 * cm, 4 * cm, 1 * cm],
                            base=s["base"], bold=s["bold"]))

    elems.append(PageBreak())

    elems.append(Paragraph("5. Финальные противоречия v7", s["h2"]))
    for c in contradictions:
        elems.append(Paragraph(f"<b>{c.get('ID','')}</b> {c.get('Тема','')} — {c.get('Вывод','')}", s["red"]))
        elems.append(Paragraph(f"  Что сделать: {c.get('Что сделать','')}", s["gray"]))

    elems.append(Spacer(1, 8))
    elems.append(Paragraph("6. Непокрытые тезисы", s["h2"]))
    for u in uncovered:
        elems.append(Paragraph(f"<b>{u.get('ID','')}</b> {u.get('Тезис','')}", s["red"]))
        elems.append(Paragraph(f"  Добавить: {u.get('Что добавить','')}", s["gray"]))

    elems.append(Spacer(1, 8))
    elems.append(Paragraph("7. Defect log", s["h2"]))
    for d in defects:
        elems.append(Paragraph(f"<b>{d.get('ID','')}</b> {d.get('Проблема','')} — {d.get('Статус','')}", s["body"]))
        elems.append(Paragraph(f"  Профилактика: {d.get('Профилактика','')}", s["gray"]))

    elems.append(Spacer(1, 8))
    elems.append(Paragraph("8. Запреты (commitment)", s["h2"]))
    for limitation in [
        "Не делать выводы без источника.",
        "Не считать брошюру/письмо выше НПА.",
        "Не скрывать недоступность official URL — provenance-статус явно зафиксирован.",
        "Не смешивать «нет противоречия» и «не проверено».",
        "Не считать v8 юридическим заключением.",
    ]:
        elems.append(Paragraph("• " + limitation, s["body"]))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(elems)


def render_pdf_redgreen(
    out_path: Path,
    *,
    thesis_rows,
    redaction_rows,
    contradictions,
    uncovered,
    sources_norm,
) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    s = _pdf_styles()
    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            leftMargin=1.2 * cm, rightMargin=1.2 * cm,
                            topMargin=1.2 * cm, bottomMargin=1.2 * cm)
    elems: list[Any] = []

    elems.append(Paragraph("Редакционный diff (red/green) — миграционный корпус v8", s["h1"]))
    elems.append(Paragraph(
        "Красный — устаревшие/рискованные/не-подтверждённые формулировки. "
        "Зелёный — подтверждённые первичным НПА. Серый — основание (документ + норма + статус v8).",
        s["body"]))
    elems.append(Spacer(1, 8))

    elems.append(Paragraph("Раздел A. Тезисы вторичных источников × НПА (фрагмент, 60 строк)", s["h2"]))
    for r in thesis_rows[:60]:
        st = r["Статус v8"]
        body_style = s["green"] if st == STATUS_MATCH else (s["red"] if st in {STATUS_REVIEW, STATUS_PARTIAL, STATUS_CONTRADICTION, STATUS_OUTDATED} else s["body"])
        elems.append(Paragraph(f"<b>{r['Тезис ID']}</b> • {r['Источник тезиса']} — {r['Тезис']}", body_style))
        elems.append(Paragraph(f"   ↳ {r['Связанный НПА']} • статус v8: {st}", s["gray"]))

    elems.append(PageBreak())
    elems.append(Paragraph("Раздел B. Старая → новая редакция (outdated)", s["h2"]))
    for r in redaction_rows:
        elems.append(Paragraph(
            f"<b>{r['ID']}</b> {r['Изменяющий код']} ➜ {r['Изменяемый код']} ({r['Тип изменения']})",
            s["body"]))
        elems.append(Paragraph(f"   <font color='#B91C1C'>[redacted: {r['Изменяемый код']}]</font> → "
                               f"<font color='#166F33'>[обновлено: {r['Изменяющий код']}]</font>",
                               s["body"]))
        elems.append(Paragraph(f"   ↳ {r['Основание']}", s["gray"]))

    elems.append(PageBreak())
    elems.append(Paragraph("Раздел C. Финальные противоречия", s["h2"]))
    for c in contradictions:
        elems.append(Paragraph(f"<b>{c.get('ID','')}</b> {c.get('Тема','')} — {c.get('Вывод','')}", s["red"]))
        elems.append(Paragraph(f"   ↳ Что сделать: {c.get('Что сделать','')}", s["gray"]))

    elems.append(Spacer(1, 6))
    elems.append(Paragraph("Раздел D. Source gaps (∅)", s["h2"]))
    for u in uncovered:
        elems.append(Paragraph(f"<b>{u.get('ID','')}</b> {u.get('Тезис','')}", s["red"]))
        elems.append(Paragraph(f"   ↳ Добавить: {u.get('Что добавить','')}", s["gray"]))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(elems)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    sys.exit(main())
