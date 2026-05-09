#!/usr/bin/env python3
"""v8.3 — sign-off form + cross-reference index + JSON Schema embed + updated NAVIGATION.

Outputs:
  docs/Лист_согласования.docx                — signature table per FA action
  docs/Лист_согласования.pdf                 — printable
  data/18_doc_xref.csv                       — for each D01..D26: pairs/themes/actions/amendments
  data/v8_bundle.schema.json                 — exported JSON Schema (json-schema draft-07)
  logs/qa_v8_3.json
"""
from __future__ import annotations

import csv
import datetime as dt
import json
import sys
from pathlib import Path
from typing import Any

ROOT = Path("/home/dev/diff/migration_v8_out")
DATA = ROOT / "data"
DOCS = ROOT / "docs"
LOGS = ROOT / "logs"

GENERATED_AT = dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%d %H:%M:%SZ")

sys.path.insert(0, "/home/dev/diff/docdiffops_mvp")
from docdiffops.forensic_actions import (  # noqa: E402
    DEFAULT_ACTIONS, raci_for_action, actions_for_pair,
)
from docdiffops.forensic_schema import BUNDLE_SCHEMA_DICT  # noqa: E402


# ---------------------------------------------------------------------------
# 1. Export JSON Schema as a standalone artifact
# ---------------------------------------------------------------------------

def export_schema() -> Path:
    p = DATA / "v8_bundle.schema.json"
    p.write_text(json.dumps(BUNDLE_SCHEMA_DICT, indent=2, ensure_ascii=False),
                 encoding="utf-8")
    return p


# ---------------------------------------------------------------------------
# 2. Build cross-reference index per document
# ---------------------------------------------------------------------------

def _read_csv(path: Path) -> list[dict[str, str]]:
    return list(csv.DictReader(path.open(encoding="utf-8-sig"), delimiter=";"))


def build_doc_xref() -> Path:
    sources = _read_csv(DATA / "01_источники_v8.csv")
    pairs = _read_csv(DATA / "02_pairs_v8.csv")
    amendments = _read_csv(DATA / "06_old_vs_new_redactions.csv")

    rows: list[dict[str, Any]] = []
    for s in sources:
        did = s["ИД"]
        # Pairs participated in
        my_pairs = [p for p in pairs if p["L"] == did or p["R"] == did]
        # By status
        by_status: dict[str, int] = {}
        for p in my_pairs:
            by_status[p["Статус v8"]] = by_status.get(p["Статус v8"], 0) + 1
        # Themes touched (collapse all topic strings from my_pairs)
        themes_set: set[str] = set()
        for p in my_pairs:
            for t in (p.get("Темы v8") or "").split("; "):
                if t:
                    themes_set.add(t)
        # Actions where this doc is referenced
        my_actions = [a for a in DEFAULT_ACTIONS if did in a.related_docs]
        # Amendment links
        amend_in = [r for r in amendments
                    if r["Изменяющий документ"] == did or r["Изменяемый документ"] == did]
        rows.append({
            "ИД": did,
            "Код": s["Код"],
            "Тип": s["Тип"],
            "Ранг": s["Ранг"],
            "Пар (всего)": len(my_pairs),
            "Пар match": by_status.get("match", 0),
            "Пар partial_overlap": by_status.get("partial_overlap", 0),
            "Пар manual_review": by_status.get("manual_review", 0),
            "Пар contradiction": by_status.get("contradiction", 0),
            "Пар not_comparable": by_status.get("not_comparable", 0),
            "Темы (число)": len(themes_set),
            "Темы (список)": "; ".join(sorted(themes_set)),
            "FA-actions": ", ".join(a.id for a in my_actions),
            "Изменяет / изменяется": ", ".join(
                f"{r['Изменяющий код']}→{r['Изменяемый код']}" for r in amend_in
            ),
        })

    p = DATA / "18_doc_xref.csv"
    with p.open("w", encoding="utf-8-sig", newline="") as fh:
        w = csv.DictWriter(
            fh, delimiter=";",
            fieldnames=["ИД", "Код", "Тип", "Ранг",
                        "Пар (всего)", "Пар match", "Пар partial_overlap",
                        "Пар manual_review", "Пар contradiction", "Пар not_comparable",
                        "Темы (число)", "Темы (список)",
                        "FA-actions", "Изменяет / изменяется"],
        )
        w.writeheader()
        for row in rows:
            w.writerow(row)
    return p


# ---------------------------------------------------------------------------
# 3. Sign-off form DOCX + PDF
# ---------------------------------------------------------------------------

def render_signoff_docx() -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    GRAY = RGBColor(0x4B, 0x55, 0x63)
    RED = RGBColor(0xB9, 0x1C, 0x1C)

    doc.add_heading("Лист согласования — Forensic v8 пакет", level=1)
    doc.add_paragraph(
        f"Корпус: миграционный (РФ, 2018–2026). Источников: 26. Пар: 325. Событий: 281. "
        f"Генерация: {GENERATED_AT}."
    )
    doc.add_paragraph(
        "Цель документа: зафиксировать согласование каждого выявленного несоответствия "
        "(FA-01..FA-10) ответственными лицами по матрице RACI. После подписания вернуть "
        "сканированную копию владельцу пакета."
    )

    p = doc.add_paragraph()
    warn = p.add_run(
        "ВАЖНО: пакет не является юридическим заключением. Каждый FA должен пройти "
        "проверку Accountable-юристом, прежде чем правка применится в продакшен-материалах."
    )
    warn.font.color.rgb = RED
    warn.bold = True

    doc.add_heading("Таблица согласования", level=2)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate([
        "FA", "Серьёзность", "R (исполнитель)", "A (подписывает)",
        "Решение", "Подпись", "Дата", "Комментарий"
    ]):
        hdr[i].text = h

    for a in DEFAULT_ACTIONS:
        raci = raci_for_action(a.id)
        cells = table.add_row().cells
        cells[0].text = a.id
        cells[1].text = a.severity
        cells[2].text = raci["R"]
        cells[3].text = raci["A"]
        cells[4].text = "□ APPROVED  □ NEEDS FIX  □ REJECTED"
        cells[5].text = ""  # подпись
        cells[6].text = ""  # дата
        cells[7].text = ""  # комментарий

    doc.add_heading("Подтверждение владельца пакета", level=2)
    doc.add_paragraph("Я подтверждаю, что:")
    for line in [
        "проверил(а) применимость каждого FA к актуальной редакции корпуса;",
        "получил(а) подписи всех Accountable-лиц до распространения исправлений;",
        "сохранил(а) provenance-цепочки и фиксации mirror-источников;",
        "учёл(а) source gaps U-01..U-03 как незакрытые до отдельного решения;",
        "обеспечил(а), что rank-3 источники (D01, D02, D03, D09) не используются для "
        "опровержения первичных НПА.",
    ]:
        p = doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph()
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.style = "Light Grid"
    hdr = sig_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Должность", "ФИО", "Подпись / дата"
    sig_table.rows[1].cells[0].text = "Владелец пакета"

    p = doc.add_paragraph()
    p.add_run(
        "Контактное лицо для уточнений: укажите ответственного за корпус (имя, email, телефон)."
    ).font.color.rgb = GRAY

    out = DOCS / "Лист_согласования.docx"
    doc.save(str(out))
    return out


def render_signoff_pdf() -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (PageBreak, Paragraph, SimpleDocTemplate,
                                     Spacer, Table, TableStyle)

    candidates = [
        ("/usr/share/fonts/noto/NotoSans-Regular.ttf",
         "/usr/share/fonts/noto/NotoSans-Bold.ttf"),
        ("/usr/share/fonts/liberation/LiberationSans-Regular.ttf",
         "/usr/share/fonts/liberation/LiberationSans-Bold.ttf"),
    ]
    base, bold = "Helvetica", "Helvetica-Bold"
    for r_path, b_path in candidates:
        if Path(r_path).exists() and Path(b_path).exists():
            try:
                pdfmetrics.registerFont(TTFont("V8Sans", r_path))
                pdfmetrics.registerFont(TTFont("V8Sans-Bold", b_path))
                base, bold = "V8Sans", "V8Sans-Bold"
                break
            except Exception:
                continue

    out = DOCS / "Лист_согласования.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=1.2 * cm, rightMargin=1.2 * cm,
                            topMargin=1.2 * cm, bottomMargin=1.2 * cm)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName=base, fontSize=9, leading=12)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName=bold, fontSize=18, leading=22, spaceAfter=10)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName=bold, fontSize=12, leading=15, spaceAfter=6)
    red = ParagraphStyle("red", parent=body, textColor=colors.HexColor("#B91C1C"), fontName=bold)
    elems: list[Any] = []
    elems.append(Paragraph("Лист согласования — Forensic v8 пакет", h1))
    elems.append(Paragraph(
        f"Источников: 26. Пар: 325. Событий: 281. Сгенерировано: {GENERATED_AT}.",
        body))
    elems.append(Paragraph(
        "Не является юридическим заключением. Каждый FA должен пройти проверку "
        "Accountable-юристом перед применением.", red))
    elems.append(Spacer(1, 8))

    elems.append(Paragraph("Таблица согласования FA-01..FA-10", h2))
    rows = [["FA", "Sev", "R", "A", "Решение", "Подпись/дата"]]
    for a in DEFAULT_ACTIONS:
        raci = raci_for_action(a.id)
        rows.append([
            a.id, a.severity,
            raci["R"][:30], raci["A"][:30],
            "□ APPROVED  □ NEEDS FIX", " ",
        ])
    t = Table(rows, colWidths=[1.4 * cm, 1.4 * cm, 4.5 * cm, 4.5 * cm, 4 * cm, 3 * cm], repeatRows=1)
    t.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), bold, 9),
        ("FONT", (0, 1), (-1, -1), base, 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.HexColor("#FFFFFF"), colors.HexColor("#F9FAFB")]),
    ]))
    elems.append(t)

    elems.append(PageBreak())
    elems.append(Paragraph("Подтверждение владельца пакета", h2))
    elems.append(Paragraph("Я подтверждаю, что:", body))
    for line in [
        "проверил(а) применимость каждого FA к актуальной редакции корпуса;",
        "получил(а) подписи всех Accountable-лиц до распространения исправлений;",
        "сохранил(а) provenance-цепочки и фиксации mirror-источников;",
        "учёл(а) source gaps U-01..U-03 как незакрытые;",
        "обеспечил(а), что rank-3 источники не используются против rank-1.",
    ]:
        elems.append(Paragraph("• " + line, body))

    elems.append(Spacer(1, 14))
    sig_rows = [
        ["Должность", "ФИО", "Подпись / дата"],
        ["Владелец пакета", " ", " "],
        ["Юрист (главный)", " ", " "],
        ["Главный аналитик", " ", " "],
    ]
    sig_t = Table(sig_rows, colWidths=[5 * cm, 6 * cm, 6 * cm])
    sig_t.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), bold, 9),
        ("FONT", (0, 1), (-1, -1), base, 9),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
        ("MINIMUMHEIGHT", (0, 1), (-1, -1), 1.5 * cm),
    ]))
    elems.append(sig_t)

    doc.build(elems)
    return out


# ---------------------------------------------------------------------------
# 4. NAVIGATION.md update
# ---------------------------------------------------------------------------

def update_navigation() -> Path:
    nav = ROOT / "NAVIGATION.md"
    content = f"""# Forensic v8 — Навигатор пакета (v8.3)

**Корпус:** миграционный, Россия, 2018–2026. **Сгенерировано:** {GENERATED_AT}
**Источник данных:** `/home/dev/diff/migration_v7_evidence/` (read-only)
**Этот пакет:** `/home/dev/diff/migration_v8_out/` (writable)

## Контрольные числа (все сходятся ✓)

| Параметр | Значение |
|---|---|
| Документов | 26 |
| Пар | 325 |
| Событий | 281 |
| Очередь ручной проверки | 183 |
| Финальные противоречия v7 | 3 |
| Source gaps v7 | 3 |
| Defect log v7 | 3 |
| Fallback-зеркал | 71 |
| FA-actions | 10 |
| Brochure red/green правок | 6 |
| Klerk → НПА footnotes | 6 |
| ЕАЭС split групп | 3 |
| Amendment chains | 5 |
| RACI ролей × FA | 40 |

## Релиз-таймлайн

| Версия | Дата | Что добавлено |
|---|---|---|
| **v8.0** | base | 15-листовой Excel + explanatory DOCX/PDF + red/green DOCX/PDF + 9 CSV + JSON bundle |
| **v8.1** | actionable | FA-01..FA-10 каталог; brochure red/green; Klerk→NPA links; ЕАЭС split; amendment chain; provenance actions; top-20 priority; Что делать DOCX/PDF; Несоответствия XLSX |
| **v8.2** | visual + system | 5 PNG visuals (heatmap, pie, bars, cover); cover PDF; RACI matrix CSV + Excel sheet; system module `docdiffops/forensic_actions.py` (16 tests); offline CLI `forensic_cli rebuild` |
| **v8.3** | schema + signoff | Formal JSON Schema (draft-07, 9 tests); doc xref index; sign-off form DOCX/PDF; updated NAVIGATION |

## Основные документы (читать сначала)

| Файл | Что внутри | Размер |
|---|---|---|
| **`docs/Forensic_v8_cover.pdf`** | Обложка с heatmap + pie + контрольные числа + RACI | ~480 KB |
| **`docs/Что_делать.docx`** + `.pdf` | План действий: FA-01..FA-10 с координатами и владельцами | 45/52 KB |
| **`docs/Лист_согласования.docx`** + `.pdf` | Sign-off form для юриста | new in v8.3 |
| **`docs/Интегральное_перекрестное_сравнение.xlsx`** | 15-листовой основной workbook | 89 KB |
| **`docs/Несоответствия_и_действия.xlsx`** | 9-листовой supplementary с heatmap + RACI | 152 KB |
| **`docs/Пояснительная_записка.docx`** + `.pdf` | Методика, ограничения, источники | 41/43 KB |
| **`docs/Редакционный_diff.docx`** + `.pdf` | Red/green editorial diff с основанием | 42/52 KB |

## Визуальный слой

| Файл | Что показывает |
|---|---|
| `docs/visuals/heatmap_doc_x_doc.png` | 26×26 цветная карта статусов |
| `docs/visuals/cover_summary.png` | Композит: pie + ranks + контрольные + top FA |
| `docs/visuals/topic_bar.png` | Bar-chart покрытия 17 тем |
| `docs/visuals/status_pie.png` | Распределение пар по 5 статусам |
| `docs/visuals/rank_pair_bar.png` | Распределение по rank-pair |

## Машинно-читаемые данные (`data/`)

| CSV/JSON | Что |
|---|---|
| `01_источники_v8.csv` | 26 источников + provenance статус |
| `02_pairs_v8.csv` | 325 пар с агрегированным v8-статусом и основанием |
| `03_doc_x_doc_matrix.csv` | 26×26 символьная матрица |
| `04_topic_x_doc.csv` | 17 кластеров × 26 документов |
| `05_thesis_x_npa.csv` | 128 тезисов вторичных vs НПА |
| `06_old_vs_new_redactions.csv` | 12 amendment-связок |
| `07_regime_x_regime.csv` | 8 правовых режимов |
| `08_provenance_risk.csv` | 27 строк + fallback aggregate |
| `09_manual_review_queue.csv` | 183 ручных проверки |
| `10_actions_catalogue.csv` | 10 FA-actions с WHERE/WHAT/FIX |
| `11_brochure_redgreen_diff.csv` | 6 cells: «более» → «не менее» |
| `12_klerk_npa_links.csv` | 6 footnotes |
| `13_eaeu_split.csv` | 3 группы (ЕАЭС / безвиз-патент / визовые-разрешение) |
| `14_amendment_chain.csv` | 5 цепочек поправок |
| `15_provenance_actions.csv` | 4 fallback-плана |
| `16_top_priority_review.csv` | 20 пар, ranked, с дедлайнами |
| `17_raci_matrix.csv` | RACI × FA-01..FA-10 |
| `18_doc_xref.csv` | new v8.3: для каждого D01..D26 — пары, темы, FA, поправки |
| `integral_cross_comparison.json` | Полный JSON snapshot (schema v8.0) |
| `v8_bundle.schema.json` | new v8.3: формальная JSON Schema (draft-07) |

## Системный код DocDiffOps (`/home/dev/diff/docdiffops_mvp/docdiffops/`)

| Модуль | Что делает | Тестов |
|---|---|---|
| `forensic.py` | V8_STATUSES, aggregate_pair_status_v8, build_forensic_bundle | 27 |
| `forensic_render.py` | render_v8_xlsx, render_v8_docx_explanatory/_redgreen, render_v8_pdf_summary | 4 |
| `forensic_actions.py` | DEFAULT_ACTIONS (10), apply_actions_to_bundle, raci_for_action | 16 |
| `forensic_schema.py` | BUNDLE_SCHEMA_DICT, validate_bundle | 9 |
| `forensic_cli.py` | `python -m docdiffops.forensic_cli rebuild bundle.json --out dir/` | (CLI) |
| `pipeline.py` | hook в render_global_reports → 5 артефактов | 6 |
| `main.py` | API: GET /batches/{{id}}/forensic[/{{kind}}] | — |
| **Итого** | | **62 forensic tests** |

## Шкала статусов v8

- `match` ✓ — совпадает
- `partial_overlap` ≈ — частично совпадает
- `contradiction` ⚠ — противоречие
- `outdated` ↻ — устарело после поправки
- `source_gap` ∅ — тезис без первичного подтверждения
- `manual_review` ? — нужен юрист
- `not_comparable` — — содержательно несопоставимо

## Иерархия источников (invariant)

- **rank-1** — первичный НПА (114-ФЗ, 115-ФЗ, 109-ФЗ, 260-ФЗ, 270-ФЗ, 271-ФЗ, 281-ФЗ, 121-ФЗ, ПП 2573, ПП 1510, ПП 1562, ПП 468, Указ 467, Концепция 2026/2019, План 4171-р/30-р, КоАП, Договор ЕАЭС, НК)
- **rank-2** — ведомственные материалы Минэка (D10, D18)
- **rank-3** — аналитика (D01 Нейрон, D02 пример выгрузки, D03 ВЦИОМ, D09 Клерк)
- **Правило C-02:** rank-3 не опровергает rank-1; пересечение rank-3 ↔ rank-1 → `manual_review`.

## CLI: offline rebuild

```bash
python -m docdiffops.forensic_cli rebuild bundle.json --out dir/ --with-actions
# rebuilds 5 artifacts under dir/
```

## Re-build всего пакета

```bash
# v8.0 базовый пакет
.venv/bin/python /home/dev/diff/migration_v8_out/scripts/build_integral.py

# v8.1 enhancement (FA-каталог, brochure, Klerk, ЕАЭС, amendments)
.venv/bin/python /home/dev/diff/migration_v8_out/scripts/enhance_v8.py

# v8.2 visuals + RACI + cover PDF
.venv/bin/python /home/dev/diff/migration_v8_out/scripts/build_visuals.py
.venv/bin/python /home/dev/diff/migration_v8_out/scripts/build_v8_2.py

# v8.3 schema + signoff + doc xref + NAVIGATION
.venv/bin/python /home/dev/diff/migration_v8_out/scripts/build_v8_3.py
```

## Ключевые риски — короткое резюме

- **C-01 / FA-02** Минэк-проект «Работа в ЕАЭС» включает Узбекистан/Таджикистан — НЕ члены ЕАЭС, должны быть в группе «патент» по 115-ФЗ.
- **C-03 / FA-01 / BR-01..BR-06** Брошюра Минэка использует «более X», ПП №2573 — «не менее X». Инвестор с пороговой суммой выпадает.
- **C-02 / FA-03 / KL-01..KL-06** Клерк (D09, rank-3) даёт 6 фактов без footnote-ссылок на НПА.
- **FA-04** Концепция 2019–2025 цитируется без отметки «утратила силу с D04».
- **FA-05 / AC-01** ПП №1510 цитируется без редакции 1562/468.
- **FA-06 / AC-02** 115/109-ФЗ без отметки «в ред. 270-ФЗ».
- **FA-07 / AC-03** КоАП ст.18.x без «в ред. 281-ФЗ».
- **FA-08** ВЦИОМ-claims смешаны с НПА.
- **FA-09 / PV-01..PV-04** Provenance: consultant.ru/pravo.gov.ru/mvd.ru нестабильны → fallback-mirrors.
- **FA-10 / U-01..U-03** Source gaps: внешние юрисдикции, статьи НК, постатейные планы.
"""
    nav.write_text(content, encoding="utf-8")
    return nav


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    schema_path = export_schema()
    print(f"  ✓ {schema_path.relative_to(ROOT)}: {schema_path.stat().st_size:,}b")

    xref = build_doc_xref()
    print(f"  ✓ {xref.relative_to(ROOT)}: {xref.stat().st_size:,}b")

    so_docx = render_signoff_docx()
    print(f"  ✓ {so_docx.relative_to(ROOT)}: {so_docx.stat().st_size:,}b")

    so_pdf = render_signoff_pdf()
    print(f"  ✓ {so_pdf.relative_to(ROOT)}: {so_pdf.stat().st_size:,}b")

    nav = update_navigation()
    print(f"  ✓ {nav.relative_to(ROOT)}: {nav.stat().st_size:,}b")

    qa = {
        "generated_at": GENERATED_AT,
        "schema": "v8.3",
        "added_artifacts": {
            "v8_bundle.schema.json": str(schema_path.relative_to(ROOT)),
            "doc_xref_csv": str(xref.relative_to(ROOT)),
            "signoff_docx": str(so_docx.relative_to(ROOT)),
            "signoff_pdf": str(so_pdf.relative_to(ROOT)),
            "navigation_md": str(nav.relative_to(ROOT)),
        },
    }
    (LOGS / "qa_v8_3.json").write_text(
        json.dumps(qa, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
