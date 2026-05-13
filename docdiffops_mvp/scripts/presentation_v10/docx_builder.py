"""Build a native Word DOCX companion to the DocDiffOps v10 PPTX presentation.

Same storyline, same data — but as a readable text document (A4 portrait).
Run as:
    cd docdiffops_mvp
    python -m scripts.presentation_v10.docx_builder
"""
from __future__ import annotations

import sys
from collections import Counter
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Make ``docdiffops`` importable when invoked from docdiffops_mvp/
# ---------------------------------------------------------------------------
_REPO_PKG = Path(__file__).resolve().parents[2]
if str(_REPO_PKG) not in sys.path:
    sys.path.insert(0, str(_REPO_PKG))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .data_loader import V10Data, load_data
from .theme import OCEAN, STATUS_RU, STATUS_TINT_BG, V8_STATUSES

REPO_ROOT = Path(__file__).resolve().parents[3]
ASSETS_DIR = REPO_ROOT / "migration_v10_out" / "presentation" / "assets"
DEFAULT_OUT = REPO_ROOT / "migration_v10_out" / "presentation" / "DocDiffOps_v10_presentation.docx"

# ---------------------------------------------------------------------------
# Category / severity translations (mirror slides_part4.py)
# ---------------------------------------------------------------------------
CAT_RU: dict[str, str] = {
    "brochure_vs_npa":        "брошюра против НПА",
    "department_page_split":  "раздробление ведомственной страницы",
    "secondary_digest_links": "сноски на первичные НПА",
    "concept_supersession":   "замещение концепций",
    "amendment_chain":        "цепочка поправок",
    "amendment_to_law":       "поправки к закону",
    "amendment_to_koap":      "поправки к КоАП",
    "analytic_separation":    "разделение аналитики и НПА",
    "provenance_risk":        "риск provenance",
    "source_gap":             "пробел источника",
}
SEV_RU: dict[str, str] = {
    "high": "высокая",
    "medium": "средняя",
    "low": "низкая",
}
# Severity → STATUS_TINT_BG key for cell color
SEV_STATUS_KEY: dict[str, str] = {
    "high": "contradiction",
    "medium": "partial_overlap",
    "low": "match",
}
# Priority → STATUS_TINT_BG key
PRIO_STATUS_KEY: dict[str, str] = {
    "P0": "contradiction",
    "P1": "partial_overlap",
    "P2": "not_comparable",
}


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _trunc(text: str, n: int) -> str:
    if len(text) <= n:
        return text
    return text[: n - 1] + "…"


def _set_cell_bg(cell: Any, hex_color: str) -> None:
    """Set cell shading fill (mirrors forensic_render._docx_set_cell_bg)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _repeat_table_header(table: Any) -> None:
    """Mark the first row as a repeating header row across pages."""
    tr = table.rows[0]._tr
    tr_pr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tr_pr.append(tbl_header)


def _set_default_font(doc: Document, font_name: str, size_pt: int = 11) -> None:
    """Apply default font to Normal style."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)


def _set_a4_portrait(doc: Document) -> None:
    """Set A4 portrait with 2 cm margins."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def _add_page_field(para: Any) -> None:
    """Insert Word PAGE / NUMPAGES field pair into paragraph."""
    run = para.add_run()
    # PAGE
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)

    run2 = para.add_run(" из ")
    run2.font.name = "DejaVu Sans"

    run3 = para.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    run3._r.append(fldChar3)
    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = " NUMPAGES "
    run3._r.append(instrText2)
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar4)


def _add_toc_field(doc: Document) -> None:
    """Insert a Word TOC field (levels 1-3) that auto-populates on open."""
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)


def _add_footer(doc: Document) -> None:
    """Add footer with doc title + date + page number."""
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run("DocDiffOps v10 · 2026-05-09 · стр. ")
    run.font.name = "DejaVu Sans"
    run.font.size = Pt(9)
    run.font.color.rgb = _hex_to_rgb(OCEAN["muted"])
    _add_page_field(fp)
    for r in fp.runs:
        r.font.name = "DejaVu Sans"
        r.font.size = Pt(9)


def _heading(doc: Document, text: str, level: int) -> Any:
    """Add heading with Ocean brand colors and DejaVu Sans font."""
    para = doc.add_heading(text, level=level)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.name = "DejaVu Sans"
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = _hex_to_rgb(OCEAN["primary"])
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = _hex_to_rgb(OCEAN["secondary"])
        run.font.bold = True
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = _hex_to_rgb(OCEAN["accent"])
        run.font.bold = True
    return para


def _para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
          size_pt: int = 11, color_hex: str | None = None) -> Any:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "DejaVu Sans"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = _hex_to_rgb(color_hex)
    return para


def _bullet(doc: Document, text: str, *, size_pt: int = 11) -> Any:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "DejaVu Sans"
    run.font.size = Pt(size_pt)
    return para


def _table_header_row(table: Any, headers: list[str],
                      bg_hex: str = "065A82") -> None:
    """Style the first row as a header row."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        _set_cell_bg(cell, bg_hex)
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(header)
        run.font.name = "DejaVu Sans"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)


def _table_body_row(table: Any, row_idx: int, values: list[str],
                    status_col: int | None = None,
                    status_map: dict[str, str] | None = None) -> None:
    """Populate a data row; optionally colour a status column."""
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(val)
        run.font.name = "DejaVu Sans"
        run.font.size = Pt(8)
        if status_col is not None and i == status_col and status_map:
            # Try to find bg color; val may be RU label or eng key
            bg = status_map.get(val)
            if not bg:
                # Reverse lookup: RU label → eng key
                ru_to_eng = {v: k for k, v in STATUS_RU.items()}
                eng = ru_to_eng.get(val)
                if eng:
                    bg = status_map.get(eng)
            if bg:
                _set_cell_bg(cell, bg)


def _add_image(doc: Document, img_path: Path, width_cm: float = 15.0) -> bool:
    """Add image paragraph centered; return True if added."""
    if not img_path.exists():
        _para(doc, f"[Изображение недоступно: {img_path.name}]",
              italic=True, color_hex=OCEAN["muted"])
        return False
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    return True


def _page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Chapter builders
# ---------------------------------------------------------------------------

def _build_cover(doc: Document, data: V10Data) -> None:
    cn = data.control_numbers
    doc_count = cn.get("documents", 27)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("DocDiffOps v10 — Сквозная презентация")
    run.font.name = "DejaVu Sans"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(OCEAN["primary"])

    doc.add_paragraph()  # spacer

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        "Криминалистическое "
        "сравнение корпуса "
        "нормативных и "
        "аналитических "
        "документов "
        "миграционной "
        "политики РФ"
    )
    sub_run.font.name = "DejaVu Sans"
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True

    doc.add_paragraph()

    # KPI summary
    kpi_para = doc.add_paragraph()
    kpi_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kpi_run = kpi_para.add_run(
        f"{doc_count} документов"
        f" · {cn.get('pairs', 351)} пар"
        f" · {cn.get('events', 312)} событий"
    )
    kpi_run.font.name = "DejaVu Sans"
    kpi_run.font.size = Pt(13)
    kpi_run.font.bold = True
    kpi_run.font.color.rgb = _hex_to_rgb(OCEAN["secondary"])

    doc.add_paragraph()

    # Footer line
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "Версия 10.0.0 · 2026-05-09 · QA-гейт 12/12 PASS"
    )
    footer_run.font.name = "DejaVu Sans"
    footer_run.font.size = Pt(11)
    footer_run.font.color.rgb = _hex_to_rgb(OCEAN["muted"])

    _page_break(doc)


def _build_toc(doc: Document, data: V10Data) -> None:
    _heading(doc, "Содержание", 1)
    _add_toc_field(doc)
    _page_break(doc)


def _build_executive(doc: Document, data: V10Data) -> None:
    _heading(doc, "Executive Summary", 1)

    # 2.1 Corpus & events KPI
    _heading(doc, "2.1 Корпус и события", 2)
    cn = data.control_numbers
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.style = "Table Grid"
    kpi_headers = ["27 документов",
                   "351 пара",
                   "312 diff-событий",
                   "12/12 QA PASS"]
    kpi_labels = ["Документов",
                  "Пар сравнения",
                  "Diff-событий",
                  "Гейт QA"]
    for i, (val, lbl) in enumerate(zip(kpi_headers, kpi_labels)):
        c_val = kpi_table.rows[0].cells[i]
        c_lbl = kpi_table.rows[1].cells[i]
        c_val.text = val
        c_lbl.text = lbl
        _set_cell_bg(c_val, OCEAN["tile_bg"])
        for cell in (c_val, c_lbl):
            para = cell.paragraphs[0]
            if para.runs:
                para.runs[0].font.name = "DejaVu Sans"
                para.runs[0].font.size = Pt(11)
                para.runs[0].font.bold = True

    doc.add_paragraph()

    # 2.2 Pair status distribution
    _heading(doc, "2.2 Распределение пар по статусам", 2)
    pairs_dist = data.pairs_by_status()
    total_pairs = sum(pairs_dist.values()) or 351
    status_table = doc.add_table(rows=1 + len(V8_STATUSES), cols=4)
    status_table.style = "Table Grid"
    _table_header_row(status_table,
                      ["Статус (eng)",
                       "Статус (RU)",
                       "Количество",
                       "%"])
    _repeat_table_header(status_table)
    for ri, status in enumerate(V8_STATUSES):
        count = pairs_dist.get(status, 0)
        pct = f"{count / total_pairs * 100:.1f}%"
        row = status_table.rows[ri + 1]
        row.cells[0].text = status
        row.cells[1].text = STATUS_RU.get(status, status)
        row.cells[2].text = str(count)
        row.cells[3].text = pct
        bg = STATUS_TINT_BG.get(status)
        if bg:
            _set_cell_bg(row.cells[1], bg)
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()
    _para(doc,
          "Из 351 пары: 86 на "
          "ручную проверку, "
          "202 несопоставимы, "
          "53 частичные, "
          "8 устаревших, "
          "1 совпадение, "
          "1 противоречие.")
    _add_image(doc, ASSETS_DIR / "chart_status_pie.png", width_cm=15.0)
    doc.add_paragraph()

    # 2.3 Trend
    _heading(doc, "2.3 Тренд по итерациям", 2)
    _para(doc,
          "Падение match_share с 8% (v8) "
          "до 0.28% (v9) — расширение "
          "корпуса, не регресс.")
    _add_image(doc, ASSETS_DIR / "chart_trend_match_share.png", width_cm=15.0)
    doc.add_paragraph()

    # 2.4 Methodology
    _heading(doc, "2.4 Методология: ранговый шлюз", 2)
    meth_bullets = [
        "Все события классифицируются по 7-балльной шкале v8",
        "rank-3 ↔ rank-1 события статуса contradiction → автоматическое понижение до manual_review",
        "Аналитика не может юридически опровергнуть закон — базовое доказательственное правило",
        "AC-02 в QA-гейте: 0 нарушений шлюза в v10",
    ]
    for b in meth_bullets:
        _bullet(doc, b)
    doc.add_paragraph()

    # 2.5 What's new
    _heading(doc, "2.5 Что нового в v10", 2)
    new_bullets = [
        "4 новых артефакта: correlation_matrix, dependency_graph, coverage_heatmap, claim_provenance",
        "0 новых документов (корпус не изменился от v9)",
        "0 новых пар сравнения",
        "0 новых событий (v10 — рендеринг-релиз)",
    ]
    for b in new_bullets:
        _bullet(doc, b)

    _page_break(doc)


def _build_corpus(doc: Document, data: V10Data) -> None:
    _heading(doc, "Глава 3. Корпус: 27 документов", 1)

    # 3.1 Rank distribution
    _heading(doc, "3.1 Распределение по рангам", 2)
    dbr = data.docs_by_rank()
    r1 = len(dbr.get(1, []))
    r2 = len(dbr.get(2, []))
    r3 = len(dbr.get(3, []))
    _para(doc, f"Реальное распределение: "
              f"rank-1 = {r1}, rank-2 = {r2}, rank-3 = {r3}.")
    _add_image(doc, ASSETS_DIR / "chart_rank_distribution.png", width_cm=12.0)
    doc.add_paragraph()
    kpi3 = doc.add_table(rows=2, cols=3)
    kpi3.style = "Table Grid"
    for i, (lbl, val) in enumerate([
        ("Rank-1 (НПА)", str(r1)),
        ("Rank-2 (ведомственные)", str(r2)),
        ("Rank-3 (аналитика)", str(r3)),
    ]):
        kpi3.rows[0].cells[i].text = val
        kpi3.rows[1].cells[i].text = lbl
        _set_cell_bg(kpi3.rows[0].cells[i], OCEAN["tile_bg"])
        for cell in (kpi3.rows[0].cells[i], kpi3.rows[1].cells[i]):
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

    # 3.2 Rank-1 documents
    _heading(doc, "3.2 Список rank-1 (21 документ)", 2)
    rank1 = dbr.get(1, [])
    tbl1 = doc.add_table(rows=1 + len(rank1), cols=5)
    tbl1.style = "Table Grid"
    _table_header_row(tbl1, ["ID", "Код", "Заголовок", "Тип", "Источник"])
    _repeat_table_header(tbl1)
    for ri, d in enumerate(rank1):
        url = d.get("url") or ""
        domain = url.split("/")[2] if url.count("/") >= 2 else url[:30]
        vals = [
            d.get("id", ""),
            _trunc(d.get("code", ""), 18),
            _trunc(d.get("title", ""), 50),
            _trunc(d.get("type", ""), 22),
            _trunc(domain, 30),
        ]
        _table_body_row(tbl1, ri + 1, vals)
    doc.add_paragraph()

    # 3.3 Rank-2 + Rank-3
    _heading(doc, "3.3 Список rank-2 + rank-3 (6 документов)", 2)
    rank23 = dbr.get(2, []) + dbr.get(3, [])
    tbl23 = doc.add_table(rows=1 + len(rank23), cols=6)
    tbl23.style = "Table Grid"
    _table_header_row(tbl23, ["ID", "Ранг", "Код", "Заголовок", "Тип", "Источник"])
    for ri, d in enumerate(rank23):
        url = d.get("url") or ""
        domain = url.split("/")[2] if url.count("/") >= 2 else url[:20]
        vals = [
            d.get("id", ""),
            str(d.get("rank", "")),
            _trunc(d.get("code", ""), 15),
            _trunc(d.get("title", ""), 40),
            _trunc(d.get("type", ""), 20),
            _trunc(domain, 25),
        ]
        _table_body_row(tbl23, ri + 1, vals)
    doc.add_paragraph()

    # 3.4 Document types
    _heading(doc, "3.4 Типы документов", 2)
    type_counts = Counter(d.get("type", "—") for d in data.documents)
    for dtype, cnt in type_counts.most_common(5):
        _bullet(doc, f"{dtype}: {cnt}")
    doc.add_paragraph()

    # 3.5 Thematic clusters
    _heading(doc, "3.5 Тематические кластеры", 2)
    # Unique theme names from theme_doc
    themes_seen: dict[str, str] = {}
    for td in data.theme_doc:
        tid = td.get("theme_id", "")
        tname = td.get("theme", "")
        if tid and tname:
            themes_seen[tid] = tname
    for tid in sorted(themes_seen.keys())[:15]:
        _bullet(doc, f"{tid}: {themes_seen[tid]}")
    _add_image(doc, ASSETS_DIR / "chart_themes_distribution.png", width_cm=14.0)
    doc.add_paragraph()

    # 3.6 Provenance
    _heading(doc, "3.6 Provenance", 2)
    fetch_counts = Counter(p.get("fetch_status", "—") for p in data.provenance)
    parts = []
    for status, cnt in fetch_counts.most_common():
        parts.append(f"{status}: {cnt}")
    _para(doc, "Сводка fetch_status: " + ", ".join(parts) + ".")

    _page_break(doc)


def _build_pair_matrix(doc: Document, data: V10Data) -> None:
    _heading(doc, "Глава 4. Матрица пар (351)", 1)

    _heading(doc, "4.1 Описание", 2)
    _para(doc,
          "Все 351 пара C(27,2). "
          "Каждая отнесена "
          "к одному из 7 статусов v8.")

    _heading(doc, "4.2 Таблица всех 351 пар", 2)
    headers = ["ID", "Левый", "Правый",
               "Темы", "Статус",
               "Соб.", "Ранги"]
    pair_tbl = doc.add_table(rows=1 + len(data.pairs), cols=7)
    pair_tbl.style = "Table Grid"
    _table_header_row(pair_tbl, headers)
    _repeat_table_header(pair_tbl)

    for ri, p in enumerate(data.pairs):
        status = p.get("v8_status", "").strip()
        status_ru = STATUS_RU.get(status, status)
        topics = _trunc(p.get("topics", ""), 40)
        left_doc = _trunc(p.get("left", ""), 22)
        right_doc = _trunc(p.get("right", ""), 22)
        vals = [
            p.get("id", ""),
            left_doc,
            right_doc,
            topics,
            status_ru,
            str(p.get("events_count", "")),
            p.get("rank_pair", ""),
        ]
        row = pair_tbl.rows[ri + 1]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(7)
            if ci == 4:  # status column
                bg = STATUS_TINT_BG.get(status)
                if bg:
                    _set_cell_bg(cell, bg)

    _page_break(doc)


def _build_events(doc: Document, data: V10Data) -> None:
    _heading(doc, "Глава 5. Diff-события (312)", 1)

    _heading(doc, "5.1 Описание", 2)
    _para(doc,
          "312 событий извлечены "
          "из попарного сравнения; "
          "каждое имеет утверждение "
          "и доказательство.")

    _heading(doc, "5.2 Таблица всех 312 событий", 2)
    events_sorted = sorted(
        data.events_all,
        key=lambda e: int(e["event_id"][1:]) if e["event_id"][1:].isdigit() else 0,
    )
    ev_tbl = doc.add_table(rows=1 + len(events_sorted), cols=7)
    ev_tbl.style = "Table Grid"
    _table_header_row(ev_tbl,
                      ["ID", "Тема", "Левый",
                       "Правый",
                       "Утверждение",
                       "Доказательство",
                       "Статус"])
    _repeat_table_header(ev_tbl)

    for ri, ev in enumerate(events_sorted):
        status = ev.get("status", "").strip()
        status_ru = STATUS_RU.get(status, status)
        vals = [
            ev.get("event_id", ""),
            _trunc(ev.get("theme", ""), 22),
            _trunc(ev.get("left_doc", ev.get("left_id", "")), 18),
            _trunc(ev.get("right_doc", ev.get("right_id", "")), 18),
            _trunc(ev.get("claim_left", ""), 60),
            _trunc(ev.get("evidence_right", ""), 60),
            status_ru,
        ]
        row = ev_tbl.rows[ri + 1]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(7)
            if ci == 6:
                bg = STATUS_TINT_BG.get(status)
                if bg:
                    _set_cell_bg(cell, bg)

    _page_break(doc)


def _build_themes(doc: Document, data: V10Data) -> None:
    _heading(doc, "Глава 6. Темы и корреляция", 1)

    # 6.1 Correlation
    _heading(doc, "6.1 Корреляционная матрица", 2)
    _para(doc, "14 тем × 27 документов — "
               "бинарная карта присутствия.")
    _add_image(doc, ASSETS_DIR / "chart_correlation_heatmap.png", width_cm=15.0)
    doc.add_paragraph()

    # 6.2 Coverage
    _heading(doc, "6.2 Покрытие", 2)
    _para(doc, "Глубина покрытия "
               "по 4 рангам — где аналитика "
               "разговаривает, а НПА молчит.")
    _add_image(doc, ASSETS_DIR / "chart_coverage_heatmap.png", width_cm=15.0)
    doc.add_paragraph()

    # 6.3 Dependency graph
    _heading(doc, "6.3 Граф зависимостей", 2)
    _para(doc, "85 рёбер: документы "
               "ссылаются друг на друга.")
    _add_image(doc, ASSETS_DIR / "chart_dependency_graph.png", width_cm=15.0)
    doc.add_paragraph()
    rel_counts = Counter(e.get("relation_type", "") for e in data.dependency_graph)
    degree: dict[str, int] = {}
    for edge in data.dependency_graph:
        for key in ("from_doc_id", "to_doc_id"):
            node = edge.get(key, "")
            if node:
                degree[node] = degree.get(node, 0) + 1
    top_nodes = sorted(degree.items(), key=lambda x: -x[1])[:3]
    dep_bullets = [
        f"Всего рёбер: {len(data.dependency_graph)}",
        f"Тип «references»: {rel_counts.get('references', 0)} (ссылки)",
        f"Тип «amends»: {rel_counts.get('amends', 0)} (поправки)",
        "Топ-узлы (наибольшая степень): " +
        _trunc(", ".join(f"{data.doc_short(n)} ({c})" for n, c in top_nodes), 90),
    ]
    for b in dep_bullets:
        _bullet(doc, b)
    doc.add_paragraph()

    # 6.4 Theme catalogue
    _heading(doc, "6.4 Каталог тем", 2)
    themes_seen2: dict[str, str] = {}
    theme_doc_count: dict[str, set] = {}
    for td in data.theme_doc:
        tid = td.get("theme_id", "")
        tname = td.get("theme", "")
        did = td.get("doc_id", "")
        if tid:
            themes_seen2[tid] = tname
            theme_doc_count.setdefault(tid, set()).add(did)
    theme_rows_list = sorted(themes_seen2.keys())
    th_tbl = doc.add_table(rows=1 + len(theme_rows_list), cols=3)
    th_tbl.style = "Table Grid"
    _table_header_row(th_tbl, ["theme_id", "Название темы",
                                "Количество документов"])
    for ri, tid in enumerate(theme_rows_list):
        vals = [tid, _trunc(themes_seen2.get(tid, ""), 45), str(len(theme_doc_count.get(tid, set())))]
        _table_body_row(th_tbl, ri + 1, vals)

    _page_break(doc)


def _build_review_queue(doc: Document, data: V10Data) -> None:
    _heading(doc, "Глава 7. Очередь ручной проверки (103 задачи)", 1)

    # 7.1 Structure
    _heading(doc, "7.1 Структура очереди", 2)
    _para(doc,
          "103 = 54 baseline + 49 partial_overlap из 08_close_ac10.py")
    by_priority = data.review_by_priority()
    p0 = by_priority.get("P0", 0)
    p1 = by_priority.get("P1", 0)
    p2 = by_priority.get("P2", 0)
    prio_kpi = doc.add_table(rows=2, cols=4)
    prio_kpi.style = "Table Grid"
    for i, (val, lbl, color) in enumerate([
        (str(p0), "P0 — критично", "FEE2E2"),
        (str(p1), "P1 — важно", "FEF3C7"),
        (str(p2), "P2 — плановая", "DBEAFE"),
        (str(p0 + p1 + p2), "Итого", OCEAN["tile_bg"]),
    ]):
        prio_kpi.rows[0].cells[i].text = val
        prio_kpi.rows[1].cells[i].text = lbl
        _set_cell_bg(prio_kpi.rows[0].cells[i], color)
        for cell in (prio_kpi.rows[0].cells[i], prio_kpi.rows[1].cells[i]):
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()
    _add_image(doc, ASSETS_DIR / "chart_priority_split.png", width_cm=12.0)
    doc.add_paragraph()

    # Review queue table headers reused
    rq_headers = ["RV-ID", "Приоритет",
                  "Тема", "Что проверить",
                  "Источник",
                  "Дедлайн",
                  "Владелец"]

    # 7.2 Top P0+P1
    _heading(doc, "7.2 Топ-приоритеты P0 + P1", 2)
    p0p1_items = [r for r in data.review_queue if r.get("priority") in ("P0", "P1")]
    rq_top = doc.add_table(rows=1 + len(p0p1_items), cols=7)
    rq_top.style = "Table Grid"
    _table_header_row(rq_top, rq_headers)
    for ri, r in enumerate(p0p1_items):
        prio = r.get("priority", "")
        vals = [
            r.get("review_id", ""),
            prio,
            _trunc(r.get("theme", ""), 22),
            _trunc(r.get("what_to_check", ""), 50),
            _trunc(r.get("source", ""), 22),
            _trunc(r.get("deadline", ""), 18),
            _trunc(r.get("owner", ""), 20),
        ]
        row = rq_top.rows[ri + 1]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(8)
            if ci == 1:  # priority column
                bg = PRIO_STATUS_KEY.get(prio)
                if bg:
                    _set_cell_bg(cell, STATUS_TINT_BG.get(bg, ""))
    doc.add_paragraph()

    # 7.3 Full P2 queue
    _heading(doc, "7.3 Полная очередь P2 (97 задач)", 2)
    p2_items = [r for r in data.review_queue if r.get("priority") == "P2"]
    rq_p2 = doc.add_table(rows=1 + len(p2_items), cols=7)
    rq_p2.style = "Table Grid"
    _table_header_row(rq_p2, rq_headers)
    _repeat_table_header(rq_p2)
    for ri, r in enumerate(p2_items):
        prio = r.get("priority", "")
        vals = [
            r.get("review_id", ""),
            prio,
            _trunc(r.get("theme", ""), 22),
            _trunc(r.get("what_to_check", ""), 50),
            _trunc(r.get("source", ""), 22),
            _trunc(r.get("deadline", ""), 18),
            _trunc(r.get("owner", ""), 20),
        ]
        row = rq_p2.rows[ri + 1]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(8)
            if ci == 1:
                bg_key = PRIO_STATUS_KEY.get(prio)
                if bg_key:
                    _set_cell_bg(cell, STATUS_TINT_BG.get(bg_key, ""))
    doc.add_paragraph()

    # 7.4 Owner distribution
    _heading(doc, "7.4 Распределение по владельцам", 2)
    owner_counts = Counter(r.get("owner", "—") for r in data.review_queue)
    for owner, cnt in owner_counts.most_common(5):
        _bullet(doc, f"{owner}: {cnt}")

    _page_break(doc)


def _build_trend_qa(doc: Document, data: V10Data) -> None:
    _heading(doc, "Глава 8. Тренд и QA", 1)

    # 8.1 Trend table + charts
    _heading(doc, "8.1 Тренд", 2)
    timeline = data.trend.get("timeline", [])
    trend_tbl = doc.add_table(rows=1 + len(timeline), cols=8)
    trend_tbl.style = "Table Grid"
    _table_header_row(trend_tbl,
                      ["Версия", "Дата",
                       "Docs", "Пары", "События",
                       "Match", "Review", "Match%"])
    for ri, t in enumerate(timeline):
        vals = [
            str(t.get("version", "")),
            str(t.get("date", ""))[:10],
            str(t.get("docs", "")),
            str(t.get("pairs", "")),
            str(t.get("events", "")),
            str(t.get("status_match", "")),
            str(t.get("review_queue", "")),
            f"{t.get('match_share', '')}%",
        ]
        _table_body_row(trend_tbl, ri + 1, vals)
    doc.add_paragraph()
    _add_image(doc, ASSETS_DIR / "chart_trend_match_share.png", width_cm=14.0)
    doc.add_paragraph()
    _add_image(doc, ASSETS_DIR / "chart_trend_review_queue.png", width_cm=14.0)
    doc.add_paragraph()

    # 8.2 Degrading direction
    _heading(doc, "8.2 Обоснование «degrading direction»", 2)
    deg_bullets = [
        "trend_direction=«degrading» — намеренная метка, не сигнал ухудшения",
        "v9: добавлен D27 (ВЦИОМ, rank-3) — 26 новых пар → not_comparable",
        "match_share упал 8.0% → 0.28%: знаменатель вырос (325→351), числитель остался (1)",
        "review_queue сократился 183→54: 129 manual_review переведены в not_comparable",
    ]
    for b in deg_bullets:
        _bullet(doc, b)
    doc.add_paragraph()

    # 8.3 QA gate table
    _heading(doc, "8.3 QA-гейт", 2)
    checks = data.qa.get("checks", [])
    qa_color_map = {
        "PASS": STATUS_TINT_BG.get("match", "DCFCE7"),
        "WARN": STATUS_TINT_BG.get("partial_overlap", "FEF3C7"),
        "FAIL": STATUS_TINT_BG.get("contradiction", "FEE2E2"),
    }
    qa_tbl = doc.add_table(rows=1 + len(checks), cols=4)
    qa_tbl.style = "Table Grid"
    _table_header_row(qa_tbl,
                      ["AC-ID", "Описание",
                       "Статус",
                       "Доказательство"])
    _repeat_table_header(qa_tbl)
    for ri, c in enumerate(checks):
        status_val = c.get("status", "PASS")
        vals = [
            c.get("name", ""),
            _trunc(c.get("description", ""), 70),
            status_val,
            _trunc(str(c.get("evidence", "")), 65),
        ]
        row = qa_tbl.rows[ri + 1]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(8)
            if ci == 2:
                bg = qa_color_map.get(status_val)
                if bg:
                    _set_cell_bg(cell, bg)
    doc.add_paragraph()

    # 8.4 Caveats
    _heading(doc, "8.4 Известные оговорки", 2)
    caveat_bullets = [
        "lhs_page/rhs_page = null у ряда событий v9-наследия — допустимо",
        "rank-4 → rank-3 нормализация: не влияет на ранговый шлюз",
        "Provenance: 17 URL-строк без статуса (Гарант, КонсультантПлюс) — документы присутствуют локально",
        "v10 — рендеринг-релиз: новые артефакты не потребовали перезапуска LLM",
    ]
    for b in caveat_bullets:
        _bullet(doc, b)

    _page_break(doc)


def _build_actions(doc: Document, data: V10Data) -> None:
    _heading(doc, "Глава 9. Действия и заключение", 1)

    # 9.1 Actions catalogue
    _heading(doc, "9.1 Каталог действий (FA-01..FA-10)", 2)
    _add_image(doc, ASSETS_DIR / "chart_actions_severity.png", width_cm=12.0)
    doc.add_paragraph()
    fa_tbl = doc.add_table(rows=1 + len(data.actions), cols=7)
    fa_tbl.style = "Table Grid"
    _table_header_row(fa_tbl,
                      ["ID", "Серьёзность",
                       "Категория",
                       "Где", "Что не так",
                       "Что сделать",
                       "Владелец"])
    for ri, a in enumerate(data.actions):
        sev = a.get("severity", "")
        cat = a.get("category", "")
        cat_label = CAT_RU.get(cat, cat.replace("_", " ") if cat else "")
        vals = [
            a.get("id", ""),
            SEV_RU.get(sev, sev),
            _trunc(cat_label, 25),
            _trunc(a.get("where", ""), 25),
            _trunc(a.get("what_is_wrong", ""), 40),
            _trunc(a.get("what_to_do", ""), 40),
            _trunc(a.get("owner", ""), 20),
        ]
        row = fa_tbl.rows[ri + 1]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(8)
            if ci == 1:  # severity column
                bg_key = SEV_STATUS_KEY.get(sev)
                if bg_key:
                    _set_cell_bg(cell, STATUS_TINT_BG.get(bg_key, ""))
    doc.add_paragraph()

    # 9.2 Per-action detail cards
    for idx, a in enumerate(data.actions):
        fa_id = a.get("id", f"FA-{idx+1:02d}")
        sev = a.get("severity", "")
        cat = a.get("category", "")
        cat_label = CAT_RU.get(cat, cat.replace("_", " ") if cat else "редакционное действие")
        _heading(doc, f"{fa_id} — {_trunc(cat_label, 50)}", 3)
        _para(doc, f"Где: {a.get('where', '')}")
        _para(doc, f"Что не так: {a.get('what_is_wrong', '')}")
        _para(doc, f"Что сделать: {a.get('what_to_do', '')}")
        _para(doc, f"Ответственный: {a.get('owner', '')}")
    doc.add_paragraph()

    # 9.3 Prioritisation
    _heading(doc, "9.3 Приоритизация", 2)
    by_sev = data.actions_by_severity()
    for sev_key, sev_label in [("high", "Высокая"), ("medium", "Средняя"), ("low", "Низкая")]:
        cnt = by_sev.get(sev_key, 0)
        if cnt:
            _bullet(doc, f"{sev_label}: {cnt} действий")
    doc.add_paragraph()

    # 9.4 v11 roadmap
    _heading(doc, "9.4 Следующие итерации (v11)", 2)
    v11_bullets = [
        "Гиперссылки в XLSX на источники: восстановить URL-колонку в bundle/pairs.csv",
        "lhs_page / rhs_page: восстановить номера страниц — повторная экстракция PDF",
        "Расширение review_queue: T08 (режим высылки) и T10 (КоАП)",
        "LLM-вердикт: SEMANTIC_COMPARATOR_ENABLED=true для A/B-сравнения с fuzzy-результатами",
    ]
    for b in v11_bullets:
        _bullet(doc, b)
    doc.add_paragraph()

    # 9.5 Contacts
    _heading(doc, "9.5 Контакты и ссылки", 2)
    _para(doc,
          "Полный bundle.json: migration_v10_out/bundle/bundle.json | "
          "Машинное приложение: "
          "migration_v10_out/machine_appendix/ (14 CSV) | "
          "QA-гейт: migration_v10_out/qa_report.json | "
          "Word-документ собран scripts/presentation_v10/docx_builder.py | "
          "Версия корпуса: v10.0.0 · Дата сборки: 2026-05-09")

    _page_break(doc)


def _build_outro(doc: Document, data: V10Data) -> None:
    # This is intentionally minimal - the outro is embedded in actions chapter (9.4, 9.5)
    pass


# ---------------------------------------------------------------------------
# Wave F: New sections (F1–F5)
# ---------------------------------------------------------------------------

def _build_hero_block(doc: Document, data: V10Data) -> None:
    """F1: Hero stat & journey — short H2 block with key numbers + chart."""
    _heading(doc, "Hero: ключевые цифры", 2)
    pairs_dist = data.pairs_by_status()
    contradiction_n = pairs_dist.get("contradiction", 1)
    match_n = pairs_dist.get("match", 1)
    manual_n = pairs_dist.get("manual_review", 86)
    nc_n = pairs_dist.get("not_comparable", 202)
    _para(
        doc,
        f"Из {sum(pairs_dist.values()) or 351} пар — "
        f"{match_n} совпадение и {contradiction_n} противоречие. "
        f"{manual_n} на ручную проверку, {nc_n} несопоставимы.",
    )
    _add_image(doc, ASSETS_DIR / "chart_hero_visualization.png", width_cm=15.0)
    doc.add_paragraph()


def _build_critical_pairs(doc: Document, data: V10Data) -> None:
    """F2: Critical pairs — all contradiction + partial_overlap pairs."""
    critical = [
        p for p in data.pairs
        if p.get("v8_status") in ("contradiction", "partial_overlap")
    ]
    _heading(
        doc,
        f"Критические пары: contradiction + partial_overlap ({len(critical)} пары)",
        2,
    )
    _para(
        doc,
        f"Все {len(critical)} пар со статусами contradiction и partial_overlap — "
        "требуют приоритетной ручной проверки.",
    )
    headers = ["ID", "Левый", "Правый", "Темы", "Статус", "Соб.", "Ранги"]
    tbl = doc.add_table(rows=1 + len(critical), cols=7)
    tbl.style = "Table Grid"
    _table_header_row(tbl, headers)
    _repeat_table_header(tbl)
    for ri, p in enumerate(critical):
        status = p.get("v8_status", "").strip()
        status_ru = STATUS_RU.get(status, status)
        vals = [
            p.get("id", ""),
            _trunc(p.get("left", ""), 22),
            _trunc(p.get("right", ""), 22),
            _trunc(p.get("topics", ""), 38),
            status_ru,
            str(p.get("events_count", "")),
            p.get("rank_pair", ""),
        ]
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.name = "DejaVu Sans"
                cell.paragraphs[0].runs[0].font.size = Pt(8)
            if ci == 4:
                bg = STATUS_TINT_BG.get(status)
                if bg:
                    _set_cell_bg(cell, bg)
    doc.add_paragraph()
    _page_break(doc)


def _build_event_detail_cards(doc: Document, data: V10Data) -> None:
    """F3: Event detail cards — 10 sampled events (1 contradiction+6 partial+3 outdated)."""
    _heading(doc, "Детальные карточки событий", 2)
    _para(
        doc,
        "Выборка: 1 противоречие, 6 частичных перекрытий, 3 устаревших события — "
        "топ по confidence в каждой группе.",
    )
    doc.add_paragraph()

    # Build selection inline (mirrors slides_part5._select_top_events)
    from .theme import V8_STATUSES as _V8S
    by_status: dict[str, list[dict]] = {s: [] for s in _V8S}
    for ev in data.events_all:
        s = ev.get("status", "")
        if s in by_status:
            by_status[s].append(ev)
    for s in by_status:
        by_status[s].sort(
            key=lambda e: float(e.get("confidence", "0") or "0"),
            reverse=True,
        )
    selected: list[dict] = []
    selected.extend(by_status.get("contradiction", [])[:1])
    selected.extend(by_status.get("partial_overlap", [])[:6])
    selected.extend(by_status.get("outdated", [])[:3])
    selected = selected[:10]

    for ev in selected:
        ev_id = ev.get("event_id", "")
        theme = ev.get("theme", "")
        _heading(doc, f"{ev_id} — {_trunc(theme, 60)}", 3)
        _para(doc,
              f"Утверждение: {_trunc(ev.get('claim_left', ''), 200)} "
              f"(источник: {ev.get('left_id', '')} {_trunc(ev.get('left_doc', ''), 40)})")
        _para(doc,
              f"Доказательство: {_trunc(ev.get('evidence_right', ''), 200)} "
              f"(источник: {ev.get('right_id', '')} {_trunc(ev.get('right_doc', ''), 40)})")
        _para(doc, f"Заключение: {_trunc(ev.get('conclusion', ''), 200)}")
        _para(doc,
              f"Координата: {ev.get('legal_coordinate', '')} · "
              f"Confidence: {ev.get('confidence', '')}")
        doc.add_paragraph()

    _page_break(doc)


def _build_theme_cards(doc: Document, data: V10Data) -> None:
    """F4: Theme cards — 14 core themes with docs table + theses."""
    _heading(doc, "Карточки тем", 2)
    _para(doc, "14 тематических кластеров корпуса — документный состав и ключевые тезисы.")
    doc.add_paragraph()

    # Build theme meta (mirrors slides_part5._select_core_themes)
    theme_meta: dict[str, dict] = {}
    for row in data.theme_doc:
        tid = row.get("theme_id", "")
        if not tid:
            continue
        if tid not in theme_meta:
            theme_meta[tid] = {
                "id": tid,
                "name": row.get("theme", ""),
                "doc_count": 0,
                "event_count": 0,
            }
        theme_meta[tid]["doc_count"] += 1
    for ev in data.events_all:
        tid = ev.get("theme_id", "")
        if tid in theme_meta:
            theme_meta[tid]["event_count"] += 1
    themes = sorted(theme_meta.values(), key=lambda t: -t["event_count"])[:14]

    for theme in themes:
        tid = theme["id"]
        _heading(doc, f"Тема {tid}: {_trunc(theme['name'], 70)}", 3)
        # Docs in theme
        docs_in_theme = []
        for row in data.theme_doc:
            if row.get("theme_id") != tid:
                continue
            doc_id = row.get("doc_id", "")
            d = data.doc_by_id(doc_id)
            if d:
                docs_in_theme.append({
                    "id": doc_id,
                    "code": _trunc(d.get("code", ""), 18),
                    "rank": d.get("rank", ""),
                    "role": _trunc(row.get("role", ""), 16),
                })
        docs_in_theme = docs_in_theme[:8]

        # Count review queue items for this theme
        review_count = sum(
            1 for r in data.review_queue if r.get("theme", "") == theme["name"]
        )
        events_count = theme["event_count"]
        _para(
            doc,
            f"Документов в теме: {theme['doc_count']} · "
            f"Событий: {events_count} · "
            f"Очередь проверки: {review_count}",
        )

        # Small docs table
        if docs_in_theme:
            dt = doc.add_table(rows=1 + len(docs_in_theme), cols=4)
            dt.style = "Table Grid"
            _table_header_row(dt, ["ID", "Код", "Ранг", "Роль"])
            for ri, drow in enumerate(docs_in_theme):
                vals = [drow["id"], drow["code"], str(drow["rank"]), drow["role"]]
                _table_body_row(dt, ri + 1, vals)

        # Theses bullets
        theses_for_theme = [
            t for t in data.theses if t.get("theme", "") == theme["name"]
        ][:3]
        for t in theses_for_theme:
            _bullet(doc, _trunc(t.get("thesis", t.get("claim_text", "")), 120))
        doc.add_paragraph()

    _page_break(doc)


def _build_doc_spotlights(doc: Document, data: V10Data) -> None:
    """F5: Document spotlights — D18, D24, D27 detailed."""
    _heading(doc, "Document Spotlights", 2)
    _para(doc, "Три ключевых документа корпуса — паспорт, тезисы и граф зависимостей.")
    doc.add_paragraph()

    target_ids = ("D18", "D24", "D27")
    ordered: dict[str, dict | None] = {doc_id: None for doc_id in target_ids}
    for d in data.documents:
        if d.get("id") in target_ids:
            ordered[d["id"]] = d  # type: ignore[assignment]
    spotlight_docs = [v for v in ordered.values() if v is not None]

    for d in spotlight_docs:
        doc_id = d.get("id", "")
        doc_code = d.get("code", "")
        doc_title = _trunc(d.get("title", ""), 60)
        _heading(doc, f"{doc_id} {doc_code} — {doc_title}", 3)

        # Metadata block
        _para(doc, f"Тип: {d.get('type', '')} · Ранг: {d.get('rank', '')} · URL: {_trunc(d.get('url', ''), 60)}")

        # Top theses
        theses_for_doc = [
            t for t in data.theses if t.get("source_doc", "") == doc_code
        ][:3]
        if theses_for_doc:
            _para(doc, "Ключевые тезисы:", bold=True)
            for t in theses_for_doc:
                _bullet(doc, _trunc(t.get("thesis", t.get("claim_text", "")), 120))

        # Refs in / out
        refs_out = [e for e in data.dependency_graph if e.get("from_doc_id") == doc_id][:5]
        refs_in = [e for e in data.dependency_graph if e.get("to_doc_id") == doc_id][:5]
        if refs_out:
            _para(doc, "Ссылается на: " + ", ".join(
                e.get("to_doc_id", "") for e in refs_out
            ))
        if refs_in:
            _para(doc, "На него ссылаются: " + ", ".join(
                e.get("from_doc_id", "") for e in refs_in
            ))
        doc.add_paragraph()

    _page_break(doc)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def build_docx(out_path: Path, *, data: V10Data | None = None) -> Path:
    """Build the Word companion document and save it to *out_path*.

    Args:
        out_path: Destination .docx path.
        data: Pre-loaded V10Data; loaded from default bundle dir if None.

    Returns:
        The resolved output path.
    """
    if data is None:
        data = load_data()

    doc = Document()
    _set_default_font(doc, "DejaVu Sans", 11)
    _set_a4_portrait(doc)
    _add_footer(doc)

    _build_cover(doc, data)
    _build_toc(doc, data)
    _build_executive(doc, data)
    _build_hero_block(doc, data)
    _build_corpus(doc, data)
    _build_pair_matrix(doc, data)
    _build_critical_pairs(doc, data)
    _build_events(doc, data)
    _build_event_detail_cards(doc, data)
    _build_themes(doc, data)
    _build_theme_cards(doc, data)
    _build_review_queue(doc, data)
    _build_trend_qa(doc, data)
    _build_actions(doc, data)
    _build_doc_spotlights(doc, data)
    _build_outro(doc, data)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    import time

    t0 = time.time()
    out = build_docx(DEFAULT_OUT)
    elapsed = time.time() - t0
    size_mb = out.stat().st_size / 1_048_576
    print(f"[OK] {out}")
    print(f"     size={size_mb:.2f} MB  elapsed={elapsed:.1f}s")
