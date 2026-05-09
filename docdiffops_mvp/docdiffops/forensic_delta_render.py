"""ULTRA-HQ Russian-language renderers for forensic delta reports.

Three render targets, all consuming the dict produced by
``forensic_delta.compare_bundles`` (schema_version="v8-delta"):

  * ``render_delta_xlsx(delta, out)`` — many-sheet XLSX with KPI tiles,
    status_changes table, new/removed pair tables, distribution diff.
  * ``render_delta_docx(delta, out)`` — narrative DOCX with cover page,
    executive summary, color-coded status-shift list.
  * ``render_delta_pdf(delta, out)`` — paginated PDF with cover, KPI
    tiles, distribution diff bar chart, status-shift table.

All renderers reuse the design vocabulary from ``forensic_render``:
``PALETTE`` (palette keys), ``STATUS_RU`` (Russian status labels),
``STATUS_PALETTE`` (status → palette key), and the helpers
``_docx_set_cell_bg`` / ``_docx_add_page_numbers``.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Mapping

from .forensic import STATUS_TO_MARK, V8_STATUSES
from .forensic_delta import (
    DIRECTION_DEGRADED,
    DIRECTION_IMPROVED,
    DIRECTION_UNCHANGED,
)
from .forensic_render import (
    DOC_TITLE_RU,
    PALETTE,
    STATUS_PALETTE,
    STATUS_RU,
    _docx_add_page_numbers,
    _docx_kpi_card,
    _docx_set_cell_bg,
    _pdf_page_decoration,
    _pdf_styles,
)


DELTA_TITLE_RU = "Дельта-отчёт между запусками — DocDiffOps Forensic v8"
DELTA_SUBTITLE_RU = (
    "Сопоставление двух последовательных бандлов; отражает динамику пар "
    "и статусов. Документ предназначен для аналитики качества корпуса."
)

DIRECTION_RU: dict[str, str] = {
    DIRECTION_IMPROVED:  "Улучшилось",
    DIRECTION_DEGRADED:  "Ухудшилось",
    DIRECTION_UNCHANGED: "Без изменений",
}

DIRECTION_PALETTE: dict[str, str] = {
    DIRECTION_IMPROVED:  "match",
    DIRECTION_DEGRADED:  "contradict",
    DIRECTION_UNCHANGED: "nc",
}


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def render_delta_xlsx(delta: Mapping[str, Any], out_path: Path | str) -> None:
    """Multi-sheet XLSX delta report with cover, status changes, distribution diff."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    head_fill = PatternFill(start_color=PALETTE["primary"],
                            end_color=PALETTE["primary"], fill_type="solid")
    head_font = Font(bold=True, color="FFFFFF")
    wrap = Alignment(wrap_text=True, vertical="top")
    thin = Side(border_style="thin", color=PALETTE["border"])
    border = Border(top=thin, bottom=thin, left=thin, right=thin)

    cn = delta.get("control_numbers") or {}

    def _style_header(ws: Any, row: int = 1) -> None:
        for cell in ws[row]:
            cell.font = head_font
            cell.fill = head_fill
            cell.alignment = wrap
            cell.border = border

    def _table(ws: Any, header: list[str], rows: list[dict[str, Any]],
               col_widths: dict[str, int] | None = None) -> None:
        ws.append(header)
        _style_header(ws)
        for row in rows:
            ws.append([row.get(h, "") for h in header])
        for i, h in enumerate(header, 1):
            ws.column_dimensions[get_column_letter(i)].width = (
                (col_widths or {}).get(h, max(12, min(48, len(h) + 2)))
            )
        ws.freeze_panes = "A2"
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            for cell in row:
                cell.alignment = wrap
                cell.border = border

    # ---------- Sheet 00 — Обложка ----------
    ws = wb.active
    ws.title = "00 Обложка"

    title_fill = PatternFill(start_color=PALETTE["accent"],
                             end_color=PALETTE["accent"], fill_type="solid")
    ws.append([DELTA_TITLE_RU])
    ws.merge_cells(start_row=1, end_row=1, start_column=1, end_column=5)
    ws["A1"].font = Font(bold=True, size=18, color="FFFFFF")
    ws["A1"].fill = title_fill
    ws["A1"].alignment = Alignment(vertical="center", horizontal="left",
                                    indent=1)
    ws.row_dimensions[1].height = 36

    ws.append([DELTA_SUBTITLE_RU])
    ws.merge_cells(start_row=2, end_row=2, start_column=1, end_column=5)
    ws["A2"].font = Font(italic=True, size=10, color=PALETTE["nc"])
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="top", indent=1)
    ws.row_dimensions[2].height = 30

    ws.append([
        f"Базовый запуск (old): {delta.get('baseline_generated_at','—')}",
        "", "",
        f"Текущий запуск (new): {delta.get('current_generated_at','—')}",
        "",
    ])
    for c in ws[3]:
        c.font = Font(size=9, color=PALETTE["nc"])

    ws.append([])

    # KPI tiles — 5 colored cells in row 5, captions in row 6
    kpi_specs = [
        (cn.get("pairs_total", 0),    "Пар (всего)",    PALETTE["primary"]),
        (cn.get("pairs_changed", 0),  "Изменено",       PALETTE["partial"]),
        (cn.get("pairs_resolved", 0), "Закрыто (→ ✓)",  PALETTE["match"]),
        (cn.get("pairs_new", 0),      "Новых",          PALETTE["accent"]),
        (cn.get("pairs_removed", 0),  "Удалено",        PALETTE["nc"]),
    ]
    ws.append([str(spec[0]) for spec in kpi_specs])
    ws.append([spec[1] for spec in kpi_specs])
    for i, (_, _, hex_color) in enumerate(kpi_specs, start=1):
        v_cell = ws.cell(row=5, column=i)
        v_cell.font = Font(bold=True, size=22, color="FFFFFF")
        v_cell.alignment = Alignment(horizontal="center", vertical="center")
        v_cell.fill = PatternFill(start_color=hex_color,
                                  end_color=hex_color, fill_type="solid")
        v_cell.border = border
        l_cell = ws.cell(row=6, column=i)
        l_cell.font = Font(bold=True, size=9, color=PALETTE["ink"])
        l_cell.alignment = Alignment(horizontal="center", vertical="center")
        l_cell.fill = PatternFill(start_color=PALETTE["muted"],
                                  end_color=PALETTE["muted"], fill_type="solid")
        l_cell.border = border
        ws.column_dimensions[get_column_letter(i)].width = 22
    ws.row_dimensions[5].height = 50
    ws.row_dimensions[6].height = 22

    ws.append([])
    if delta.get("asymmetric_actions_warning"):
        ws.append(["⚠ " + delta["asymmetric_actions_warning"]])
        ws.cell(row=ws.max_row, column=1).font = Font(
            italic=True, size=9, color=PALETTE["review"]
        )
        ws.append([])

    ws.append(["Покрытие действиями:"])
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=10,
                                                   color=PALETTE["ink"])
    ws.append([f"actions_coverage = {delta.get('actions_coverage','—')}"])

    # ---------- Sheet 01 — Изменения статусов ----------
    ws = wb.create_sheet("01 Изменения статусов")
    fills = {
        "match":     PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid"),
        "partial":   PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid"),
        "contradict":PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid"),
        "outdated":  PatternFill(start_color="DBEAFE", end_color="DBEAFE", fill_type="solid"),
        "gap":       PatternFill(start_color="EDE9FE", end_color="EDE9FE", fill_type="solid"),
        "review":    PatternFill(start_color="FFEDD5", end_color="FFEDD5", fill_type="solid"),
        "nc":        PatternFill(start_color="F3F4F6", end_color="F3F4F6", fill_type="solid"),
    }
    cols = ["pair_id", "left_id", "right_id",
            "old_status", "old_status_ru", "new_status", "new_status_ru",
            "direction", "direction_ru"]
    rows_view = []
    for ch in delta.get("status_changes", []):
        rows_view.append({
            "pair_id":       ch.get("pair_id", ""),
            "left_id":       ch.get("left_id", ""),
            "right_id":      ch.get("right_id", ""),
            "old_status":    ch.get("old_status", ""),
            "old_status_ru": STATUS_RU.get(ch.get("old_status", ""), ""),
            "new_status":    ch.get("new_status", ""),
            "new_status_ru": STATUS_RU.get(ch.get("new_status", ""), ""),
            "direction":     ch.get("direction", ""),
            "direction_ru":  DIRECTION_RU.get(ch.get("direction", ""), ""),
        })
    _table(ws, cols, rows_view, col_widths={
        "old_status": 18, "old_status_ru": 22,
        "new_status": 18, "new_status_ru": 22,
        "direction_ru": 18,
    })
    # Tint direction cells
    dir_idx = cols.index("direction") + 1
    for r, ch in enumerate(delta.get("status_changes", []), start=2):
        pal_key = DIRECTION_PALETTE.get(ch.get("direction", ""), "nc")
        cell = ws.cell(row=r, column=dir_idx)
        cell.fill = fills.get(pal_key)

    # ---------- Sheet 02 — Распределение (дельта) ----------
    ws = wb.create_sheet("02 Распределение (Δ)")
    cols = ["status", "status_ru", "delta"]
    rows_view = []
    for st, d in delta.get("distribution_diff", {}).items():
        rows_view.append({
            "status": st,
            "status_ru": STATUS_RU.get(st, st),
            "delta": d,
        })
    rows_view.sort(key=lambda r: -abs(r["delta"]))
    _table(ws, cols, rows_view, col_widths={"status": 22, "status_ru": 28,
                                            "delta": 14})
    # Color delta column: positive=green, negative=red
    delta_idx = cols.index("delta") + 1
    pos_fill = PatternFill(start_color="DCFCE7", end_color="DCFCE7",
                           fill_type="solid")
    neg_fill = PatternFill(start_color="FEE2E2", end_color="FEE2E2",
                           fill_type="solid")
    for r, row in enumerate(rows_view, start=2):
        cell = ws.cell(row=r, column=delta_idx)
        cell.fill = pos_fill if row["delta"] > 0 else (
            neg_fill if row["delta"] < 0 else None
        )

    # ---------- Sheet 03 — Новые пары ----------
    new_pairs = delta.get("new_pairs", [])
    if new_pairs:
        ws = wb.create_sheet("03 Новые пары")
        cols = ["id", "left", "right", "v8_status", "v8_status_ru",
                "events_count", "topics"]
        rows_view = [{
            "id": p.get("id", ""),
            "left": p.get("left", ""),
            "right": p.get("right", ""),
            "v8_status": p.get("v8_status", ""),
            "v8_status_ru": STATUS_RU.get(p.get("v8_status", ""), ""),
            "events_count": p.get("events_count", 0),
            "topics": "; ".join(p.get("topics", [])),
        } for p in new_pairs]
        _table(ws, cols, rows_view, col_widths={"v8_status": 20,
                                                "v8_status_ru": 22,
                                                "topics": 36})

    # ---------- Sheet 04 — Удалённые пары ----------
    removed_pairs = delta.get("removed_pairs", [])
    if removed_pairs:
        ws = wb.create_sheet("04 Удалённые пары")
        cols = ["id", "left", "right", "v8_status", "v8_status_ru",
                "events_count", "topics"]
        rows_view = [{
            "id": p.get("id", ""),
            "left": p.get("left", ""),
            "right": p.get("right", ""),
            "v8_status": p.get("v8_status", ""),
            "v8_status_ru": STATUS_RU.get(p.get("v8_status", ""), ""),
            "events_count": p.get("events_count", 0),
            "topics": "; ".join(p.get("topics", [])),
        } for p in removed_pairs]
        _table(ws, cols, rows_view, col_widths={"v8_status": 20,
                                                "v8_status_ru": 22,
                                                "topics": 36})

    # Print setup + document properties
    for sheet_name in wb.sheetnames:
        sh = wb[sheet_name]
        sh.page_setup.orientation = sh.ORIENTATION_LANDSCAPE
        sh.page_setup.fitToWidth = 1
        sh.page_setup.fitToHeight = 0
        sh.print_options.horizontalCentered = True
        sh.sheet_properties.pageSetUpPr.fitToPage = True
        sh.print_title_rows = "1:1"
        sh.oddHeader.center.text = "DocDiffOps · Forensic v8 · Δ-отчёт"
        sh.oddHeader.center.size = 9
        sh.oddFooter.right.text = "&P / &N"
        sh.oddFooter.right.size = 8

    wb.properties.title = DELTA_TITLE_RU
    wb.properties.subject = "DocDiffOps Forensic v8 — delta report"
    wb.properties.creator = "DocDiffOps Forensic v8"
    wb.properties.language = "ru-RU"

    wb.save(str(out_path))


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def render_delta_docx(delta: Mapping[str, Any], out_path: Path | str) -> None:
    """Narrative DOCX delta report with cover, executive summary, status shifts."""
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt, RGBColor

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    def _hex_to_rgb(hex_str: str) -> RGBColor:
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16),
                        int(hex_str[4:6], 16))

    INK   = _hex_to_rgb(PALETTE["ink"])
    NAVY  = _hex_to_rgb(PALETTE["accent"])
    GRAY  = _hex_to_rgb(PALETTE["nc"])

    cn = delta.get("control_numbers") or {}

    doc = Document()
    doc.core_properties.title = DELTA_TITLE_RU
    doc.core_properties.subject = "DocDiffOps Forensic v8 — delta report"
    doc.core_properties.language = "ru-RU"
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ----- Cover -----------------------------------------------------------
    band = doc.add_paragraph()
    band_run = band.add_run("DOCDIFFOPS · FORENSIC v8 · ДЕЛЬТА-ОТЧЁТ")
    band_run.bold = True
    band_run.font.size = Pt(11)
    band_run.font.color.rgb = NAVY

    title = doc.add_paragraph()
    t_run = title.add_run(DELTA_TITLE_RU)
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    s_run = subtitle.add_run(DELTA_SUBTITLE_RU)
    s_run.italic = True
    s_run.font.size = Pt(10)
    s_run.font.color.rgb = GRAY

    # Metadata table
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Light Grid"
    meta.rows[0].cells[0].text = "Базовый запуск"
    meta.rows[0].cells[1].text = "Текущий запуск"
    meta.rows[0].cells[2].text = "Тип отчёта"
    meta.rows[0].cells[3].text = "Покрытие действиями"
    meta.rows[1].cells[0].text = delta.get("baseline_generated_at", "—")
    meta.rows[1].cells[1].text = delta.get("current_generated_at", "—")
    meta.rows[1].cells[2].text = delta.get("schema_version", "v8-delta")
    meta.rows[1].cells[3].text = delta.get("actions_coverage", "—")
    for cell in meta.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _docx_set_cell_bg(cell, PALETTE["primary"])
    for cell in meta.rows[1].cells:
        _docx_set_cell_bg(cell, PALETTE["muted"])

    doc.add_paragraph()

    # KPI tiles
    kpi_specs = [
        (cn.get("pairs_total", 0),    "Пар (всего)",    PALETTE["primary"]),
        (cn.get("pairs_changed", 0),  "Изменено",       PALETTE["partial"]),
        (cn.get("pairs_resolved", 0), "Закрыто (→ ✓)",  PALETTE["match"]),
        (cn.get("pairs_new", 0),      "Новых",          PALETTE["accent"]),
    ]
    kpi_table = doc.add_table(rows=2, cols=len(kpi_specs))
    for i, (val, lbl, color) in enumerate(kpi_specs):
        _docx_kpi_card(kpi_table, 0, i, str(val), lbl, color)

    doc.add_paragraph()

    # Executive summary line
    summary = doc.add_paragraph()
    sum_run = summary.add_run(
        f"Краткая сводка: всего {cn.get('pairs_total', 0)} пар; "
        f"изменилось — {cn.get('pairs_changed', 0)}, "
        f"закрыто — {cn.get('pairs_resolved', 0)}, "
        f"новых — {cn.get('pairs_new', 0)}, "
        f"удалено — {cn.get('pairs_removed', 0)}."
    )
    sum_run.bold = True
    sum_run.font.size = Pt(10)
    sum_run.font.color.rgb = NAVY

    if delta.get("asymmetric_actions_warning"):
        warn = doc.add_paragraph()
        w_run = warn.add_run("⚠ " + delta["asymmetric_actions_warning"])
        w_run.italic = True
        w_run.font.size = Pt(9)
        w_run.font.color.rgb = _hex_to_rgb(PALETTE["review"])

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ----- Раздел 1. Изменения статусов -----------------------------------
    doc.add_heading("Раздел 1. Изменения статусов", level=1)
    changes = delta.get("status_changes", [])
    if changes:
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid"
        for i, h in enumerate(("Пара", "Слева ↔ Справа",
                               "Было", "Стало", "Направление")):
            t.rows[0].cells[i].text = h
            _docx_set_cell_bg(t.rows[0].cells[i], PALETTE["primary"])
            for run in t.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for ch in changes:
            cells = t.add_row().cells
            cells[0].text = ch.get("pair_id", "")
            cells[1].text = f"{ch.get('left_id','')} ↔ {ch.get('right_id','')}"
            cells[2].text = STATUS_RU.get(ch.get("old_status", ""),
                                          ch.get("old_status", ""))
            cells[3].text = STATUS_RU.get(ch.get("new_status", ""),
                                          ch.get("new_status", ""))
            cells[4].text = DIRECTION_RU.get(ch.get("direction", ""),
                                             ch.get("direction", ""))
            dir_pal = DIRECTION_PALETTE.get(ch.get("direction", ""), "nc")
            _docx_set_cell_bg(cells[4], PALETTE.get(dir_pal, PALETTE["nc"]))
            for run in cells[4].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    else:
        doc.add_paragraph("Изменений статусов не зафиксировано.")

    # ----- Раздел 2. Распределение (Δ) -----------------------------------
    doc.add_heading("Раздел 2. Распределение пар по статусам (Δ)", level=1)
    dist = delta.get("distribution_diff") or {}
    if dist:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid"
        for i, h in enumerate(("Статус (код)", "Русское название", "Δ")):
            t.rows[0].cells[i].text = h
            _docx_set_cell_bg(t.rows[0].cells[i], PALETTE["primary"])
            for run in t.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for st, d in sorted(dist.items(), key=lambda kv: -abs(kv[1])):
            cells = t.add_row().cells
            cells[0].text = st
            cells[1].text = STATUS_RU.get(st, st)
            cells[2].text = f"{d:+d}"
            sign_pal = "match" if d > 0 else ("contradict" if d < 0 else "nc")
            _docx_set_cell_bg(cells[2], PALETTE[sign_pal])
            for run in cells[2].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    else:
        doc.add_paragraph("Распределение статусов не изменилось.")

    # ----- Раздел 3. Новые / удалённые пары --------------------------------
    new_pairs = delta.get("new_pairs", [])
    removed_pairs = delta.get("removed_pairs", [])
    if new_pairs or removed_pairs:
        doc.add_heading("Раздел 3. Новые и удалённые пары", level=1)
        if new_pairs:
            doc.add_heading(f"3.1 Новые пары ({len(new_pairs)})", level=2)
            for p in new_pairs:
                para = doc.add_paragraph()
                head = para.add_run(
                    f"{p.get('id','')} • {p.get('left','')} ↔ {p.get('right','')} — "
                )
                head.bold = True
                body = para.add_run(STATUS_RU.get(p.get("v8_status", ""),
                                                   p.get("v8_status", "")))
                body.font.color.rgb = _hex_to_rgb(PALETTE.get(
                    STATUS_PALETTE.get(p.get("v8_status", ""), "nc"),
                    PALETTE["nc"]
                ))
                body.bold = True
        if removed_pairs:
            doc.add_heading(f"3.2 Удалённые пары ({len(removed_pairs)})",
                            level=2)
            for p in removed_pairs:
                para = doc.add_paragraph()
                head = para.add_run(
                    f"{p.get('id','')} • {p.get('left','')} ↔ {p.get('right','')} — "
                )
                head.bold = True
                body = para.add_run(STATUS_RU.get(p.get("v8_status", ""),
                                                   p.get("v8_status", "")))
                body.font.color.rgb = GRAY
                body.italic = True

    _docx_add_page_numbers(doc)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def render_delta_pdf(delta: Mapping[str, Any], out_path: Path | str) -> None:
    """Paginated PDF delta report with cover, KPI tiles, distribution chart, table."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import (PageBreak, Paragraph, SimpleDocTemplate,
                                     Spacer, Table, TableStyle)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    s = _pdf_styles()
    colors = s["colors"]
    base_font = s["base"]
    bold_font = s["bold"]

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=1.2 * cm, rightMargin=1.2 * cm,
        topMargin=2.0 * cm, bottomMargin=1.6 * cm,
        title=DELTA_TITLE_RU,
        author="DocDiffOps Forensic v8",
        subject="DocDiffOps Forensic v8 — delta report",
    )

    def _on_page(canvas: Any, doc_: Any) -> None:
        _pdf_page_decoration(canvas, doc_, base_font, bold_font)

    elems: list[Any] = []
    cn = delta.get("control_numbers") or {}

    # ----- Cover -----------------------------------------------------------
    elems.append(Spacer(1, 16))
    elems.append(Paragraph(DELTA_TITLE_RU, s["h1"]))
    elems.append(Paragraph(
        f'<font color="#{PALETTE["nc"]}"><i>{DELTA_SUBTITLE_RU}</i></font>',
        s["body"]))
    elems.append(Spacer(1, 12))

    # KPI tiles — 5 cards
    kpi_data = [[
        Paragraph(f'<font size="20" color="white"><b>{cn.get("pairs_total", 0)}</b></font><br/>'
                  f'<font size="8" color="white">Пар (всего)</font>', s["body"]),
        Paragraph(f'<font size="20" color="white"><b>{cn.get("pairs_changed", 0)}</b></font><br/>'
                  f'<font size="8" color="white">Изменено</font>', s["body"]),
        Paragraph(f'<font size="20" color="white"><b>{cn.get("pairs_resolved", 0)}</b></font><br/>'
                  f'<font size="8" color="white">Закрыто</font>', s["body"]),
        Paragraph(f'<font size="20" color="white"><b>{cn.get("pairs_new", 0)}</b></font><br/>'
                  f'<font size="8" color="white">Новых</font>', s["body"]),
        Paragraph(f'<font size="20" color="white"><b>{cn.get("pairs_removed", 0)}</b></font><br/>'
                  f'<font size="8" color="white">Удалено</font>', s["body"]),
    ]]
    kpi_table = Table(kpi_data, colWidths=[3.5 * cm] * 5, rowHeights=[2.0 * cm])
    kpi_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#" + PALETTE["primary"])),
        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#" + PALETTE["partial"])),
        ("BACKGROUND", (2, 0), (2, 0), colors.HexColor("#" + PALETTE["match"])),
        ("BACKGROUND", (3, 0), (3, 0), colors.HexColor("#" + PALETTE["accent"])),
        ("BACKGROUND", (4, 0), (4, 0), colors.HexColor("#" + PALETTE["nc"])),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elems.append(kpi_table)
    elems.append(Spacer(1, 12))

    elems.append(Paragraph(
        f'<font color="#{PALETTE["nc"]}">'
        f'Базовый запуск: {delta.get("baseline_generated_at","—")}<br/>'
        f'Текущий запуск: {delta.get("current_generated_at","—")}<br/>'
        f'Покрытие действиями: {delta.get("actions_coverage","—")}'
        f'</font>',
        s["body"]))

    if delta.get("asymmetric_actions_warning"):
        elems.append(Paragraph(
            f'<font color="#{PALETTE["review"]}">⚠ {delta["asymmetric_actions_warning"]}</font>',
            s["body"]))

    elems.append(Spacer(1, 16))

    # Distribution diff bar chart (text-based)
    dist = delta.get("distribution_diff") or {}
    if dist:
        elems.append(Paragraph("Распределение пар по статусам (Δ)", s["h2"]))
        rows = [["Статус (код)", "Русское название", "Δ"]]
        for st, d in sorted(dist.items(), key=lambda kv: -abs(kv[1])):
            rows.append([st, STATUS_RU.get(st, st), f"{d:+d}"])
        t = Table(rows, colWidths=[5 * cm, 7 * cm, 2 * cm], repeatRows=1)
        dist_style = [
            ("FONT", (0, 0), (-1, 0), bold_font, 9),
            ("FONT", (0, 1), (-1, -1), base_font, 8),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + PALETTE["primary"])),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (2, 1), (2, -1), "RIGHT"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#" + PALETTE["border"])),
        ]
        for idx, (_, d) in enumerate(sorted(dist.items(),
                                            key=lambda kv: -abs(kv[1])),
                                     start=1):
            sign_pal = "match" if d > 0 else ("contradict" if d < 0 else "nc")
            dist_style.append((
                "BACKGROUND", (2, idx), (2, idx),
                colors.HexColor("#" + PALETTE[sign_pal]),
            ))
            dist_style.append(("TEXTCOLOR", (2, idx), (2, idx), colors.white))
        t.setStyle(TableStyle(dist_style))
        elems.append(t)
    elems.append(PageBreak())

    # ---------- Section: Status changes ----------
    elems.append(Paragraph("Раздел 1. Изменения статусов", s["h2"]))
    changes = delta.get("status_changes", [])
    if changes:
        rows = [["Пара", "Слева ↔ Справа", "Было", "Стало", "Направление"]]
        for ch in changes:
            rows.append([
                ch.get("pair_id", ""),
                f"{ch.get('left_id','')} ↔ {ch.get('right_id','')}",
                STATUS_RU.get(ch.get("old_status", ""), ch.get("old_status", "")),
                STATUS_RU.get(ch.get("new_status", ""), ch.get("new_status", "")),
                DIRECTION_RU.get(ch.get("direction", ""), ch.get("direction", "")),
            ])
        t = Table(rows, colWidths=[1.6 * cm, 3.6 * cm, 4 * cm, 4 * cm, 3 * cm],
                  repeatRows=1)
        ch_style = [
            ("FONT", (0, 0), (-1, 0), bold_font, 9),
            ("FONT", (0, 1), (-1, -1), base_font, 8),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + PALETTE["primary"])),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#" + PALETTE["muted"])]),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#" + PALETTE["border"])),
        ]
        for idx, ch in enumerate(changes, start=1):
            dir_pal = DIRECTION_PALETTE.get(ch.get("direction", ""), "nc")
            ch_style.append((
                "BACKGROUND", (4, idx), (4, idx),
                colors.HexColor("#" + PALETTE[dir_pal]),
            ))
            ch_style.append(("TEXTCOLOR", (4, idx), (4, idx), colors.white))
        t.setStyle(TableStyle(ch_style))
        elems.append(t)
    else:
        elems.append(Paragraph(
            "<i>Изменений статусов не зафиксировано.</i>", s["body"]))

    doc.build(elems, onFirstPage=_on_page, onLaterPages=_on_page)


__all__ = [
    "render_delta_xlsx",
    "render_delta_docx",
    "render_delta_pdf",
    "DIRECTION_RU",
    "DIRECTION_PALETTE",
    "DELTA_TITLE_RU",
    "DELTA_SUBTITLE_RU",
]
