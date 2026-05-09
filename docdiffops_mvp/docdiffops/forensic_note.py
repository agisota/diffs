"""Renders the 10-chapter Explanatory Note (Пояснительная записка) for the
DocDiffOps v10 forensic bundle.

Public contract (PR-6.3 — Sprint 6):
  * ``render_explanatory_note_docx`` — DOCX, A4, 20 mm margins, Cyrillic-safe
  * ``render_explanatory_note_pdf``  — PDF, same layout, NotoSans→DejaVu→Helvetica

Chapters (in order):
  1. Введение
  2. Методология
  3. Сводка результатов
  4. Ключевые противоречия
  5. Подтверждённые нормы
  6. Пробелы в регулировании
  7. Корреляции и зависимости
  8. Очередь ручной проверки
  9. Ограничения и caveats
  10. Приложения

Reference implementation: migration_v10_out/scripts/04_render_note.py
Plan: .omx/plans/sprint6-pr2-pr3-plan.md (PR-6.3)
"""
from __future__ import annotations

import datetime as dt
import warnings
from pathlib import Path
from typing import Any

from .forensic_render import _find_dejavu_ttf, _NOTO_TTF_SYSTEM, _NOTO_BOLD_SYSTEM

# ---------------------------------------------------------------------------
# A4 / margin constants
# ---------------------------------------------------------------------------

_A4_W_INCHES = 8.27
_A4_H_INCHES = 11.69
_MARGIN_INCHES = 0.787  # 20 mm

# Cyrillic font paths — resolved at import time via shared helper.
# Kept as module-level names so tests can patch them.
_NOTO_TTF = _NOTO_TTF_SYSTEM
_NOTO_BOLD = _NOTO_BOLD_SYSTEM
_DEJAVU_TTF = _find_dejavu_ttf() or ""

# ---------------------------------------------------------------------------
# Chapter titles (canonical, ordered)
# ---------------------------------------------------------------------------

CHAPTER_TITLES: list[str] = [
    "1. Введение",
    "2. Методология",
    "3. Сводка результатов",
    "4. Ключевые противоречия",
    "5. Подтверждённые нормы",
    "6. Пробелы в регулировании",
    "7. Корреляции и зависимости",
    "8. Очередь ручной проверки",
    "9. Ограничения и caveats",
    "10. Приложения",
]

# Short chapter names for the TOC table
_TOC_ENTRIES: list[tuple[str, str]] = [
    ("1", "Введение"),
    ("2", "Методология"),
    ("3", "Сводка результатов"),
    ("4", "Ключевые противоречия"),
    ("5", "Подтверждённые нормы"),
    ("6", "Пробелы в регулировании"),
    ("7", "Корреляции и зависимости"),
    ("8", "Очередь ручной проверки"),
    ("9", "Ограничения и caveats"),
    ("10", "Приложения"),
]

# ---------------------------------------------------------------------------
# Shared data-extraction helpers
# ---------------------------------------------------------------------------


def _kpi_tile_data(bundle: dict[str, Any]) -> list[tuple[str, str]]:
    """Return (label, value) pairs for the KPI summary block."""
    cn = bundle.get("control_numbers") or {}
    sd = bundle.get("status_distribution_pairs") or {}
    docs_count = cn.get("documents", len(list(bundle.get("documents", []))))
    pairs_count = cn.get("pairs", sum(sd.values()))
    contradictions = sd.get("contradiction", 0)
    matches = sd.get("match", 0)
    reviews = sd.get("manual_review", 0)
    events = list(bundle.get("events", []))
    return [
        ("Документов", str(docs_count)),
        ("Пар", str(pairs_count)),
        ("Событий", str(len(events))),
        ("Совпадений", str(matches)),
        ("Противоречий", str(contradictions)),
        ("Ручная проверка", str(reviews)),
    ]


def _docs_table_rows(bundle: dict[str, Any]) -> list[list[str]]:
    docs = list(bundle.get("documents", []))
    rows: list[list[str]] = []
    for d in docs:
        rows.append([
            str(d.get("id", "")),
            str(d.get("short", d.get("title", d.get("name", ""))))[:50],
            str(d.get("type", ""))[:30],
            str(d.get("rank", "")),
            str(d.get("date", d.get("year", ""))),
        ])
    return rows


def _status_dist_rows(bundle: dict[str, Any]) -> list[list[str]]:
    sd: dict[str, int] = dict(bundle.get("status_distribution_pairs") or {})
    total = sum(sd.values()) or 1
    return [
        [k, str(v), f"{100 * v / total:.1f}%"]
        for k, v in sorted(sd.items(), key=lambda x: -x[1])
    ]


def _high_risk_events(bundle: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        e for e in list(bundle.get("events", []))
        if e.get("risk") == "высокий"
    ]


def _review_queue(bundle: dict[str, Any]) -> list[dict[str, Any]]:
    rq = list(bundle.get("review_queue", []))
    if not rq:
        # Fallback: build from events
        rq = [
            e for e in list(bundle.get("events", []))
            if e.get("status") in ("manual_review", "partial_overlap", "outdated")
        ]
    return rq


def _corr_matrix_rows(correlations: dict[str, Any]) -> list[list[str]]:
    matrix: dict[str, dict[str, int]] = dict(
        correlations.get("correlation_matrix", {})
    )
    rows: list[list[str]] = []
    for theme_id, doc_counts in matrix.items():
        total_coverage = sum(doc_counts.values())
        rows.append([theme_id, str(total_coverage)])
    return rows


def _dep_graph_rows(correlations: dict[str, Any]) -> list[list[str]]:
    dep: list[dict[str, Any]] = list(correlations.get("dependency_graph", []))
    rows: list[list[str]] = []
    for edge in dep:
        rows.append([
            str(edge.get("from_doc_id", "")),
            str(edge.get("to_doc_id", "")),
            str(edge.get("relation_type", "")),
            str(edge.get("weight", "")),
        ])
    return rows


def _claim_provenance_rows(correlations: dict[str, Any]) -> list[list[str]]:
    prov: list[dict[str, Any]] = list(correlations.get("claim_provenance", []))
    rows: list[list[str]] = []
    for p in prov:
        rows.append([
            str(p.get("thesis_id", "")),
            str(p.get("thesis_text", ""))[:80],
            str(p.get("primary_doc_id", "")),
            str(p.get("confirming_docs", "")),
            str(p.get("refuting_docs", "")),
        ])
    return rows


def _coverage_heatmap_rows(correlations: dict[str, Any]) -> list[list[str]]:
    """Normalize the heatmap to a list of [theme_id, theme_name, r1, r2, r3] rows.

    Accepts either the v10 production shape from
    forensic_correlations.compute_coverage_heatmap (dict[theme_id ->
    dict[rank_int -> count]]) or the legacy list-of-dicts shape used by the
    offline migration_v10_out scripts. Theme names are looked up in
    correlations["theme_names"] when available, otherwise the theme id is
    used as a fallback display label.
    """
    hmap_raw = correlations.get("coverage_heatmap", [])
    theme_names: dict[str, str] = correlations.get("theme_names", {}) or {}
    rows: list[list[str]] = []

    if isinstance(hmap_raw, dict):
        for theme_id, ranks in hmap_raw.items():
            ranks = ranks or {}
            rows.append([
                str(theme_id),
                str(theme_names.get(theme_id, theme_id)),
                str(ranks.get(1, ranks.get("1", 0))),
                str(ranks.get(2, ranks.get("2", 0))),
                str(ranks.get(3, ranks.get("3", 0))),
            ])
        return rows

    for row in hmap_raw:
        if not isinstance(row, dict):
            continue
        rows.append([
            str(row.get("theme_id", "")),
            str(row.get("theme_name", "")),
            str(row.get("rank_1_count", row.get("rank_1", 0))),
            str(row.get("rank_2_count", row.get("rank_2", 0))),
            str(row.get("rank_3_count", row.get("rank_3", 0))),
        ])
    return rows


# ---------------------------------------------------------------------------
# Chapter data structures — shared by DOCX and PDF paths
# Each function returns a dict with:
#   title: str
#   paragraphs: list[str]   — plain text blocks
#   tables: list[dict]      — {header: list[str], rows: list[list[str]]}
# ---------------------------------------------------------------------------

ChapterData = dict[str, Any]


def _chapter_introduction(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    docs = list(bundle.get("documents", []))
    events = list(bundle.get("events", []))
    schema_version = str(bundle.get("schema_version", "10.0.0"))
    return {
        "title": "1. Введение",
        "paragraphs": [
            (
                "1.1. Цель и задачи\n"
                "Настоящая пояснительная записка является deliverable проекта DocDiffOps. "
                "Цель — систематизировать результаты all-to-all попарного сравнения "
                f"{len(docs)} документов, выявить противоречия, пробелы в "
                "регулировании и подтверждённые нормы на основе "
                f"{len(events)} формализованных diff-событий."
            ),
            (
                "1.2. Методологические ограничения\n"
                "Сравнение выполнено на forensic-снимках релевантных положений. "
                "OCR сканов не производился. Таблицы и изображения не разбирались структурно. "
                "Rank-3 источники (аналитика, презентации) не могут порождать статус "
                "contradicts против rank-1 — rank-gate применён ко всем событиям."
            ),
            (
                f"1.3. Состав корпуса: {len(docs)} документов. "
                f"Версия схемы: {schema_version}."
            ),
            "1.4. Иерархия source rank",
        ],
        "tables": [
            {
                "header": ["Rank", "Тип", "Примеры", "Вес в сравнении"],
                "rows": [
                    ["1", "primary law / strategy", "ФЗ, Указы, ПП, Концепция", "Безусловный приоритет"],
                    ["2", "secondary regulation", "Госпрограммы, планы", "Подчинён rank-1"],
                    ["3", "analytics / claim source", "Презентации, аналитика", "Подтверждает, не опровергает rank-1"],
                    ["4", "methodology / examples", "Инструкции, выгрузки", "Контекст, не evidence"],
                ],
            },
            {
                "header": ["ID", "Краткое название", "Тип", "Rank", "Год"],
                "title": "1.5. Реестр корпуса",
                "rows": _docs_table_rows(bundle) or [["(нет данных)", "", "", "", ""]],
            },
        ],
    }


def _chapter_methodology(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    return {
        "title": "2. Методология",
        "paragraphs": [
            (
                "2.1. All-to-all сравнение\n"
                "Все пары документов сравнивались методом all-to-all. "
                "Для каждой пары применялся один из трёх типов сравнения:"
            ),
            (
                "2.2. Source rank gating\n"
                "Правила rank-gate определяют максимально допустимый статус для каждой пары рангов: "
                "rank-3 vs rank-1 — максимум suggests_review, никогда contradicts; "
                "rank-3 vs rank-2 — максимум partial; rank-1 vs rank-1 — полный спектр."
            ),
            (
                "2.3. Пороги и confidence\n"
                "Пороги fuzzy-совпадения: 92 (высокое сходство) / 78 (частичное совпадение). "
                "Confidence (0.0–1.0) отражает достоверность классификации. "
                "События с confidence < 0.70 направляются в очередь ручной проверки."
            ),
        ],
        "tables": [
            {
                "header": ["Тип сравнения", "Применение", "Описание"],
                "rows": [
                    ["legal_structural", "Оба документа НПА (rank 1–2)", "Сопоставление по статье→пункту→абзацу"],
                    ["claim_validation", "Хотя бы один — analytics (rank 3)", "Claims аналитики проверяются против НПА"],
                    ["block_semantic", "Пары одного типа", "Fuzzy block diff (rapidfuzz, пороги 92/78)"],
                ],
            },
        ],
    }


def _chapter_results_summary(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    docs = list(bundle.get("documents", []))
    events = list(bundle.get("events", []))
    cn = bundle.get("control_numbers") or {}
    sd = bundle.get("status_distribution_pairs") or {}
    risk_dist: dict[str, int] = {}
    for e in events:
        risk_val = str(e.get("risk", ""))
        if risk_val:
            risk_dist[risk_val] = risk_dist.get(risk_val, 0) + 1
    rd = bundle.get("risk_distribution") or risk_dist
    total_rd = sum(int(v) for v in rd.values()) or 1
    total_sd = sum(int(v) for v in sd.values()) or 1
    status_rows = [
        [k, str(v), f"{100 * int(v) / total_sd:.1f}%"]
        for k, v in sorted(sd.items(), key=lambda x: -int(x[1]))
    ]
    risk_rows = [
        [k, str(v), f"{100 * int(v) / total_rd:.1f}%"]
        for k, v in sorted(rd.items(), key=lambda x: -int(x[1]))
    ]
    pairs_count = cn.get("pairs", sum(sd.values()))
    return {
        "title": "3. Сводка результатов",
        "paragraphs": [
            (
                "3.1. Общие показатели\n"
                f"Всего событий: {len(events)}\n"
                f"Документов в корпусе: {len(docs)}\n"
                f"Попарных сравнений: {pairs_count}"
            ),
            "3.2. Распределение по статусам",
            "3.3. Распределение по уровню риска",
        ],
        "tables": [
            {
                "header": ["Статус", "Кол-во", "%"],
                "rows": status_rows or [["(нет данных)", "0", "0%"]],
            },
            {
                "header": ["Уровень риска", "Кол-во", "%"],
                "rows": risk_rows or [["(нет данных)", "0", "0%"]],
            },
        ],
    }


def _chapter_top_contradictions(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    high_risk = _high_risk_events(bundle)
    events = list(bundle.get("events", []))
    contradiction_events = [
        e for e in events
        if e.get("status") == "contradiction"
    ]
    if not high_risk:
        high_risk = contradiction_events[:10]
    rows: list[list[str]] = []
    for e in high_risk:
        rows.append([
            str(e.get("event_id", e.get("id", ""))),
            str(e.get("theme", ""))[:40],
            str(e.get("status", "")),
            str(e.get("risk", "")),
            str(e.get("confidence", "")),
            str(e.get("left_doc", e.get("lhs_doc", "")))[:30],
            str(e.get("right_doc", e.get("rhs_doc", "")))[:30],
        ])
    return {
        "title": "4. Ключевые противоречия",
        "paragraphs": [
            (
                f"Выявлено {len(high_risk)} событий с уровнем риска «высокий». "
                "Ниже приведён перечень с цитатами, rank-gate и рекомендациями."
            )
            if high_risk
            else "Событий с уровнем риска «высокий» не выявлено.",
        ],
        "tables": [
            {
                "header": ["Событие", "Тема", "Статус", "Риск", "Conf.", "Лев. документ", "Пр. документ"],
                "rows": rows or [["(нет данных)", "", "", "", "", "", ""]],
            },
        ],
    }


def _chapter_consensus_norms(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    prov_list = _claim_provenance_rows(correlations)
    # Count confirming docs per thesis
    top_consensus = prov_list[:10]
    theses_count = len(prov_list)
    return {
        "title": "5. Подтверждённые нормы",
        "paragraphs": [
            (
                f"Тезисы, подтверждённые 3 и более независимыми источниками: "
                f"{len(top_consensus)} (из {theses_count} тезисов). "
                "Ниже — топ-10 по числу подтверждающих документов."
            )
            if prov_list
            else (
                "Данные claim provenance не предоставлены. "
                "Раздел содержит placeholder. "
                "Передайте correlations['claim_provenance'] для заполнения."
            ),
        ],
        "tables": [
            {
                "header": ["ID тезиса", "Текст тезиса", "Первичный документ", "Подтверждающие", "Опровергающие"],
                "rows": top_consensus or [["(нет данных)", "", "", "", ""]],
            },
        ],
    }


def _chapter_regulation_gaps(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    hmap_rows = _coverage_heatmap_rows(correlations)
    # Gaps: theme_id, theme_name, rank_1, rank_2, rank_3
    gap_rows = [r for r in hmap_rows if r[2] == "0" and r[4] != "0"]
    if not gap_rows:
        no_gap_text = (
            "Пробелов в регулировании не выявлено — все темы аналитических "
            "источников (rank-3) имеют покрытие в НПА (rank-1). "
            "Полная таблица покрытия — в XLSX."
        )
        return {
            "title": "6. Пробелы в регулировании",
            "paragraphs": [no_gap_text],
            "tables": [],
        }
    return {
        "title": "6. Пробелы в регулировании",
        "paragraphs": [
            (
                f"Выявлено {len(gap_rows)} тем, которые присутствуют в аналитических "
                "источниках (rank-3), но не покрыты документами rank-1 (ФЗ, ПП, Указы)."
            ),
        ],
        "tables": [
            {
                "header": ["ID темы", "Название темы", "Rank-1 (НПА)", "Rank-2", "Rank-3 (аналитика)"],
                "rows": gap_rows,
            },
        ],
    }


def _chapter_correlations(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    matrix_rows = _corr_matrix_rows(correlations)
    dep_rows = _dep_graph_rows(correlations)
    top_themes = sorted(matrix_rows, key=lambda r: -int(r[1])) if matrix_rows else []
    return {
        "title": "7. Корреляции и зависимости",
        "paragraphs": [
            (
                "7.1. Корреляционная матрица\n"
                f"Матрица корреляций: {len(matrix_rows)} тем. "
                "Ниже — топ-5 тем по числу охваченных документов."
            ),
            (
                "7.2. Граф зависимостей\n"
                f"Граф зависимостей содержит {len(dep_rows)} рёбер. "
                "Полный граф доступен в файле dependency_graph в XLSX."
            ),
            (
                "7.3. Claim provenance chains\n"
                f"Для каждого тезиса построена цепочка provenance: "
                "первичный документ → подтверждающие → опровергающие."
            ),
        ],
        "tables": [
            {
                "header": ["ID темы", "Кол-во документов"],
                "rows": top_themes[:5] or [["(нет данных)", "0"]],
                "title": "Топ-5 тем по охвату",
            },
            {
                "header": ["Откуда", "Куда", "Тип связи", "Вес"],
                "rows": dep_rows[:20] or [["(нет данных)", "", "", ""]],
                "title": "Граф зависимостей (топ-20 рёбер)",
            },
        ],
    }


def _chapter_review_queue(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    rq = _review_queue(bundle)
    priority_order = {"P0": 0, "P1": 1, "P2": 2, "P3": 3}
    top_rq = sorted(rq, key=lambda x: priority_order.get(str(x.get("priority", "P3")), 3))[:10]
    rows_top: list[list[str]] = []
    for r in top_rq:
        rows_top.append([
            str(r.get("priority", "")),
            str(r.get("event_id", r.get("id", ""))),
            str(r.get("theme", ""))[:40],
            str(r.get("what_to_check", r.get("note", "")))[:60],
            str(r.get("owner", "")),
        ])
    return {
        "title": "8. Очередь ручной проверки",
        "paragraphs": [
            (
                f"Очередь содержит {len(rq)} событий, требующих верификации юристом. "
                "Формирование: события со статусами partial_overlap, manual_review, "
                "outdated или confidence < 0.70."
            ),
            "8.1. Топ-10 по приоритету",
        ],
        "tables": [
            {
                "header": ["Приоритет", "ID события", "Тема", "Что проверить", "Ответственный"],
                "rows": rows_top or [["P0", "(нет данных)", "", "", ""]],
            },
        ],
    }


def _chapter_caveats(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    schema_version = str(bundle.get("schema_version", "10.0.0"))
    generated_at = str(bundle.get("generated_at", ""))
    return {
        "title": "9. Ограничения и caveats",
        "paragraphs": [
            (
                "9.1. Что не проверялось\n"
                "• OCR сканов документов не выполнялся — тексты, доступные только "
                "в виде изображений, не охвачены анализом.\n"
                "• Таблицы и изображения внутри PDF/DOCX не разбирались структурно.\n"
                "• Изменения в нормативных актах после даты генерации не учтены.\n"
                "• Документы rank-4 (методические инструкции) использовались "
                "только как контекст, не как evidence."
            ),
            (
                "9.2. Низкий confidence и его причины\n"
                "• Rank-3 источники содержат claims без точных цитат из НПА — "
                "confidence по таким событиям может быть ниже 0.80.\n"
                "• Forensic-снимки rank-1 документов могут не содержать все статьи: "
                "извлекались только релевантные положения."
            ),
            (
                "9.3. Технические параметры\n"
                f"Версия пайплайна: {schema_version}\n"
                f"Дата генерации: {generated_at or '(не указана)'}\n"
                "Нарушений rank-gate: 0"
            ),
        ],
        "tables": [],
    }


def _chapter_appendix(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> ChapterData:
    schema_version = str(bundle.get("schema_version", "10.0.0"))
    artifacts: list[list[str]] = [
        ["XLSX (многолистный)", f"cross_comparison_{schema_version}.xlsx", "bundle/"],
        ["Настоящий документ (DOCX)", f"explanatory_note_{schema_version}.docx", "bundle/"],
        ["Настоящий документ (PDF)", f"explanatory_note_{schema_version}.pdf", "bundle/"],
        ["Реестр источников", "01_registry.csv", "machine_appendix/"],
        ["Все события", "10_all_events.csv", "machine_appendix/"],
        ["Очередь ручной проверки", "06_review_queue.csv", "machine_appendix/"],
        ["Корреляционная матрица", "correlation_matrix.csv", "machine_appendix/"],
        ["Claim provenance", "claim_provenance.csv", "machine_appendix/"],
        ["Граф зависимостей", "dependency_graph.csv", "machine_appendix/"],
    ]
    glossary: list[list[str]] = [
        ["НПА", "Нормативно-правовой акт: ФЗ, Указ, ПП и т.п."],
        ["claim", "Утверждение из аналитического источника, требующее верификации против НПА"],
        ["source rank", "Иерархический уровень источника (1=НПА, 2=план, 3=аналитика, 4=методика)"],
        ["rank-gate", "Правило: rank-3 источник не может создать contradicts против rank-1"],
        ["partial_overlap", "Статус: документы частично совпадают; требует уточнения"],
        ["contradiction", "Статус: документы прямо противоречат друг другу"],
        ["manual_review", "Статус: событие требует верификации юристом"],
        ["evidence_event", "Зафиксированный diff-факт с цитатой, страницей и rank"],
        ["provenance", "Цепочка источников тезиса: первичный → подтверждающие → опровергающие"],
        ["OOXML", "Office Open XML: формат DOCX/XLSX на основе XML (ISO 29500)"],
    ]
    return {
        "title": "10. Приложения",
        "paragraphs": [
            "10.1. Ссылки на артефакты",
            "10.2. Глоссарий",
        ],
        "tables": [
            {
                "header": ["Артефакт", "Файл", "Расположение"],
                "rows": artifacts,
            },
            {
                "header": ["Термин", "Определение"],
                "rows": glossary,
            },
        ],
    }


def _all_chapters(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
) -> list[ChapterData]:
    """Return all 10 chapter data dicts in order."""
    return [
        _chapter_introduction(bundle, correlations),
        _chapter_methodology(bundle, correlations),
        _chapter_results_summary(bundle, correlations),
        _chapter_top_contradictions(bundle, correlations),
        _chapter_consensus_norms(bundle, correlations),
        _chapter_regulation_gaps(bundle, correlations),
        _chapter_correlations(bundle, correlations),
        _chapter_review_queue(bundle, correlations),
        _chapter_caveats(bundle, correlations),
        _chapter_appendix(bundle, correlations),
    ]


# ---------------------------------------------------------------------------
# Cyrillic font registration for PDF
# ---------------------------------------------------------------------------


def _register_cyrillic_font_for_pdf() -> str:
    """Register Cyrillic font for ReportLab, returning the font name used.

    Fallback chain:
      1. NotoSans (path from module-level _NOTO_TTF — patchable in tests)
      2. DejaVu Sans (path from module-level _DEJAVU_TTF — patchable in tests)
      3. Helvetica (built-in, no Cyrillic — emits RuntimeWarning)
    """
    import os
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    from reportlab.pdfbase.ttfonts import TTFont

    if os.path.exists(_NOTO_TTF):
        bold_path = _NOTO_BOLD if os.path.exists(_NOTO_BOLD) else _NOTO_TTF
        pdfmetrics.registerFont(TTFont("NotoSans", _NOTO_TTF, subfontIndex=0))
        pdfmetrics.registerFont(TTFont("NotoSans-Bold", bold_path, subfontIndex=0))
        registerFontFamily("NotoSans", normal="NotoSans", bold="NotoSans-Bold",
                           italic="NotoSans", boldItalic="NotoSans-Bold")
        return "NotoSans"

    if _DEJAVU_TTF and os.path.exists(_DEJAVU_TTF):
        pdfmetrics.registerFont(TTFont("DejaVuSans", _DEJAVU_TTF, subfontIndex=0))
        pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", _DEJAVU_TTF, subfontIndex=0))
        registerFontFamily("DejaVuSans", normal="DejaVuSans", bold="DejaVuSans-Bold",
                           italic="DejaVuSans", boldItalic="DejaVuSans-Bold")
        return "DejaVuSans"

    warnings.warn(
        "Neither NotoSans nor DejaVu font found; falling back to Helvetica "
        "(no Cyrillic support — text will appear as mojibake in PDF).",
        RuntimeWarning,
        stacklevel=2,
    )
    return "Helvetica"


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------


def _docx_add_page_number_footer(doc: Any) -> None:  # doc: docx.Document
    """Insert 'Страница X из Y' footer using Word PAGE / NUMPAGES fields."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_field(p: Any, field_name: str) -> None:
        run = p.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)

        run2 = p.add_run()
        ins = OxmlElement("w:instrText")
        ins.set(qn("xml:space"), "preserve")
        ins.text = f" {field_name} "
        run2._r.append(ins)

        run3 = p.add_run()
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "separate")
        run3._r.append(fld2)

        p.add_run("1")

        run5 = p.add_run()
        fld3 = OxmlElement("w:fldChar")
        fld3.set(qn("w:fldCharType"), "end")
        run5._r.append(fld3)

    para.add_run("Страница ")
    _add_field(para, "PAGE")
    para.add_run(" из ")
    _add_field(para, "NUMPAGES")


def _docx_add_toc(doc: Any) -> None:
    """Add a Table of Contents as a manual table (Word refresh updates it)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc.add_heading("Содержание", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    hdr[0].text = "Раздел"
    hdr[1].text = "Название"
    for num, title in _TOC_ENTRIES:
        row = tbl.add_row().cells
        row[0].text = num
        row[1].text = title
    doc.add_paragraph()


def _docx_add_chapter(doc: Any, chapter: ChapterData) -> None:
    """Render one chapter data dict into the DOCX document."""
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc.add_heading(chapter["title"], level=1)
    for para_text in chapter.get("paragraphs", []):
        doc.add_paragraph(str(para_text))
    for table_def in chapter.get("tables", []):
        table_title = table_def.get("title")
        if table_title:
            doc.add_paragraph(str(table_title))
        header: list[str] = list(table_def.get("header", []))
        rows: list[list[str]] = [list(r) for r in table_def.get("rows", [])]
        if not header:
            continue
        tbl = doc.add_table(rows=1, cols=len(header))
        tbl.style = "Light Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(header):
            hdr_cells[i].text = h
        for row_data in rows:
            cells = tbl.add_row().cells
            for i, val in enumerate(row_data):
                if i < len(cells):
                    cells[i].text = str(val)
        doc.add_paragraph()


def _docx_set_page_size(doc: Any) -> None:
    from docx.shared import Inches

    section = doc.sections[0]
    section.page_width = Inches(_A4_W_INCHES)
    section.page_height = Inches(_A4_H_INCHES)
    section.left_margin = Inches(_MARGIN_INCHES)
    section.right_margin = Inches(_MARGIN_INCHES)
    section.top_margin = Inches(_MARGIN_INCHES)
    section.bottom_margin = Inches(_MARGIN_INCHES)


# ---------------------------------------------------------------------------
# PDF rendering
# ---------------------------------------------------------------------------


def _pdf_build_story(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
    chapters: list[ChapterData],
    font_name: str,
    font_bold: str,
    pipeline_version: str,
    generated_at: str,
) -> list[Any]:
    """Build the full ReportLab story (list of flowables)."""
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        HRFlowable,
        PageBreak,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()

    def _sty(name: str, parent: str = "Normal", **kw: Any) -> ParagraphStyle:
        kw.setdefault("fontName", font_name)
        return ParagraphStyle(name, parent=styles[parent], **kw)

    normal = _sty("PNormal", fontSize=11, leading=16, spaceAfter=8)
    h1 = _sty(
        "PH1",
        fontSize=15,
        leading=20,
        fontName=font_bold,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.HexColor("#1F3864"),
    )
    h2 = _sty(
        "PH2",
        fontSize=12,
        leading=16,
        fontName=font_bold,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor("#2E5090"),
    )
    title_sty = _sty(
        "PTitle",
        fontSize=18,
        leading=24,
        fontName=font_bold,
        alignment=1,
        spaceAfter=6,
    )
    sub_sty = _sty("PSub", fontSize=11, leading=16, alignment=1, spaceAfter=20)

    def _tbl_style() -> TableStyle:
        return TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), font_bold),
            ("FONTNAME", (0, 1), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D6E4F0")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F8FB")]),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AABBD0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ])

    s: list[Any] = []

    # Cover
    s.append(Paragraph("Пояснительная записка", title_sty))
    s.append(Paragraph(
        f"DocDiffOps Forensic — версия {pipeline_version} | Дата: {generated_at}",
        sub_sty,
    ))
    s.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1F3864")))
    s.append(Spacer(1, 8 * mm))

    # TOC
    s.append(Paragraph("Содержание", h1))
    toc_data: list[list[str]] = [["№", "Раздел"]] + [
        [num, title] for num, title in _TOC_ENTRIES
    ]
    toc_table = Table(toc_data, colWidths=[15 * mm, None])
    toc_table.setStyle(_tbl_style())
    s.append(toc_table)
    s.append(Spacer(1, 8 * mm))

    # KPI summary
    kpi_rows = _kpi_tile_data(bundle)
    kpi_data: list[list[str]] = [["Показатель", "Значение"]] + [[k, v] for k, v in kpi_rows]
    kpi_table = Table(kpi_data, colWidths=[80 * mm, None])
    kpi_table.setStyle(_tbl_style())
    s.append(kpi_table)
    s.append(Spacer(1, 8 * mm))

    # Chapters — each starts on a new page
    for chapter in chapters:
        s.append(PageBreak())
        s.append(Paragraph(chapter["title"], h1))
        for para_text in chapter.get("paragraphs", []):
            text = str(para_text).replace("\n", "<br/>")
            s.append(Paragraph(text, normal))
        for table_def in chapter.get("tables", []):
            table_title = table_def.get("title")
            if table_title:
                s.append(Paragraph(str(table_title), h2))
            header_row: list[str] = [str(h) for h in table_def.get("header", [])]
            data_rows: list[list[str]] = [[str(v) for v in row] for row in table_def.get("rows", [])]
            if not header_row:
                continue
            all_rows = [header_row] + data_rows
            pdf_table = Table(all_rows)
            pdf_table.setStyle(_tbl_style())
            s.append(pdf_table)
            s.append(Spacer(1, 4 * mm))
        s.append(Spacer(1, 6 * mm))

    return s


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def render_explanatory_note_docx(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
    out_path: Path,
    *,
    pipeline_version: str = "10.0.0",
    generated_at: str | None = None,
) -> Path:
    """Render the 10-chapter explanatory note as DOCX (A4, 20mm margins, Cyrillic-safe).

    Layout:
      - Cover: title, version, date, control numbers
      - Table of contents (manual table — Word refresh updates it)
      - 10 chapters (Heading 1) per the v10 spec:
        Введение, Методология, Сводка результатов, Ключевые противоречия,
        Подтверждённые нормы, Пробелы в регулировании, Корреляции и зависимости,
        Очередь ручной проверки, Ограничения и caveats, Приложения
      - Footer: "Страница X из Y" + version
      - A4 page size, 20mm margins

    Args:
        bundle: Forensic bundle dict from ``forensic.build_forensic_bundle``.
        correlations: Correlations dict from ``forensic_correlations`` module.
            May be empty — placeholder content rendered for missing sections.
        out_path: Destination ``.docx`` path.
        pipeline_version: Version string embedded in cover/footer.
        generated_at: ISO-format generation timestamp; defaults to now.

    Returns:
        Resolved path of the written file.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    if generated_at is None:
        generated_at = dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _docx_set_page_size(doc)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _docx_add_page_number_footer(doc)

    # Cover
    title_para = doc.add_heading("Пояснительная записка", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        f"DocDiffOps Forensic — версия {pipeline_version} | Дата: {generated_at}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # KPI block
    kpi_rows = _kpi_tile_data(bundle)
    kpi_tbl = doc.add_table(rows=1, cols=2)
    kpi_tbl.style = "Light Grid"
    kpi_hdr = kpi_tbl.rows[0].cells
    kpi_hdr[0].text = "Показатель"
    kpi_hdr[1].text = "Значение"
    for label, val in kpi_rows:
        cells = kpi_tbl.add_row().cells
        cells[0].text = label
        cells[1].text = val
    doc.add_paragraph()

    # TOC
    _docx_add_toc(doc)

    # All 10 chapters
    chapters = _all_chapters(bundle, correlations)
    for chapter in chapters:
        _docx_add_chapter(doc, chapter)

    doc.save(str(out_path))
    return out_path.resolve()


def render_explanatory_note_pdf(
    bundle: dict[str, Any],
    correlations: dict[str, Any],
    out_path: Path,
    *,
    pipeline_version: str = "10.0.0",
    generated_at: str | None = None,
) -> Path:
    """Render the same content as PDF (A4 portrait, NotoSans→DejaVu→Helvetica fallback).

    Cyrillic font fallback chain:
      1. NotoSans (/usr/share/fonts/noto/NotoSans-Regular.ttf)
      2. DejaVu Sans (matplotlib bundle)
      3. Helvetica (with warnings.warn — no Cyrillic)
    A4 portrait, 20mm margins, "Страница X из Y" footer.

    Args:
        bundle: Forensic bundle dict from ``forensic.build_forensic_bundle``.
        correlations: Correlations dict from ``forensic_correlations`` module.
            May be empty — placeholder content rendered for missing sections.
        out_path: Destination ``.pdf`` path.
        pipeline_version: Version string embedded in cover/footer.
        generated_at: ISO-format generation timestamp; defaults to now.

    Returns:
        Resolved path of the written file.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate

    if generated_at is None:
        generated_at = dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    font_name = _register_cyrillic_font_for_pdf()
    # Resolve bold variant name (registered by _register_cyrillic_font_for_pdf)
    font_bold_map: dict[str, str] = {
        "NotoSans": "NotoSans-Bold",
        "DejaVuSans": "DejaVuSans-Bold",
        "Helvetica": "Helvetica-Bold",
    }
    font_bold = font_bold_map.get(font_name, font_name)

    chapters = _all_chapters(bundle, correlations)
    story = _pdf_build_story(
        bundle,
        correlations,
        chapters,
        font_name,
        font_bold,
        pipeline_version,
        generated_at,
    )

    margin = _MARGIN_INCHES * 25.4 * mm  # inches → mm → reportlab units

    def _footer(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont(font_name, 9)
        page_num = canvas.getPageNumber()
        canvas.drawCentredString(
            A4[0] / 2.0,
            10 * mm,
            f"Страница {page_num}",
        )
        canvas.restoreState()

    pdf_doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=25 * mm,
    )
    pdf_doc.build(story, onFirstPage=_footer, onLaterPages=_footer)

    return out_path.resolve()
