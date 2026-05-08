"""Executive diff DOCX (PR-2.2).

Mirrors the executive_diff.md content but in a Word document so reviewers
who don't read Markdown have a polished artifact. python-docx is already
in requirements (used by render_docx.py for track changes).
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


_RED = RGBColor(0xC0, 0x39, 0x2B)
_AMBER = RGBColor(0xD3, 0x8C, 0x12)
_GRAY = RGBColor(0x55, 0x55, 0x55)


def _h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 0:
        for r in p.runs:
            r.font.size = Pt(20)


def _kv(doc: Document, label: str, value: str | int) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(str(value))


def _severity_color(sev: str | None) -> RGBColor:
    return {"high": _RED, "medium": _AMBER}.get(sev or "low", _GRAY)


def render_executive_docx(
    out_path: Path,
    state: dict[str, Any],
    all_events: list[dict[str, Any]],
    pair_summaries: list[dict[str, Any]],
    top_n: int = 20,
) -> Path:
    """Generate an executive_diff.docx mirroring executive_diff.md."""
    high = [e for e in all_events if e.get("severity") == "high"]
    review = [e for e in all_events if e.get("review_required")]
    partial = [e for e in all_events if e.get("status") == "partial"]
    added = [e for e in all_events if e.get("status") == "added"]
    deleted = [e for e in all_events if e.get("status") == "deleted"]

    doc = Document()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(state.get("title") or state["batch_id"])
    r.bold = True
    r.font.size = Pt(18)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"batch_id: {state['batch_id']}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = _GRAY

    # Summary
    _h(doc, "Сводка", level=1)
    _kv(doc, "Документов", len(state.get("documents", [])))
    _kv(doc, "Пар сравнения", len(pair_summaries))
    _kv(doc, "Diff-событий", len(all_events))
    _kv(doc, "High risk", len(high))
    _kv(doc, "Требуют проверки", len(review))
    _kv(
        doc,
        "Added/Deleted/Partial",
        f"{len(added)} / {len(deleted)} / {len(partial)}",
    )

    # Cache metrics — surfaces PR-1.6 effectiveness for ops.
    extract_hits = sum(
        1 for d in state.get("documents", []) if d.get("cache_extract_hit")
    )
    compare_hits = sum(1 for p in (pair_summaries or []) if p.get("cache_hit"))
    if state.get("documents") or pair_summaries:
        _h(doc, "Cache", level=2)
        _kv(
            doc,
            "Extract hits",
            f"{extract_hits}/{len(state.get('documents', []))}",
        )
        _kv(doc, "Compare hits", f"{compare_hits}/{len(pair_summaries)}")

    # Top high-risk events
    _h(doc, f"Топ high-risk событий ({min(top_n, len(high))} из {len(high)})", level=1)
    if not high:
        doc.add_paragraph("(нет high-severity событий)").italic = True
    for e in high[:top_n]:
        lhs = e.get("lhs") or {}
        rhs = e.get("rhs") or {}
        ph = doc.add_heading(
            f"{e.get('event_id')} — {e.get('status')} / {e.get('severity')}",
            level=3,
        )
        for r in ph.runs:
            r.font.color.rgb = _severity_color(e.get("severity"))

        meta = doc.add_paragraph()
        m1 = meta.add_run("Пара: ")
        m1.bold = True
        meta.add_run(
            f"{e.get('lhs_doc_id')} ↔ {e.get('rhs_doc_id')}    "
        )
        m2 = meta.add_run("score=")
        m2.bold = True
        meta.add_run(str(e.get("score") or "—"))

        if lhs.get("quote"):
            p = doc.add_paragraph()
            p.add_run(f"LHS p.{lhs.get('page_no') or '?'}: ").bold = True
            p.add_run(lhs.get("quote"))
        if rhs.get("quote"):
            p = doc.add_paragraph()
            p.add_run(f"RHS p.{rhs.get('page_no') or '?'}: ").bold = True
            p.add_run(rhs.get("quote"))
        if e.get("explanation_short"):
            doc.add_paragraph(e["explanation_short"]).italic = True

    # Artifacts
    _h(doc, "Артефакты", level=1)
    if not state.get("artifacts"):
        doc.add_paragraph("(нет)").italic = True
    for a in state.get("artifacts", []):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{a.get('type')}: ")
        r.bold = True
        p.add_run(a.get("path") or "")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
