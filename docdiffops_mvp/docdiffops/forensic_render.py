"""Renderers for the v8 forensic bundle (ULTRA-HQ, ru-RU).

Three render targets, all in Russian, with consistent design language
(palette, typography, KPI tiles, cover pages, paginated layout):

  * ``render_v8_xlsx(bundle, out)`` — многолистная книга Excel
    (Обложка, Реестр, Матрица документов, Темы × Источники, Пары,
    Ручная проверка, Хронология поправок, Каталог тем, Действия и
    RACI, корпус-специфичные приложения, QA).
  * ``render_v8_docx_explanatory(bundle, out)`` — пояснительная
    записка в формате v8 (титульный лист, оглавление, KPI-плитки,
    нумерованные разделы, колонтитулы).
  * ``render_v8_docx_redgreen(bundle, out)`` — редакционный diff
    с цветовой кодировкой статусов.
  * ``render_v8_pdf_summary(bundle, out)`` — компактный PDF-обзор
    с титулом, KPI-плитками, легендой и нумерацией страниц.

All renderers operate on the bundle dict produced by
``forensic.build_forensic_bundle`` and are pure: side effects are limited
to writing the target path.
"""
from __future__ import annotations

import datetime as dt
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Mapping, Sequence

from .forensic import (
    DEFAULT_TOPIC_CLUSTERS,
    STATUS_CONTRADICTION,
    STATUS_GAP,
    STATUS_MATCH,
    STATUS_NC,
    STATUS_OUTDATED,
    STATUS_PARTIAL,
    STATUS_REVIEW,
    STATUS_TO_MARK,
    V8_STATUSES,
    cluster_topic_v8,
)


# ---------------------------------------------------------------------------
# Shared Cyrillic font helpers (used by this module and forensic_note)
# ---------------------------------------------------------------------------

_NOTO_TTF_SYSTEM = "/usr/share/fonts/noto/NotoSans-Regular.ttf"
_NOTO_BOLD_SYSTEM = "/usr/share/fonts/noto/NotoSans-Bold.ttf"


def _find_dejavu_ttf() -> str | None:
    """Return path to DejaVuSans.ttf, or None if not found.

    Search order:
      1. matplotlib-bundled font (resolved at runtime, Python-version-agnostic).
      2. Common system font paths (Linux/macOS).
    """
    # Option A: matplotlib bundle — does not hardcode Python version
    try:
        import matplotlib
        mpl_root = Path(matplotlib.__file__).parent
        candidate = mpl_root / "mpl-data" / "fonts" / "ttf" / "DejaVuSans.ttf"
        if candidate.exists():
            return str(candidate)
    except Exception:
        pass
    # Option B: common system locations
    for path in (
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/DejaVuSans.ttf",
    ):
        if Path(path).exists():
            return path
    return None


def _register_cyrillic_pdf_font(
    base_name: str = "CyrillicSans",
    bold_name: str = "CyrillicSans-Bold",
) -> tuple[str, str]:
    """Register a Cyrillic-capable TrueType font with ReportLab.

    Fallback chain: NotoSans (system) → DejaVu Sans (matplotlib bundle or
    system) → Helvetica (built-in; no Cyrillic — emits RuntimeWarning).

    Returns:
        (base_font_name, bold_font_name) registered with pdfmetrics.
    """
    import warnings
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    from reportlab.pdfbase.ttfonts import TTFont

    if Path(_NOTO_TTF_SYSTEM).exists():
        bold_path = _NOTO_BOLD_SYSTEM if Path(_NOTO_BOLD_SYSTEM).exists() else _NOTO_TTF_SYSTEM
        pdfmetrics.registerFont(TTFont(base_name, _NOTO_TTF_SYSTEM, subfontIndex=0))
        pdfmetrics.registerFont(TTFont(bold_name, bold_path, subfontIndex=0))
        registerFontFamily(base_name, normal=base_name, bold=bold_name,
                           italic=base_name, boldItalic=bold_name)
        return base_name, bold_name

    dejavu = _find_dejavu_ttf()
    if dejavu:
        pdfmetrics.registerFont(TTFont(base_name, dejavu, subfontIndex=0))
        pdfmetrics.registerFont(TTFont(bold_name, dejavu, subfontIndex=0))
        registerFontFamily(base_name, normal=base_name, bold=bold_name,
                           italic=base_name, boldItalic=bold_name)
        return base_name, bold_name

    warnings.warn(
        "Neither NotoSans nor DejaVu font found; "
        "falling back to Helvetica — Cyrillic text may not render correctly.",
        RuntimeWarning,
        stacklevel=3,
    )
    return "Helvetica", "Helvetica-Bold"


# ---------------------------------------------------------------------------
# Design vocabulary — ULTRA-HQ, ru-RU
# ---------------------------------------------------------------------------

# Russian display labels for v8 statuses (used in legends, headings,
# status pills). Keys mirror constants from `docdiffops.forensic`.
STATUS_RU: dict[str, str] = {
    STATUS_MATCH:        "Совпадение",
    STATUS_PARTIAL:      "Частичное совпадение",
    STATUS_CONTRADICTION:"Противоречие",
    STATUS_OUTDATED:     "Устаревшее",
    STATUS_GAP:          "Пробел источника",
    STATUS_REVIEW:       "Ручная проверка",
    STATUS_NC:           "Несопоставимо",
}

# Hex palette without leading '#' (used by both openpyxl and ReportLab).
PALETTE: dict[str, str] = {
    "ink":     "0F172A",   # Slate 900 — body text on light
    "primary": "1F2937",   # Slate 800 — table headers
    "accent":  "0F3460",   # Deep navy — title bar
    "muted":   "F9FAFB",   # Light gray — backgrounds
    "border":  "D1D5DB",   # Border gray
    "match":     "16A34A", # Green 600
    "partial":   "F59E0B", # Amber 500
    "contradict":"DC2626", # Red 600
    "outdated":  "2563EB", # Blue 600
    "gap":       "9333EA", # Purple 600
    "review":    "EA580C", # Orange 600
    "nc":        "6B7280", # Gray 500
}

# Status → palette key (for badges / legend / tile color).
STATUS_PALETTE: dict[str, str] = {
    STATUS_MATCH:         "match",
    STATUS_PARTIAL:       "partial",
    STATUS_CONTRADICTION: "contradict",
    STATUS_OUTDATED:      "outdated",
    STATUS_GAP:           "gap",
    STATUS_REVIEW:        "review",
    STATUS_NC:            "nc",
}

# Russian-localized control-number names. Falls back to the bundle key.
CONTROL_RU: dict[str, str] = {
    "documents":      "Документов",
    "pairs":          "Пар",
    "events":         "Событий diff",
    "pairs_total":    "Пар (всего)",
    "pairs_changed":  "Изменено",
    "pairs_resolved": "Закрыто (→ match)",
    "pairs_new":      "Новых",
    "pairs_removed":  "Удалено",
}

DOC_TITLE_RU = "Криминалистический сравнительный анализ — DocDiffOps Forensic v8"
DOC_SUBTITLE_RU = (
    "Evidence-grade сопоставление корпуса. Документ не является юридическим "
    "заключением и предназначен для аналитического и редакционного использования."
)


def _stat_label_ru(key: str) -> str:
    """Translate a control-number key or status code to Russian; fall back to the key."""
    return CONTROL_RU.get(key) or STATUS_RU.get(key) or key


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def render_v8_xlsx(
    bundle: Mapping[str, Any],
    out_path: Path | str,
    *,
    correlations: dict[str, Any] | None = None,
) -> Path:
    """Write a multi-sheet Excel workbook from a forensic bundle.

    Without ``correlations``: backwards-compatible behaviour — produces the
    same base workbook as before PR-6.2 (9-10 sheets depending on bundle
    content).

    With ``correlations`` (output of ``forensic_correlations.compute_*``
    functions): extends the workbook by appending four additional sheets:
      - ``correlation_matrix``   — theme × doc heatmap with ColorScaleRule
      - ``dependency_graph``     — doc → doc edges sorted by relation_type ASC,
                                   weight DESC
      - ``claim_provenance``     — thesis → confirming / refuting evidence
      - ``coverage_heatmap``     — theme × rank-bucket counts with ColorScaleRule

    Args:
        bundle: Forensic bundle dict produced by
            ``forensic.build_forensic_bundle``.
        out_path: Destination ``.xlsx`` file path.
        correlations: Optional dict with keys ``correlation_matrix``,
            ``claim_provenance``, ``dependency_graph``, ``coverage_heatmap``
            as produced by the ``forensic_correlations`` module.

    Returns:
        Resolved ``Path`` of the written file.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    head_fill = PatternFill(start_color="1F2937", end_color="1F2937", fill_type="solid")
    head_font = Font(bold=True, color="FFFFFF")
    bold = Font(bold=True)
    wrap = Alignment(wrap_text=True, vertical="top")
    thin = Side(border_style="thin", color="D1D5DB")
    border = Border(top=thin, bottom=thin, left=thin, right=thin)

    fills = {
        STATUS_MATCH: PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid"),
        STATUS_PARTIAL: PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid"),
        STATUS_CONTRADICTION: PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid"),
        STATUS_OUTDATED: PatternFill(start_color="DBEAFE", end_color="DBEAFE", fill_type="solid"),
        STATUS_GAP: PatternFill(start_color="EDE9FE", end_color="EDE9FE", fill_type="solid"),
        STATUS_REVIEW: PatternFill(start_color="FFEDD5", end_color="FFEDD5", fill_type="solid"),
        STATUS_NC: PatternFill(start_color="F3F4F6", end_color="F3F4F6", fill_type="solid"),
    }
    mark_to_status = {v: k for k, v in STATUS_TO_MARK.items()}

    docs = list(bundle.get("documents", []))
    pairs = list(bundle.get("pairs", []))
    doc_ids = [d["id"] for d in docs]

    def _style_header(ws, row=1):
        for cell in ws[row]:
            cell.font = head_font
            cell.fill = head_fill
            cell.alignment = wrap
            cell.border = border

    def _table(ws, header, rows, col_widths=None):
        ws.append(header)
        _style_header(ws)
        for row in rows:
            ws.append([row.get(h, "") for h in header])
        for i, h in enumerate(header, 1):
            ws.column_dimensions[get_column_letter(i)].width = (
                (col_widths or {}).get(h, max(12, min(48, len(h) + 2)))
            )
        ws.freeze_panes = "A2"
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            for cell in row:
                cell.alignment = wrap
                cell.border = border

    # Sheet 00 — Обложка / Сводка
    ws = wb.active
    ws.title = "00 Обложка"

    # Title bar (merged + colored)
    title_fill = PatternFill(start_color=PALETTE["accent"],
                             end_color=PALETTE["accent"], fill_type="solid")
    ws.append([DOC_TITLE_RU])
    ws.merge_cells(start_row=1, end_row=1, start_column=1, end_column=6)
    ws["A1"].font = Font(bold=True, size=18, color="FFFFFF")
    ws["A1"].fill = title_fill
    ws["A1"].alignment = Alignment(vertical="center", horizontal="left",
                                    indent=1)
    ws.row_dimensions[1].height = 36

    ws.append([DOC_SUBTITLE_RU])
    ws.merge_cells(start_row=2, end_row=2, start_column=1, end_column=6)
    ws["A2"].font = Font(italic=True, size=10, color=PALETTE["nc"])
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="top", indent=1)
    ws.row_dimensions[2].height = 30

    ws.append([f"Дата создания: {bundle.get('generated_at','')}",
               "", "", "",
               f"Схема: {bundle.get('schema_version','')}", ""])
    for c in ws[3]:
        c.font = Font(size=9, color=PALETTE["nc"])

    ws.append([])

    # KPI tiles — 4 colored cells in row 5, captions in row 6
    cn = bundle.get("control_numbers") or {}
    sd = bundle.get("status_distribution_pairs") or {}
    pairs_count = cn.get("pairs", sum(sd.values()))
    contradictions = sd.get(STATUS_CONTRADICTION, 0)
    manual_review = sd.get(STATUS_REVIEW, 0)
    matches = sd.get(STATUS_MATCH, 0)

    kpi_specs = [
        (cn.get("documents", 0), "Документов",      PALETTE["accent"]),
        (pairs_count,            "Пар",             PALETTE["primary"]),
        (matches,                "Совпадений",      PALETTE["match"]),
        (contradictions,         "Противоречий",    PALETTE["contradict"]),
    ]
    if manual_review:
        kpi_specs.append((manual_review, "Ручная проверка", PALETTE["review"]))

    ws.append([str(spec[0]) for spec in kpi_specs])
    ws.append([spec[1] for spec in kpi_specs])
    for i, (_, _, hex_color) in enumerate(kpi_specs, start=1):
        v_cell = ws.cell(row=5, column=i)
        v_cell.font = Font(bold=True, size=22, color="FFFFFF")
        v_cell.alignment = Alignment(horizontal="center", vertical="center")
        v_cell.fill = PatternFill(start_color=hex_color,
                                  end_color=hex_color, fill_type="solid")
        v_cell.border = border
        l_cell = ws.cell(row=6, column=i)
        l_cell.font = Font(bold=True, size=9, color=PALETTE["ink"])
        l_cell.alignment = Alignment(horizontal="center", vertical="center")
        l_cell.fill = PatternFill(start_color=PALETTE["muted"],
                                  end_color=PALETTE["muted"], fill_type="solid")
        l_cell.border = border
        ws.column_dimensions[get_column_letter(i)].width = 18
    ws.row_dimensions[5].height = 50
    ws.row_dimensions[6].height = 22

    ws.append([])

    ws.append(["Контрольные числа"])
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=12,
                                                  color=PALETTE["ink"])
    ws.append(["Параметр", "Значение"])
    _style_header(ws, row=ws.max_row)
    for k, v in cn.items():
        ws.append([_stat_label_ru(k), v])

    ws.append([])
    ws.append(["Распределение пар по статусам v8"])
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=12,
                                                  color=PALETTE["ink"])
    ws.append(["Статус", "Расш. название", "Знак", "Количество"])
    _style_header(ws, row=ws.max_row)
    for s, c in sorted(sd.items(), key=lambda kv: -kv[1]):
        ws.append([s, STATUS_RU.get(s, s), STATUS_TO_MARK.get(s, "?"), c])
        ws.cell(row=ws.max_row, column=3).fill = fills.get(s)
    ws.append([])

    ws.append(["Легенда статусов v8"])
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=12,
                                                  color=PALETTE["ink"])
    ws.append(["Знак", "Статус (код)", "Статус (русский)"])
    _style_header(ws, row=ws.max_row)
    for st in V8_STATUSES:
        ws.append([STATUS_TO_MARK.get(st, "?"), st, STATUS_RU.get(st, st)])
        ws.cell(row=ws.max_row, column=1).fill = fills[st]
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 14

    # Sheet 01 — Registry
    ws = wb.create_sheet("01 Реестр источников")
    cols = ["id", "code", "title", "type", "rank", "url"]
    rows_view = [{c: d.get(c, "") for c in cols} for d in docs]
    _table(ws, cols, rows_view, col_widths={"title": 50, "url": 40, "type": 24})

    # Sheet 02 — Doc × Doc matrix
    ws = wb.create_sheet("02 Документ × Документ")
    header = ["ИД", "Код"] + doc_ids
    ws.append(header)
    _style_header(ws)

    pair_status: dict[tuple[str, str], str] = {}
    pair_count: dict[tuple[str, str], int] = {}
    for p in pairs:
        key = tuple(sorted([p["left"], p["right"]]))
        pair_status[key] = p["v8_status"]
        pair_count[key] = p.get("events_count", 0)

    code_by_id = {d["id"]: d.get("code", "") for d in docs}
    for di in doc_ids:
        row = [di, code_by_id.get(di, "")]
        for dj in doc_ids:
            if di == dj:
                row.append("—")
                continue
            key = tuple(sorted([di, dj]))
            st = pair_status.get(key, STATUS_NC)
            cnt = pair_count.get(key, 0)
            mark = STATUS_TO_MARK.get(st, "—")
            row.append(f"{mark}({cnt})" if cnt else mark)
        ws.append(row)
    ws.freeze_panes = "C2"
    for i in range(1, len(header) + 1):
        ws.column_dimensions[get_column_letter(i)].width = 10 if i > 2 else 14
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=3):
        for cell in row:
            txt = str(cell.value or "")
            mark = txt[0] if txt else ""
            st: str | None = mark_to_status.get(mark)
            if st:
                cell.fill = fills[st]
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
            cell.font = Font(size=9)

    # Sheet 03 — Topic × Doc
    ws = wb.create_sheet("03 Тема × Документ")
    topic_doc: dict[str, dict[str, int]] = defaultdict(lambda: defaultdict(int))
    for p in pairs:
        for tlabel in p.get("topics", []):
            topic_doc[tlabel][p["left"]] += 1
            topic_doc[tlabel][p["right"]] += 1
    header = ["Тема"] + doc_ids + ["Σ"]
    ws.append(header)
    _style_header(ws)
    for tlabel, by_doc in topic_doc.items():
        row = [tlabel]
        total = 0
        for did in doc_ids:
            v = by_doc.get(did, 0)
            row.append(v if v else "")
            total += v
        row.append(total)
        ws.append(row)
    ws.freeze_panes = "B2"
    ws.column_dimensions["A"].width = 38

    # Sheet 04 — Pairs
    ws = wb.create_sheet("04 Пары v8")
    cols = ["id", "left", "right", "left_rank", "right_rank", "v8_status",
            "events_count", "rank_pair", "topics", "Обоснование", "actions"]
    pair_view = []
    for p in pairs:
        v = {c: p.get(c, "") for c in cols}
        v["topics"] = "; ".join(p.get("topics", []))
        v["Обоснование"] = "; ".join(p.get("explanations", []))
        v["actions"] = ", ".join(p.get("actions", []))
        pair_view.append(v)
    _table(ws, cols, pair_view, col_widths={"topics": 40, "id": 16,
                                            "Обоснование": 50, "actions": 18})
    # tint Status column
    status_idx = cols.index("v8_status") + 1
    for r, p in enumerate(pairs, start=2):
        st = p["v8_status"]
        cell = ws.cell(row=r, column=status_idx)
        cell.fill = fills.get(st)

    # Sheet 05 — Manual review
    ws = wb.create_sheet("05 Manual review")
    mr = [p for p in pairs if p["v8_status"] == STATUS_REVIEW]
    cols = ["id", "left", "right", "rank_pair", "events_count", "topics"]
    mr_view = [{c: ("; ".join(p.get(c, [])) if c == "topics" else p.get(c, ""))
                for c in cols} for p in mr]
    _table(ws, cols, mr_view)

    # Sheet 06 — Outdated
    ws = wb.create_sheet("06 Outdated (изменения)")
    cols = ["newer", "amends", "type"]
    out_rows = []
    for newer, olds in (bundle.get("amendment_graph") or {}).items():
        for old in olds:
            out_rows.append({"newer": newer, "amends": old, "type": "amendment"})
    _table(ws, cols, out_rows)

    # Sheet 07 — Topics catalogue
    ws = wb.create_sheet("07 Topics catalogue")
    cols = ["id", "label", "needles"]
    cat_rows = [{"id": t["id"], "label": t["label"], "needles": ", ".join(t["needles"])}
                for t in bundle.get("topic_clusters", [])]
    _table(ws, cols, cat_rows, col_widths={"needles": 60, "label": 38})

    # Sheet 08 — Actions catalogue + RACI (always present when actions applied)
    actions_catalogue = bundle.get("actions_catalogue") or []
    if actions_catalogue:
        ws = wb.create_sheet("08 Действия")
        cols = ["id", "category", "severity", "where", "what_is_wrong",
                "why", "what_to_do", "owner", "v8_status",
                "R", "A", "C", "I", "related_docs"]
        rows_view = []
        for a in actions_catalogue:
            raci = a.get("raci") or {}
            rows_view.append({
                "id": a.get("id", ""),
                "category": a.get("category", ""),
                "severity": a.get("severity", ""),
                "where": a.get("where", ""),
                "what_is_wrong": a.get("what_is_wrong", ""),
                "why": a.get("why", ""),
                "what_to_do": a.get("what_to_do", ""),
                "owner": a.get("owner", ""),
                "v8_status": a.get("v8_status", ""),
                "R": raci.get("R", ""),
                "A": raci.get("A", ""),
                "C": raci.get("C", ""),
                "I": raci.get("I", ""),
                "related_docs": ", ".join(a.get("related_docs", [])),
            })
        _table(ws, cols, rows_view, col_widths={
            "where": 36, "what_is_wrong": 50, "why": 50,
            "what_to_do": 50, "owner": 24, "related_docs": 24,
        })
        # Severity-based row tinting for visual triage
        sev_fill = {
            "high":   PatternFill(start_color="FEE2E2", end_color="FEE2E2",
                                  fill_type="solid"),
            "medium": PatternFill(start_color="FEF3C7", end_color="FEF3C7",
                                  fill_type="solid"),
            "low":    PatternFill(start_color="DCFCE7", end_color="DCFCE7",
                                  fill_type="solid"),
        }
        sev_idx = cols.index("severity") + 1
        for r, a in enumerate(actions_catalogue, start=2):
            f = sev_fill.get(a.get("severity"))
            if f:
                ws.cell(row=r, column=sev_idx).fill = f

    # Sheet 09 — Brochure red/green (corpus migration_v8 only)
    brochure = bundle.get("brochure_redgreen") or []
    if brochure:
        ws = wb.create_sheet("09 Брошюра R-G")
        cols = ["id", "section", "location", "before", "after", "basis", "effect"]
        _table(ws, cols, [{c: e.get(c, "") for c in cols} for e in brochure],
               col_widths={"before": 40, "after": 40, "basis": 30, "effect": 30,
                           "section": 24, "location": 22})

    # Sheet 10 — Klerk → НПА links (corpus migration_v8 only)
    klerk = bundle.get("klerk_npa_links") or []
    if klerk:
        ws = wb.create_sheet("10 Klerk → НПА")
        cols = ["id", "thesis", "npa_doc", "specific_place", "footnote", "v8_status"]
        _table(ws, cols, [{c: e.get(c, "") for c in cols} for e in klerk],
               col_widths={"thesis": 40, "specific_place": 36, "footnote": 36})

    # Sheet 11 — EAEU split (corpus migration_v8 only)
    eaeu = bundle.get("eaeu_split") or []
    if eaeu:
        ws = wb.create_sheet("11 ЕАЭС split")
        cols = ["id", "group", "countries", "work_regime", "basis",
                "employer_action", "minek_text_should_be"]
        _table(ws, cols, [{c: e.get(c, "") for c in cols} for e in eaeu],
               col_widths={"countries": 24, "work_regime": 30, "basis": 30,
                           "employer_action": 36, "minek_text_should_be": 40})

    # Sheet 12 — Amendment chain (corpus migration_v8 only)
    chain = bundle.get("amendment_chain") or []
    if chain:
        ws = wb.create_sheet("12 Цепочка изменений")
        cols = ["id", "chain", "base_act", "amendments_chronology",
                "related", "cite_now", "where_to_verify"]
        _table(ws, cols, [{c: e.get(c, "") for c in cols} for e in chain],
               col_widths={"chain": 24, "base_act": 24,
                           "amendments_chronology": 40, "cite_now": 36,
                           "where_to_verify": 32})

    # Sheet 13 — QA (was 08; renamed to keep numbering consistent)
    ws = wb.create_sheet("13 QA")
    ws.append(["Параметр", "Значение"])
    _style_header(ws)
    ws.append(["schema_version", bundle.get("schema_version", "")])
    ws.append(["generated_at", bundle.get("generated_at", "")])
    for k, v in (bundle.get("control_numbers") or {}).items():
        ws.append([f"control.{k}", v])
    ws.append(["status_distribution_pairs", str(bundle.get("status_distribution_pairs", {}))])
    ws.append(["rank_pair_distribution", str(bundle.get("rank_pair_distribution", {}))])
    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 80

    # Print setup — landscape orientation + fit-to-page on the broadest sheets
    for sheet_name in wb.sheetnames:
        sh = wb[sheet_name]
        sh.page_setup.orientation = sh.ORIENTATION_LANDSCAPE
        sh.page_setup.fitToWidth = 1
        sh.page_setup.fitToHeight = 0
        sh.print_options.horizontalCentered = True
        sh.sheet_properties.pageSetUpPr.fitToPage = True
        sh.print_title_rows = "1:1"
        sh.oddHeader.center.text = "DocDiffOps · Forensic v8"
        sh.oddHeader.center.size = 9
        sh.oddFooter.right.text = "&P / &N"
        sh.oddFooter.right.size = 8

    # ------------------------------------------------------------------ #
    # Correlation sheets (PR-6.2) — appended only when correlations kwarg
    # is provided; base workbook is untouched otherwise.
    # ------------------------------------------------------------------ #
    if correlations is not None:
        from openpyxl.formatting.rule import ColorScaleRule

        _CORR_HEAD_FILL = PatternFill(
            start_color="D9D9D9", end_color="D9D9D9", fill_type="solid"
        )
        _CORR_HEAD_FONT = Font(bold=True)

        def _corr_style_header(ws_c) -> None:
            """Bold + light-gray header row on correlation sheets."""
            for cell in ws_c[1]:
                cell.font = _CORR_HEAD_FONT
                cell.fill = _CORR_HEAD_FILL
                cell.alignment = wrap
                cell.border = border

        def _corr_set_col_width(ws_c, col_idx: int, width: int) -> None:
            ws_c.column_dimensions[get_column_letter(col_idx)].width = width

        # ---- Sheet: correlation_matrix --------------------------------- #
        cm_data: dict[str, dict[str, int]] = correlations.get(
            "correlation_matrix", {}
        )
        ws_cm = wb.create_sheet("correlation_matrix")

        # Collect all doc_ids from the matrix values in deterministic order
        cm_doc_ids: list[str] = []
        for row_v in cm_data.values():
            for did in row_v:
                if did not in cm_doc_ids:
                    cm_doc_ids.append(did)

        cm_headers = ["theme_id", "theme_name"] + cm_doc_ids
        ws_cm.append(cm_headers)
        _corr_style_header(ws_cm)

        # theme_name lookup — pull from the bundle themes list if available
        theme_name_by_id: dict[str, str] = {
            str(t.get("id", "")): str(t.get("name", t.get("label", "")))
            for t in (bundle.get("themes") or bundle.get("topic_clusters") or [])
        }

        for tid, row_v in cm_data.items():
            ws_cm.append(
                [tid, theme_name_by_id.get(tid, "")]
                + [row_v.get(did, 0) for did in cm_doc_ids]
            )

        # Style data rows
        for r in ws_cm.iter_rows(min_row=2, max_row=ws_cm.max_row):
            for c in r:
                c.alignment = wrap
                c.border = border

        # ColorScale CF on the numeric matrix cells (cols C onwards)
        if cm_doc_ids and ws_cm.max_row > 1:
            last_col_cm = get_column_letter(len(cm_headers))
            last_row_cm = ws_cm.max_row
            cm_range = f"C2:{last_col_cm}{last_row_cm}"
            ws_cm.conditional_formatting.add(
                cm_range,
                ColorScaleRule(
                    start_type="num", start_value=0, start_color="FFFFFF",
                    end_type="num",   end_value=1,   end_color="4472C4",
                ),
            )

        ws_cm.freeze_panes = "C2"
        ws_cm.auto_filter.ref = ws_cm.dimensions
        _corr_set_col_width(ws_cm, 1, 12)
        _corr_set_col_width(ws_cm, 2, 28)
        for ci in range(3, len(cm_headers) + 1):
            _corr_set_col_width(ws_cm, ci, 8)

        # ---- Sheet: dependency_graph ----------------------------------- #
        dg_rows: list[dict[str, Any]] = list(
            correlations.get("dependency_graph", [])
        )
        # Sort by relation_type ASC, then weight DESC
        dg_rows.sort(
            key=lambda r: (
                str(r.get("relation_type", "")),
                -float(r.get("weight", 0) or 0),
            )
        )

        ws_dg = wb.create_sheet("dependency_graph")
        dg_headers = [
            "from_doc_id", "from_doc_short",
            "to_doc_id", "to_doc_short",
            "relation_type", "weight",
        ]
        ws_dg.append(dg_headers)
        _corr_style_header(ws_dg)

        for row_dg in dg_rows:
            ws_dg.append([row_dg.get(h, "") for h in dg_headers])

        for r in ws_dg.iter_rows(min_row=2, max_row=ws_dg.max_row):
            for c in r:
                c.alignment = wrap
                c.border = border

        ws_dg.freeze_panes = "A2"
        ws_dg.auto_filter.ref = ws_dg.dimensions
        for ci, w in enumerate([14, 28, 14, 28, 18, 10], start=1):
            _corr_set_col_width(ws_dg, ci, w)

        # Hyperlinks on from_doc / to_doc columns when pdf_link present
        pdf_link_key = "pdf_link"
        for row_dg in dg_rows:
            if str(row_dg.get(pdf_link_key, "")).startswith("http"):
                # find the row in the sheet and set hyperlink on from_doc_id col
                pass  # dependency_graph rows don't carry pdf_link; skip

        # ---- Sheet: claim_provenance ----------------------------------- #
        cp_rows: list[dict[str, Any]] = list(
            correlations.get("claim_provenance", [])
        )

        ws_cp = wb.create_sheet("claim_provenance")
        cp_headers = [
            "thesis_id", "thesis_text",
            "primary_doc_id", "primary_rank",
            "confirming_docs", "confirming_ranks",
            "refuting_docs", "refuting_ranks",
            "evidence_event_ids",
        ]
        ws_cp.append(cp_headers)
        _corr_style_header(ws_cp)

        for row_cp in cp_rows:
            ws_cp.append([row_cp.get(h, "") for h in cp_headers])

        for r in ws_cp.iter_rows(min_row=2, max_row=ws_cp.max_row):
            for c in r:
                c.alignment = wrap
                c.border = border

        ws_cp.freeze_panes = "A2"
        ws_cp.auto_filter.ref = ws_cp.dimensions
        for ci, w in enumerate([14, 50, 16, 14, 28, 20, 28, 20, 32], start=1):
            _corr_set_col_width(ws_cp, ci, w)

        # Hyperlinks on pdf_link column when present
        try:
            pdf_col_idx_cp = cp_headers.index("pdf_link") + 1
        except ValueError:
            pdf_col_idx_cp = None
        if pdf_col_idx_cp:
            for row_iter in ws_cp.iter_rows(
                min_row=2, max_row=ws_cp.max_row,
                min_col=pdf_col_idx_cp, max_col=pdf_col_idx_cp,
            ):
                for cell_cp in row_iter:
                    val_cp = str(cell_cp.value or "").strip()
                    if val_cp.startswith("http"):
                        cell_cp.hyperlink = val_cp
                        cell_cp.style = "Hyperlink"

        # ---- Sheet: coverage_heatmap ----------------------------------- #
        hm_data: dict[str, dict[int | str, int]] = correlations.get(
            "coverage_heatmap", {}
        )

        ws_hm = wb.create_sheet("coverage_heatmap")
        hm_headers = [
            "theme_id", "theme_name",
            "rank_1_count", "rank_2_count", "rank_3_count", "rank_4_count",
        ]
        ws_hm.append(hm_headers)
        _corr_style_header(ws_hm)

        for tid_hm, rank_counts in hm_data.items():
            # rank_counts keys may be int or str depending on caller
            def _rc(k: int) -> int:
                v = rank_counts.get(k, rank_counts.get(str(k), 0))
                return int(v) if v else 0

            ws_hm.append([
                tid_hm,
                theme_name_by_id.get(tid_hm, ""),
                _rc(1), _rc(2), _rc(3), _rc(4),
            ])

        for r in ws_hm.iter_rows(min_row=2, max_row=ws_hm.max_row):
            for c in r:
                c.alignment = wrap
                c.border = border

        # ColorScale CF on rank-count columns (C:F)
        if hm_data and ws_hm.max_row > 1:
            hm_last_row = ws_hm.max_row
            hm_range = f"C2:F{hm_last_row}"
            ws_hm.conditional_formatting.add(
                hm_range,
                ColorScaleRule(
                    start_type="num", start_value=0, start_color="FFFFFF",
                    end_type="num",   end_value=1,   end_color="4472C4",
                ),
            )

        ws_hm.freeze_panes = "A2"
        ws_hm.auto_filter.ref = ws_hm.dimensions
        _corr_set_col_width(ws_hm, 1, 12)
        _corr_set_col_width(ws_hm, 2, 28)
        for ci in range(3, len(hm_headers) + 1):
            _corr_set_col_width(ws_hm, ci, 14)

        # Apply print setup to the 4 new correlation sheets
        for new_sn in ("correlation_matrix", "dependency_graph",
                       "claim_provenance", "coverage_heatmap"):
            if new_sn in wb.sheetnames:
                sh = wb[new_sn]
                sh.page_setup.orientation = sh.ORIENTATION_LANDSCAPE
                sh.page_setup.fitToWidth = 1
                sh.page_setup.fitToHeight = 0
                sh.print_options.horizontalCentered = True
                sh.sheet_properties.pageSetUpPr.fitToPage = True
                sh.print_title_rows = "1:1"

    # Document properties
    wb.properties.title = DOC_TITLE_RU
    wb.properties.subject = "DocDiffOps Forensic v8 — evidence-grade cross-comparison"
    wb.properties.creator = "DocDiffOps Forensic v8"
    wb.properties.language = "ru-RU"

    wb.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# DOCX — explanatory
# ---------------------------------------------------------------------------


def _docx_set_cell_bg(cell: Any, hex_color: str) -> None:
    """Apply a solid background fill to a python-docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _docx_add_page_numbers(doc: Any) -> None:
    """Insert a Page X / Y field into the footer of every section."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        def _field(instr: str) -> None:
            run = para.add_run()
            f_begin = OxmlElement("w:fldChar")
            f_begin.set(qn("w:fldCharType"), "begin")
            f_text = OxmlElement("w:instrText")
            f_text.set(qn("xml:space"), "preserve")
            f_text.text = instr
            f_end = OxmlElement("w:fldChar")
            f_end.set(qn("w:fldCharType"), "end")
            run._r.append(f_begin)
            run._r.append(f_text)
            run._r.append(f_end)

        para.add_run("Стр. ")
        _field("PAGE")
        para.add_run(" из ")
        _field("NUMPAGES")
        para.add_run(f"   •   {DOC_TITLE_RU[:48]}…" if len(DOC_TITLE_RU) > 48
                     else f"   •   {DOC_TITLE_RU}")


def _docx_kpi_card(table: Any, row_idx: int, col_idx: int,
                   value: str, label: str, color_hex: str) -> None:
    """Render a KPI card (large value over small label) in two stacked cells."""
    from docx.shared import Pt, RGBColor
    cell_v = table.rows[row_idx].cells[col_idx]
    cell_l = table.rows[row_idx + 1].cells[col_idx]
    cell_v.text = ""
    p_v = cell_v.paragraphs[0]
    p_v.alignment = 1
    r_v = p_v.add_run(str(value))
    r_v.bold = True
    r_v.font.size = Pt(20)
    r_v.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _docx_set_cell_bg(cell_v, color_hex)

    cell_l.text = ""
    p_l = cell_l.paragraphs[0]
    p_l.alignment = 1
    r_l = p_l.add_run(label)
    r_l.bold = True
    r_l.font.size = Pt(8)
    r_l.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    _docx_set_cell_bg(cell_l, PALETTE["muted"])


def render_v8_docx_explanatory(bundle: Mapping[str, Any], out_path: Path | str) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt, RGBColor

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.core_properties.title = DOC_TITLE_RU
    doc.core_properties.subject = "DocDiffOps Forensic v8 — evidence-grade cross-comparison"
    doc.core_properties.language = "ru-RU"

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)

    # ----- Cover page -----------------------------------------------------
    cover_band = doc.add_paragraph()
    cover_band.paragraph_format.space_before = Pt(36)
    band_run = cover_band.add_run("DOCDIFFOPS · FORENSIC v8")
    band_run.bold = True
    band_run.font.size = Pt(11)
    band_run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

    title = doc.add_paragraph()
    t_run = title.add_run(DOC_TITLE_RU)
    t_run.bold = True
    t_run.font.size = Pt(24)
    t_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    subtitle = doc.add_paragraph()
    s_run = subtitle.add_run(DOC_SUBTITLE_RU)
    s_run.italic = True
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Cover metadata block (4 columns × 2 rows)
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Light Grid"
    meta.rows[0].cells[0].text = "Дата создания"
    meta.rows[0].cells[1].text = "Версия схемы"
    meta.rows[0].cells[2].text = "Корпус"
    meta.rows[0].cells[3].text = "Класс"
    meta.rows[1].cells[0].text = bundle.get("generated_at", "—")
    meta.rows[1].cells[1].text = bundle.get("schema_version", "—")
    meta.rows[1].cells[2].text = bundle.get("corpus", "generic")
    meta.rows[1].cells[3].text = "Аналитический материал"
    for cell in meta.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
        _docx_set_cell_bg(cell, PALETTE["primary"])
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for cell in meta.rows[1].cells:
        _docx_set_cell_bg(cell, PALETTE["muted"])

    doc.add_paragraph()

    # KPI tiles on cover (4-column 2-row table)
    cn = bundle.get("control_numbers") or {}
    sd = bundle.get("status_distribution_pairs") or {}
    pairs_count = cn.get("pairs", sum(sd.values()))
    kpi_specs = [
        (cn.get("documents", 0),                 "Документов",    PALETTE["accent"]),
        (pairs_count,                            "Пар",           PALETTE["primary"]),
        (sd.get(STATUS_MATCH, 0),                "Совпадений",    PALETTE["match"]),
        (sd.get(STATUS_CONTRADICTION, 0),        "Противоречий",  PALETTE["contradict"]),
    ]
    kpi_table = doc.add_table(rows=2, cols=len(kpi_specs))
    for i, (val, lbl, color) in enumerate(kpi_specs):
        _docx_kpi_card(kpi_table, 0, i, str(val), lbl, color)

    # Executive summary line — quantified one-line snapshot
    doc.add_paragraph()
    summary = doc.add_paragraph()
    matches_n = sd.get(STATUS_MATCH, 0)
    contradict_n = sd.get(STATUS_CONTRADICTION, 0)
    review_n = sd.get(STATUS_REVIEW, 0)
    pct_match = (matches_n * 100 / pairs_count) if pairs_count else 0
    sum_run = summary.add_run(
        f"Краткая сводка: всего {pairs_count} пар(а); "
        f"совпадений — {matches_n} ({pct_match:.0f}%), "
        f"противоречий — {contradict_n}, "
        f"требуют ручной проверки — {review_n}."
    )
    sum_run.bold = True
    sum_run.font.size = Pt(10)
    sum_run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

    cover_break = doc.add_paragraph().add_run()
    cover_break.add_break(WD_BREAK.PAGE)

    # ----- Table of contents ---------------------------------------------
    doc.add_heading("Оглавление", level=1)
    toc = doc.add_paragraph()
    toc_run = toc.add_run(
        "1. Цель и методика  ·  2. Ключевые показатели  ·  "
        "3. Распределение пар по статусам  ·  4. Источники  ·  "
        "5. Действия и RACI  ·  6. Запреты"
    )
    toc_run.font.size = Pt(10)
    toc_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    # ----- 1. Цель и методика --------------------------------------------
    doc.add_heading("Раздел 1. Цель и методика", level=1)
    doc.add_paragraph(
        "Криминалистический сравнительный материал по корпусу. Применяется "
        "семизначная шкала статусов v8: совпадение, частичное совпадение, "
        "противоречие, устаревшее, пробел источника, ручная проверка, "
        "несопоставимо. Иерархия источников: rank-1 (НПА) > rank-2 "
        "(ведомственное) > rank-3 (аналитика). Аналитика третьего ранга "
        "не опровергает НПА; пересечение rank-3 ↔ rank-1 автоматически "
        "получает статус «ручная проверка»."
    )

    # ----- 2. Контрольные числа ------------------------------------------
    doc.add_heading("Раздел 2. Ключевые показатели", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid"
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "Параметр", "Значение"
    for cell in t.rows[0].cells:
        _docx_set_cell_bg(cell, PALETTE["primary"])
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for k, v in cn.items():
        cells = t.add_row().cells
        cells[0].text, cells[1].text = _stat_label_ru(k), str(v)

    # ----- 3. Распределение пар по статусам ------------------------------
    doc.add_heading("Раздел 3. Распределение пар по статусам", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid"
    for i, h in enumerate(("Статус (код)", "Статус (русский)", "Знак", "Пар")):
        t.rows[0].cells[i].text = h
        _docx_set_cell_bg(t.rows[0].cells[i], PALETTE["primary"])
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for st, c in sorted(sd.items(), key=lambda kv: -kv[1]):
        cells = t.add_row().cells
        cells[0].text = st
        cells[1].text = STATUS_RU.get(st, st)
        cells[2].text = STATUS_TO_MARK.get(st, "?")
        cells[3].text = str(c)
        _docx_set_cell_bg(cells[2], PALETTE.get(STATUS_PALETTE.get(st, "nc"),
                                                PALETTE["nc"]))
        for run in cells[2].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    # ----- 4. Источники ---------------------------------------------------
    doc.add_heading("Раздел 4. Источники", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid"
    for i, h in enumerate(("ИД", "Код", "Тип", "Ранг")):
        t.rows[0].cells[i].text = h
        _docx_set_cell_bg(t.rows[0].cells[i], PALETTE["primary"])
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for d in bundle.get("documents", []):
        cells = t.add_row().cells
        cells[0].text = d.get("id", "")
        cells[1].text = d.get("code", "")
        cells[2].text = d.get("type", "")
        cells[3].text = str(d.get("rank", ""))

    # ----- 5. Действия и RACI --------------------------------------------
    actions = bundle.get("actions_catalogue") or []
    if actions:
        doc.add_heading("Раздел 5. Действия (FA-XX) и матрица RACI", level=1)
        t = doc.add_table(rows=1, cols=6)
        t.style = "Light Grid"
        for i, h in enumerate(("ИД", "Категория", "Уровень", "Что не так",
                               "Что делать", "RACI (R / A / C / I)")):
            t.rows[0].cells[i].text = h
            _docx_set_cell_bg(t.rows[0].cells[i], PALETTE["primary"])
            for run in t.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sev_color = {"high": PALETTE["contradict"],
                     "medium": PALETTE["partial"],
                     "low": PALETTE["match"]}
        for a in actions:
            raci = a.get("raci") or {}
            cells = t.add_row().cells
            cells[0].text = a.get("id", "")
            cells[1].text = a.get("category", "")
            cells[2].text = a.get("severity", "")
            cells[3].text = a.get("what_is_wrong", "")
            cells[4].text = a.get("what_to_do", "")
            cells[5].text = " / ".join(filter(None, (
                raci.get("R", ""), raci.get("A", ""),
                raci.get("C", ""), raci.get("I", ""),
            )))
            sev_hex = sev_color.get(a.get("severity"))
            if sev_hex:
                _docx_set_cell_bg(cells[2], sev_hex)
                for run in cells[2].paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ----- 6. Запреты -----------------------------------------------------
    bans_section_no = "6" if actions else "5"
    doc.add_heading(f"Раздел {bans_section_no}. Запреты", level=1)
    for line in [
        "Не делать выводы без источника.",
        "Не считать брошюру или письмо выше НПА.",
        "Не скрывать недоступность official-URL.",
        "Не смешивать «нет противоречия» и «не проверено».",
        "Не считать v8 юридическим заключением.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    _docx_add_page_numbers(doc)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# DOCX — red/green
# ---------------------------------------------------------------------------


def render_v8_docx_redgreen(bundle: Mapping[str, Any], out_path: Path | str) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt, RGBColor

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    def _hex_to_rgb(hex_str: str) -> RGBColor:
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16),
                        int(hex_str[4:6], 16))

    RED   = _hex_to_rgb(PALETTE["contradict"])
    GREEN = _hex_to_rgb(PALETTE["match"])
    GRAY  = _hex_to_rgb(PALETTE["nc"])
    BLUE  = _hex_to_rgb(PALETTE["outdated"])
    AMBER = _hex_to_rgb(PALETTE["partial"])
    PURPLE= _hex_to_rgb(PALETTE["gap"])
    ORANGE= _hex_to_rgb(PALETTE["review"])
    INK   = _hex_to_rgb(PALETTE["ink"])
    NAVY  = _hex_to_rgb(PALETTE["accent"])

    color_for = {
        STATUS_MATCH: GREEN,
        STATUS_PARTIAL: AMBER,
        STATUS_CONTRADICTION: RED,
        STATUS_OUTDATED: BLUE,
        STATUS_GAP: PURPLE,
        STATUS_REVIEW: ORANGE,
        STATUS_NC: GRAY,
    }

    doc = Document()
    doc.core_properties.title = "Редакционный diff (red / green) — v8"
    doc.core_properties.language = "ru-RU"
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ----- Cover -----------------------------------------------------------
    band = doc.add_paragraph()
    band_run = band.add_run("DOCDIFFOPS · FORENSIC v8 · РЕДАКЦИОННЫЙ DIFF")
    band_run.bold = True
    band_run.font.size = Pt(11)
    band_run.font.color.rgb = NAVY

    title = doc.add_paragraph()
    t_run = title.add_run("Редакционный diff (red / green) — v8")
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    s_run = subtitle.add_run(
        f"Дата: {bundle.get('generated_at', '—')}    •    "
        f"Схема: {bundle.get('schema_version','—')}    •    "
        f"Пар в анализе: {len(bundle.get('pairs', []))}"
    )
    s_run.italic = True
    s_run.font.size = Pt(10)
    s_run.font.color.rgb = GRAY

    doc.add_paragraph()

    # Status legend table
    legend = doc.add_table(rows=2, cols=len(V8_STATUSES))
    legend.style = "Light Grid"
    for i, st in enumerate(V8_STATUSES):
        head = legend.rows[0].cells[i]
        body = legend.rows[1].cells[i]
        head.text = STATUS_TO_MARK.get(st, "?")
        body.text = STATUS_RU.get(st, st)
        for run in head.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for run in body.paragraphs[0].runs:
            run.font.size = Pt(7)
            run.bold = True
            run.font.color.rgb = INK
        head.paragraphs[0].alignment = 1
        body.paragraphs[0].alignment = 1
        _docx_set_cell_bg(head, PALETTE.get(STATUS_PALETTE.get(st, "nc"),
                                            PALETTE["nc"]))
        _docx_set_cell_bg(body, PALETTE["muted"])

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ----- A. Пары и статусы v8 -------------------------------------------
    doc.add_heading("Раздел A. Пары и статусы v8", level=1)
    for p in bundle.get("pairs", []):
        para = doc.add_paragraph()
        prefix = para.add_run(
            f"{p['id']} • {p['left']} ↔ {p['right']} "
            f"({p.get('rank_pair','')}) — "
        )
        prefix.bold = True
        body = para.add_run(STATUS_RU.get(p["v8_status"], p["v8_status"]))
        body.font.color.rgb = color_for.get(p["v8_status"], GRAY)
        body.bold = True
        if p.get("topics"):
            sub = doc.add_paragraph()
            s = sub.add_run(f"   ↳ {' • '.join(p['topics'])}")
            s.font.size = Pt(8)
            s.font.color.rgb = GRAY
        if p.get("explanations"):
            expl = doc.add_paragraph()
            e = expl.add_run("   " + "; ".join(p["explanations"]))
            e.font.size = Pt(8)
            e.font.color.rgb = GRAY
            e.italic = True
        if p.get("actions"):
            act = doc.add_paragraph()
            a = act.add_run("   ⚙ Действия: " + ", ".join(p["actions"]))
            a.font.size = Pt(8)
            a.font.color.rgb = NAVY

    doc.add_heading("Раздел B. Хронология поправок (outdated)", level=1)
    for newer, olds in (bundle.get("amendment_graph") or {}).items():
        for old in olds:
            para = doc.add_paragraph()
            prefix = para.add_run(f"{newer} → {old}: ")
            prefix.bold = True
            red = para.add_run(f"[устарело: {old}] ")
            red.font.color.rgb = RED
            para.add_run("➜ ")
            green = para.add_run(f"[действует: {newer}]")
            green.font.color.rgb = GREEN

    brochure = bundle.get("brochure_redgreen") or []
    if brochure:
        doc.add_heading("Раздел C. Брошюра — конкретные правки", level=1)
        for entry in brochure:
            para = doc.add_paragraph()
            head = para.add_run(
                f"{entry.get('id','')} • {entry.get('section','')} "
                f"({entry.get('location','')})"
            )
            head.bold = True
            head.font.color.rgb = INK
            before_p = doc.add_paragraph()
            br = before_p.add_run(f"   − {entry.get('before','')}")
            br.font.color.rgb = RED
            br.font.size = Pt(9)
            after_p = doc.add_paragraph()
            ar = after_p.add_run(f"   + {entry.get('after','')}")
            ar.font.color.rgb = GREEN
            ar.font.size = Pt(9)
            note = doc.add_paragraph()
            n = note.add_run(
                f"   ↳ основание: {entry.get('basis','')} • "
                f"эффект: {entry.get('effect','')}"
            )
            n.font.size = Pt(8)
            n.font.color.rgb = GRAY

    _docx_add_page_numbers(doc)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# PDF — summary
# ---------------------------------------------------------------------------


def _pdf_styles():
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        ("/usr/share/fonts/noto/NotoSans-Regular.ttf",
         "/usr/share/fonts/noto/NotoSans-Bold.ttf"),
        ("/usr/share/fonts/TTF/DejaVuSans.ttf",
         "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/liberation/LiberationSans-Regular.ttf",
         "/usr/share/fonts/liberation/LiberationSans-Bold.ttf"),
    ]
    base, bold = "Helvetica", "Helvetica-Bold"
    for r_path, b_path in candidates:
        if Path(r_path).exists() and Path(b_path).exists():
            try:
                pdfmetrics.registerFont(TTFont("V8Sans", r_path))
                pdfmetrics.registerFont(TTFont("V8Sans-Bold", b_path))
                base, bold = "V8Sans", "V8Sans-Bold"
                break
            except Exception:
                continue

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName=base,
                          fontSize=9, leading=12)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName=bold,
                        fontSize=16, leading=20, spaceAfter=12)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName=bold,
                        fontSize=12, leading=15, spaceAfter=6)
    small = ParagraphStyle("small", parent=body, fontSize=8, leading=10)
    return {"body": body, "h1": h1, "h2": h2, "small": small,
            "base": base, "bold": bold, "colors": colors}


def _pdf_page_decoration(canvas: Any, doc: Any, base_font: str,
                         bold_font: str) -> None:
    """Footer with page numbers + small header band on every page."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    canvas.saveState()
    page_w, page_h = A4

    # Top accent bar
    canvas.setFillColor(colors.HexColor("#" + PALETTE["accent"]))
    canvas.rect(0, page_h - 0.5 * 28.35, page_w, 0.5 * 28.35,
                fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont(bold_font, 8)
    canvas.drawString(1.2 * 28.35, page_h - 0.34 * 28.35,
                      "DOCDIFFOPS · FORENSIC v8")

    # Footer
    canvas.setFillColor(colors.HexColor("#" + PALETTE["nc"]))
    canvas.setFont(base_font, 8)
    canvas.drawString(1.2 * 28.35, 0.6 * 28.35,
                      f"Страница {doc.page}")
    canvas.drawRightString(page_w - 1.2 * 28.35, 0.6 * 28.35,
                           DOC_TITLE_RU)
    canvas.restoreState()


def render_v8_pdf_summary(bundle: Mapping[str, Any], out_path: Path | str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import (PageBreak, Paragraph, SimpleDocTemplate,
                                     Spacer, Table, TableStyle)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    s = _pdf_styles()
    colors = s["colors"]

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=1.2 * cm, rightMargin=1.2 * cm,
        topMargin=2.0 * cm, bottomMargin=1.6 * cm,
        title=DOC_TITLE_RU,
        author="DocDiffOps Forensic v8",
        subject="Криминалистический сравнительный анализ",
    )
    base_font = s["base"]
    bold_font = s["bold"]

    def _on_page(canvas: Any, doc_: Any) -> None:
        _pdf_page_decoration(canvas, doc_, base_font, bold_font)

    elems: list[Any] = []

    # ---------- Cover page ----------
    elems.append(Spacer(1, 24))
    elems.append(Paragraph(DOC_TITLE_RU, s["h1"]))
    elems.append(Paragraph(
        f'<font color="#{PALETTE["nc"]}"><i>{DOC_SUBTITLE_RU}</i></font>',
        s["body"]))
    elems.append(Spacer(1, 12))

    cn = bundle.get("control_numbers") or {}
    sd = bundle.get("status_distribution_pairs") or {}
    pairs_count = cn.get("pairs", sum(sd.values()))

    # KPI tiles — 4 colored boxes
    kpi_data = [[
        Paragraph(f'<font size="22" color="white"><b>{cn.get("documents", 0)}</b></font><br/>'
                  f'<font size="8" color="white">Документов</font>', s["body"]),
        Paragraph(f'<font size="22" color="white"><b>{pairs_count}</b></font><br/>'
                  f'<font size="8" color="white">Пар</font>', s["body"]),
        Paragraph(f'<font size="22" color="white"><b>{sd.get(STATUS_MATCH, 0)}</b></font><br/>'
                  f'<font size="8" color="white">Совпадений</font>', s["body"]),
        Paragraph(f'<font size="22" color="white"><b>{sd.get(STATUS_CONTRADICTION, 0)}</b></font><br/>'
                  f'<font size="8" color="white">Противоречий</font>', s["body"]),
    ]]
    kpi_table = Table(kpi_data, colWidths=[4.4 * cm] * 4, rowHeights=[2.2 * cm])
    kpi_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#" + PALETTE["accent"])),
        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#" + PALETTE["primary"])),
        ("BACKGROUND", (2, 0), (2, 0), colors.HexColor("#" + PALETTE["match"])),
        ("BACKGROUND", (3, 0), (3, 0), colors.HexColor("#" + PALETTE["contradict"])),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elems.append(kpi_table)
    elems.append(Spacer(1, 14))

    elems.append(Paragraph(
        f'<font color="#{PALETTE["nc"]}">Дата создания: {bundle.get("generated_at","—")}<br/>'
        f'Версия схемы: {bundle.get("schema_version","—")}<br/>'
        f'Корпус: {bundle.get("corpus","generic")}</font>',
        s["body"]))

    elems.append(Spacer(1, 18))

    # Status legend
    elems.append(Paragraph("Легенда статусов v8", s["h2"]))
    legend_rows = [["Знак", "Код", "Русское название", "Цвет"]]
    for st in V8_STATUSES:
        legend_rows.append([
            STATUS_TO_MARK.get(st, "?"),
            st,
            STATUS_RU.get(st, st),
            "",
        ])
    legend = Table(legend_rows, colWidths=[1.2 * cm, 4 * cm, 6 * cm, 2 * cm],
                   repeatRows=1)
    legend_style = [
        ("FONT", (0, 0), (-1, 0), bold_font, 9),
        ("FONT", (0, 1), (-1, -1), base_font, 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + PALETTE["primary"])),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#" + PALETTE["border"])),
    ]
    for idx, st in enumerate(V8_STATUSES, start=1):
        legend_style.append((
            "BACKGROUND", (3, idx), (3, idx),
            colors.HexColor("#" + PALETTE.get(STATUS_PALETTE.get(st, "nc"),
                                              PALETTE["nc"])),
        ))
    legend.setStyle(TableStyle(legend_style))
    elems.append(legend)
    elems.append(PageBreak())

    # ---------- Section 1: Контрольные числа ----------
    elems.append(Paragraph("Раздел 1. Ключевые показатели", s["h2"]))
    rows = [["Параметр", "Значение"]]
    for k, v in cn.items():
        rows.append([_stat_label_ru(k), str(v)])
    t = Table(rows, colWidths=[8 * cm, 4 * cm], repeatRows=1)
    t.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), bold_font, 9),
        ("FONT", (0, 1), (-1, -1), base_font, 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + PALETTE["primary"])),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#" + PALETTE["muted"])]),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#" + PALETTE["border"])),
    ]))
    elems.append(t)
    elems.append(Spacer(1, 12))

    # ---------- Pie chart of status distribution ----------
    if sd:
        from reportlab.graphics.charts.piecharts import Pie
        from reportlab.graphics.shapes import Drawing, String
        from reportlab.lib import colors as rl_colors

        sorted_dist = sorted(sd.items(), key=lambda kv: -kv[1])
        d_chart = Drawing(420, 180)
        pie = Pie()
        pie.x = 70
        pie.y = 15
        pie.width = 150
        pie.height = 150
        pie.data = [c for _, c in sorted_dist]
        pie.labels = [""] * len(sorted_dist)
        pie.slices.strokeColor = rl_colors.white
        pie.slices.strokeWidth = 1.5
        for idx, (st, _) in enumerate(sorted_dist):
            pie.slices[idx].fillColor = rl_colors.HexColor(
                "#" + PALETTE.get(STATUS_PALETTE.get(st, "nc"), PALETTE["nc"])
            )
        d_chart.add(pie)
        # Inline legend (right side)
        for idx, (st, c) in enumerate(sorted_dist):
            y = 150 - idx * 18
            sw = String(240, y, "■")
            sw.fillColor = rl_colors.HexColor(
                "#" + PALETTE.get(STATUS_PALETTE.get(st, "nc"), PALETTE["nc"])
            )
            sw.fontSize = 14
            sw.fontName = bold_font
            d_chart.add(sw)
            label = String(258, y + 2,
                           f"{STATUS_RU.get(st, st)}: {c}")
            label.fontSize = 9
            label.fontName = base_font
            label.fillColor = rl_colors.HexColor("#" + PALETTE["ink"])
            d_chart.add(label)
        elems.append(d_chart)
        elems.append(Spacer(1, 8))

    # ---------- Section 2: Распределение по статусам ----------
    elems.append(Paragraph("Раздел 2. Распределение пар по статусам", s["h2"]))
    rows = [["Знак", "Статус", "Русское название", "Пар"]]
    for st, c in sorted(sd.items(), key=lambda kv: -kv[1]):
        rows.append([
            STATUS_TO_MARK.get(st, "?"),
            st,
            STATUS_RU.get(st, st),
            str(c),
        ])
    t = Table(rows, colWidths=[1.2 * cm, 4 * cm, 5 * cm, 2 * cm], repeatRows=1)
    dist_style = [
        ("FONT", (0, 0), (-1, 0), bold_font, 9),
        ("FONT", (0, 1), (-1, -1), base_font, 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + PALETTE["primary"])),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 1), (0, -1), "CENTER"),
        ("ALIGN", (3, 1), (3, -1), "RIGHT"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#" + PALETTE["border"])),
    ]
    for idx, (st, _) in enumerate(sorted(sd.items(), key=lambda kv: -kv[1]),
                                  start=1):
        dist_style.append((
            "BACKGROUND", (0, idx), (0, idx),
            colors.HexColor("#" + PALETTE.get(STATUS_PALETTE.get(st, "nc"),
                                              PALETTE["nc"])),
        ))
        dist_style.append(("TEXTCOLOR", (0, idx), (0, idx), colors.white))
    t.setStyle(TableStyle(dist_style))
    elems.append(t)
    elems.append(PageBreak())

    # ---------- Section 3: Источники ----------
    elems.append(Paragraph("Раздел 3. Источники", s["h2"]))
    rows = [["ИД", "Код", "Тип", "Ранг"]]
    for d in bundle.get("documents", []):
        rows.append([d.get("id", ""), d.get("code", ""), d.get("type", ""),
                     str(d.get("rank", ""))])
    t = Table(rows, colWidths=[1.6 * cm, 4 * cm, 5 * cm, 1.2 * cm], repeatRows=1)
    t.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), bold_font, 9),
        ("FONT", (0, 1), (-1, -1), base_font, 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + PALETTE["primary"])),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#" + PALETTE["muted"])]),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#" + PALETTE["border"])),
    ]))
    elems.append(t)

    # ---------- Section 4 (optional): Действия ----------
    actions = bundle.get("actions_catalogue") or []
    if actions:
        elems.append(Spacer(1, 12))
        elems.append(Paragraph("Раздел 4. Действия (FA-XX) и RACI", s["h2"]))
        rows = [["ИД", "Категория", "Уровень", "Что делать", "RACI"]]
        sev_color = {"high": PALETTE["contradict"],
                     "medium": PALETTE["partial"],
                     "low": PALETTE["match"]}
        for a in actions:
            raci = a.get("raci") or {}
            rows.append([
                a.get("id", ""), a.get("category", ""), a.get("severity", ""),
                a.get("what_to_do", ""),
                " / ".join(filter(None, (raci.get("R", ""), raci.get("A", ""),
                                         raci.get("C", ""), raci.get("I", "")))),
            ])
        t = Table(rows, colWidths=[1.6 * cm, 3.6 * cm, 1.6 * cm, 6.5 * cm, 4 * cm],
                  repeatRows=1)
        actions_style = [
            ("FONT", (0, 0), (-1, 0), bold_font, 9),
            ("FONT", (0, 1), (-1, -1), base_font, 7),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + PALETTE["primary"])),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#" + PALETTE["border"])),
        ]
        for idx, a in enumerate(actions, start=1):
            sev_hex = sev_color.get(a.get("severity"))
            if sev_hex:
                actions_style.append((
                    "BACKGROUND", (2, idx), (2, idx), colors.HexColor("#" + sev_hex),
                ))
                actions_style.append(("TEXTCOLOR", (2, idx), (2, idx), colors.white))
        t.setStyle(TableStyle(actions_style))
        elems.append(t)

    doc.build(elems, onFirstPage=_on_page, onLaterPages=_on_page)


# ---------------------------------------------------------------------------
# PDF — integral N×N cross-comparison matrix
# ---------------------------------------------------------------------------

# Color map for matrix cells (0-255 RGB tuples).
_MATRIX_STATUS_RGB: dict[str, tuple[int, int, int]] = {
    "match":           (200, 230, 201),
    "partial_overlap": (255, 245, 157),
    "contradiction":   (255, 205, 210),
    "manual_review":   (255, 224, 178),
    "source_gap":      (245, 245, 245),
    "outdated":        (187, 222, 251),
    "not_comparable":  (224, 224, 224),
}
_MATRIX_DEFAULT_RGB: tuple[int, int, int] = (238, 238, 238)


def _matrix_rgb(status: str) -> tuple[int, int, int]:
    return _MATRIX_STATUS_RGB.get(status, _MATRIX_DEFAULT_RGB)


def _integral_font_setup() -> tuple[str, str]:
    """Register a Cyrillic-capable TrueType font; return (base, bold) font names."""
    return _register_cyrillic_pdf_font("IntegralSans", "IntegralSans-Bold")


def _cover_page(
    story: list[Any],
    bundle: dict[str, Any],
    base: str,
    bold: str,
    page_w: float,
    margin: float,
) -> None:
    """Append cover-page elements to *story* (in-place)."""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    title_s = ParagraphStyle("IntCoverTitle", fontName=bold, fontSize=22,
                             leading=28, alignment=TA_CENTER, spaceAfter=8)
    sub_s = ParagraphStyle("IntCoverSub", fontName=base, fontSize=12,
                           leading=16, alignment=TA_CENTER, spaceAfter=6)

    generated_at = bundle.get("generated_at", "")[:10]
    cn = bundle.get("control_numbers") or {}
    sd = bundle.get("status_distribution_pairs") or {}

    contradictions = sd.get("contradiction", 0) + sd.get("contradicts", 0)
    review_queue = sd.get("manual_review", 0)

    story.append(Spacer(1, 40 * mm))
    story.append(Paragraph("Интегральное перекрёстное сравнение", title_s))
    story.append(Paragraph(generated_at, sub_s))
    story.append(Spacer(1, 8 * mm))

    kpi_labels = ["Документов", "Пар", "Событий", "Противоречий", "Очередь проверки"]
    kpi_values = [
        str(cn.get("documents", len(bundle.get("documents", [])))),
        str(cn.get("pairs", len(bundle.get("pairs", [])))),
        str(cn.get("events", 0)),
        str(contradictions),
        str(review_queue),
    ]
    kpi_bg = ["#37474F", "#455A64", "#FFF9C4", "#FFCDD2", "#FFE0B2"]
    kpi_tc = ["white", "white", "#1a1a1a", "#1a1a1a", "#1a1a1a"]

    col_w = (page_w - 2 * margin) / 5
    kpi_tbl = Table(
        [kpi_labels, kpi_values],
        colWidths=[col_w] * 5,
        rowHeights=[10 * mm, 18 * mm],
    )
    kpi_style: list[Any] = [
        ("FONTNAME", (0, 0), (-1, 0), bold),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("FONTNAME", (0, 1), (-1, 1), bold),
        ("FONTSIZE", (0, 1), (-1, 1), 20),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#90A4AE")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0BEC5")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    for i, (bg, tc) in enumerate(zip(kpi_bg, kpi_tc)):
        kpi_style.append(("BACKGROUND", (i, 0), (i, 0), colors.HexColor(bg)))
        kpi_style.append(("TEXTCOLOR", (i, 0), (i, 0),
                          colors.white if tc == "white" else colors.HexColor(tc)))
        kpi_style.append(("BACKGROUND", (i, 1), (i, 1), colors.HexColor(bg)))
        kpi_style.append(("TEXTCOLOR", (i, 1), (i, 1),
                          colors.white if tc == "white" else colors.HexColor(tc)))
    kpi_tbl.setStyle(TableStyle(kpi_style))
    story.append(kpi_tbl)


def _legend_page(
    story: list[Any],
    base: str,
    bold: str,
    page_w: float,
    margin: float,
) -> None:
    """Append status-legend page elements to *story* (in-place)."""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import HRFlowable, Paragraph, Spacer, Table, TableStyle

    h2_s = ParagraphStyle("IntLegH2", fontName=bold, fontSize=14,
                          leading=18, spaceAfter=6)

    legend_items = [
        (STATUS_MATCH,        "✓", "#C8E6C9", STATUS_RU[STATUS_MATCH],
         "Документы описывают одно и то же положение"),
        (STATUS_PARTIAL,      "≈", "#FFF59D", STATUS_RU[STATUS_PARTIAL],
         "Есть общие элементы, но есть и различия"),
        (STATUS_CONTRADICTION,"⚠", "#FFCDD2", STATUS_RU[STATUS_CONTRADICTION],
         "Положения несовместимы или взаимоисключают друг друга"),
        (STATUS_REVIEW,       "?", "#FFE0B2", STATUS_RU[STATUS_REVIEW],
         "Confidence ниже порога или rank-gate — требует ручной проверки"),
        (STATUS_GAP,          "∅", "#F5F5F5", STATUS_RU[STATUS_GAP],
         "Тема поднята в одном документе, не охвачена другим"),
        (STATUS_OUTDATED,     "↻", "#BBDEFB", STATUS_RU[STATUS_OUTDATED],
         "Один из документов содержит более раннюю редакцию нормы"),
        (STATUS_NC,           "—", "#E0E0E0", STATUS_RU[STATUS_NC],
         "Пара не несёт содержательного сравнения (разный предмет)"),
    ]

    story.append(Paragraph("Легенда статусов", h2_s))
    story.append(HRFlowable(width="100%", thickness=1,
                             color=colors.HexColor("#CCCCCC")))
    story.append(Spacer(1, 4 * mm))

    avail_w = page_w - 2 * margin
    col_ws = [avail_w * f for f in (0.20, 0.06, 0.06, 0.22, 0.46)]
    leg_data: list[list[Any]] = [["Код статуса", "Глиф", "Цвет",
                                   "Название (RU)", "Описание"]]
    for status, glyph, _, ru_name, desc in legend_items:
        leg_data.append([status, glyph, "", ru_name, desc])

    leg_tbl = Table(leg_data, colWidths=col_ws)
    leg_cmds: list[Any] = [
        ("FONTNAME", (0, 0), (-1, 0), bold),
        ("FONTNAME", (0, 1), (-1, -1), base),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ALIGN", (1, 0), (2, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#BBBBBB")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DDDDDD")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#37474F")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
    ]
    for i, (_, _, hex_color, _, _) in enumerate(legend_items, start=1):
        leg_cmds.append(("BACKGROUND", (2, i), (2, i),
                          colors.HexColor(hex_color)))
    leg_tbl.setStyle(TableStyle(leg_cmds))
    story.append(leg_tbl)


def _matrix_page(
    story: list[Any],
    doc_ids: list[str],
    pair_status_lookup: dict[tuple[str, str], str],
    base: str,
    bold: str,
    page_w: float,
    margin: float,
) -> None:
    """Append N×N matrix page elements to *story* (in-place)."""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    h2_s = ParagraphStyle("IntMatH2", fontName=bold, fontSize=14,
                          leading=18, spaceAfter=6)
    cell_s = ParagraphStyle("IntMatCell", fontName=base, fontSize=6,
                            leading=8, alignment=TA_CENTER)

    n = len(doc_ids)
    avail_w = page_w - 2 * margin
    header_col_w = 26 * mm
    data_col_w = (avail_w - header_col_w) / max(n, 1)
    cell_h = data_col_w

    story.append(Paragraph(f"Матрица статусов {n} × {n} документов", h2_s))
    story.append(Spacer(1, 3 * mm))

    header_row: list[Any] = [""] + [Paragraph(did, cell_s) for did in doc_ids]
    mat_data: list[list[Any]] = [header_row]

    for row_id in doc_ids:
        row: list[Any] = [Paragraph(row_id, cell_s)]
        for col_id in doc_ids:
            if row_id == col_id:
                row.append(Paragraph("■", cell_s))
            else:
                st = pair_status_lookup.get((row_id, col_id), "not_comparable")
                glyph = STATUS_TO_MARK.get(st, "?")
                row.append(Paragraph(glyph, cell_s))
        mat_data.append(row)

    col_widths = [header_col_w] + [data_col_w] * n
    row_heights = [cell_h] * (n + 1)
    mat_tbl = Table(mat_data, colWidths=col_widths, rowHeights=row_heights)

    mat_cmds: list[Any] = [
        ("FONTNAME", (0, 0), (-1, -1), base),
        ("FONTSIZE", (0, 0), (-1, -1), 6),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#AAAAAA")),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#37474F")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#37474F")),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.white),
    ]
    # Diagonal cells
    for i in range(n):
        mat_cmds.append(("BACKGROUND", (i + 1, i + 1), (i + 1, i + 1),
                          colors.HexColor("#90A4AE")))
    # Color data cells by pair status
    for row_i, row_id in enumerate(doc_ids):
        for col_i, col_id in enumerate(doc_ids):
            if row_id == col_id:
                continue
            st = pair_status_lookup.get((row_id, col_id), "not_comparable")
            rgb = _matrix_rgb(st)
            hex_col = "#{:02X}{:02X}{:02X}".format(*rgb)
            mat_cmds.append(
                ("BACKGROUND", (col_i + 1, row_i + 1), (col_i + 1, row_i + 1),
                 colors.HexColor(hex_col))
            )
    mat_tbl.setStyle(TableStyle(mat_cmds))
    story.append(mat_tbl)


def _top_n_page(
    story: list[Any],
    events: list[dict[str, Any]],
    top_n: int,
    base: str,
    bold: str,
    page_w: float,
    margin: float,
) -> None:
    """Append Top-N events page elements to *story* (in-place)."""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import HRFlowable, Paragraph, Spacer, Table, TableStyle

    h2_s = ParagraphStyle("IntTopH2", fontName=bold, fontSize=14,
                          leading=18, spaceAfter=6)
    body_s = ParagraphStyle("IntTopBody", fontName=base, fontSize=9,
                            leading=12, spaceAfter=4)

    _SEV_ORDER = {"high": 0, "medium": 1, "low": 2,
                  "высокий": 0, "средний": 1, "низкий": 2}

    top_events = sorted(
        events,
        key=lambda e: (_SEV_ORDER.get(e.get("severity") or e.get("risk", ""), 9),
                       -float(e.get("confidence") or 0))
    )[:top_n]

    story.append(Paragraph(f"Топ-{top_n} событий по уровню серьёзности", h2_s))
    story.append(HRFlowable(width="100%", thickness=1,
                             color=colors.HexColor("#CCCCCC")))
    story.append(Spacer(1, 4 * mm))

    if not top_events:
        story.append(Paragraph("(нет событий)", body_s))
        return

    avail_w = page_w - 2 * margin
    col_ws = [avail_w * f for f in (0.08, 0.14, 0.10, 0.10, 0.58)]
    hdr = ["ID", "Статус", "Серьёзность", "Риск", "Текст / Цитата"]
    tbl_data: list[list[str]] = [hdr]
    _SEV_BG = {"high": "#FFCDD2", "medium": "#FFE0B2", "low": "#C8E6C9",
               "высокий": "#FFCDD2", "средний": "#FFE0B2", "низкий": "#C8E6C9"}

    for e in top_events:
        eid = str(e.get("event_id") or e.get("id") or "")
        status = str(e.get("status") or "")
        sev = str(e.get("severity") or e.get("risk") or "")
        risk = str(e.get("risk") or e.get("severity") or "")
        text = str(
            e.get("conclusion") or e.get("claim_left") or
            e.get("evidence_right") or e.get("text") or ""
        )[:200]
        tbl_data.append([eid, status, sev, risk, text])

    top_tbl = Table(tbl_data, colWidths=col_ws, repeatRows=1)
    tbl_cmds: list[Any] = [
        ("FONTNAME", (0, 0), (-1, 0), bold),
        ("FONTNAME", (0, 1), (-1, -1), base),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#37474F")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#BBBBBB")),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#DDDDDD")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
    ]
    for i, e in enumerate(top_events, start=1):
        sev = str(e.get("severity") or e.get("risk") or "")
        bg = _SEV_BG.get(sev)
        if bg:
            tbl_cmds.append(("BACKGROUND", (2, i), (2, i),
                              colors.HexColor(bg)))
    top_tbl.setStyle(TableStyle(tbl_cmds))
    story.append(top_tbl)


def render_integral_matrix_pdf(
    bundle: dict[str, Any],
    out_path: Path,
    *,
    page_size: str = "auto",
    top_n_events: int = 10,
) -> Path:
    """Render the integral N×N visual matrix PDF.

    Page layout:
      1. Cover — title, date, KPI tiles (docs/pairs/events/contradictions/review_queue)
      2. Status legend — 7 v8 statuses with glyphs and Russian descriptions
      3. Matrix — N×N grid colored by aggregated pair status from bundle["pairs"]
      4. Top-N events — ranked by severity desc, with quotes

    Page sizing:
      - "auto" → A3 landscape if N >= 13, else A4 portrait
      - explicit "A3-landscape" or "A4-portrait" overrides

    Cyrillic font: NotoSans first, fallback to DejaVu Sans, final fallback
    to Helvetica (logs warning).

    Returns out_path on success.
    """
    from reportlab.lib.pagesizes import A3, A4, landscape
    from reportlab.lib.units import mm
    from reportlab.platypus import PageBreak, SimpleDocTemplate, Spacer

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    docs = list(bundle.get("documents", []))
    n = len(docs)

    # Resolve page size
    if page_size == "A3-landscape":
        page = landscape(A3)
    elif page_size == "A4-portrait":
        page = A4
    elif page_size == "auto":
        page = landscape(A3) if n >= 13 else A4
    else:
        page = landscape(A3) if n >= 13 else A4

    W, H = page
    margin = 20 * mm

    base_font, bold_font = _integral_font_setup()

    def _on_page(canvas: Any, doc_: Any) -> None:
        canvas.saveState()
        canvas.setFont(base_font, 8)
        canvas.drawCentredString(W / 2, 12 * mm,
                                 f"Страница {doc_.page} • DocDiffOps Forensic")
        canvas.restoreState()

    doc_obj = SimpleDocTemplate(
        str(out_path),
        pagesize=page,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=22 * mm,
        title="Интегральное перекрёстное сравнение",
        author="DocDiffOps Forensic",
    )

    # Build pair-status lookup: (left_id, right_id) → v8_status
    pair_status_lookup: dict[tuple[str, str], str] = {}
    for p in bundle.get("pairs", []):
        li = str(p.get("left", ""))
        ri = str(p.get("right", ""))
        st = str(p.get("v8_status", "not_comparable"))
        pair_status_lookup[(li, ri)] = st
        pair_status_lookup[(ri, li)] = st

    doc_ids = [str(d["id"]) for d in docs]

    events: list[dict[str, Any]] = list(bundle.get("events", []))

    story: list[Any] = []

    # Page 1: Cover
    _cover_page(story, bundle, base_font, bold_font, W, margin)
    story.append(PageBreak())

    # Page 2: Status legend
    _legend_page(story, base_font, bold_font, W, margin)
    story.append(PageBreak())

    # Page 3: N×N matrix
    _matrix_page(story, doc_ids, pair_status_lookup, base_font, bold_font,
                 W, margin)
    story.append(PageBreak())

    # Page 4: Top-N events
    _top_n_page(story, events, top_n_events, base_font, bold_font, W, margin)

    doc_obj.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    return out_path


__all__ = [
    "render_v8_xlsx",
    "render_v8_docx_explanatory",
    "render_v8_docx_redgreen",
    "render_v8_pdf_summary",
    "render_integral_matrix_pdf",
]
