"""Export a 'merged' DOCX with applied accept/reject decisions.

For each event in the pair:
- confirmed → write the kept side as normal paragraph
- rejected  → write the original (LHS) side as normal paragraph
- pending   → keep as Word track-change (w:ins/w:del) so reviewer can
              still see what's unresolved

The output is a fresh document, not an edit of any original — caller
chains this with download endpoint.
"""

from __future__ import annotations

import datetime as _dt
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt


def _ins(paragraph, text: str, author: str = "DocDiffOps", rev_id: int = 1):
    """Add a w:ins run (track-change insertion) to paragraph."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1f7a3a")
    rPr.append(color)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text or ""
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    ins.append(run)
    paragraph._p.append(ins)


def _del(paragraph, text: str, author: str = "DocDiffOps", rev_id: int = 1):
    """Add a w:del run (track-change deletion) to paragraph."""
    delE = OxmlElement("w:del")
    delE.set(qn("w:id"), str(rev_id))
    delE.set(qn("w:author"), author)
    delE.set(qn("w:date"), _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "8b1d1d")
    strike = OxmlElement("w:strike")
    strike.set(qn("w:val"), "true")
    rPr.append(color)
    rPr.append(strike)
    run.append(rPr)
    delText = OxmlElement("w:delText")
    delText.text = text or ""
    delText.set(qn("xml:space"), "preserve")
    run.append(delText)
    delE.append(run)
    paragraph._p.append(delE)


def _classify_decision(event: dict[str, Any]) -> str:
    """Return 'confirmed' | 'rejected' | 'pending' for the event."""
    lr = event.get("last_review") or {}
    dec = (lr.get("decision") or "").lower()
    if dec == "confirmed":
        return "confirmed"
    if dec == "rejected":
        return "rejected"
    return "pending"


def _sort_key(event: dict[str, Any]):
    """Best-effort document order: lhs_page, then rhs_page, then event_id."""
    lhs = event.get("lhs") or {}
    rhs = event.get("rhs") or {}
    return (
        lhs.get("page_no") or rhs.get("page_no") or 999,
        (lhs.get("bbox") or [0, 0, 0, 0])[1] if isinstance(lhs.get("bbox"), list) else 0,
        event.get("event_id") or "",
    )


def render_merged_docx(
    out_path: Path,
    pair_id: str,
    events: list[dict[str, Any]],
    lhs_doc: dict[str, Any] | None = None,
    rhs_doc: dict[str, Any] | None = None,
) -> dict[str, int]:
    """Render a merged DOCX. Returns counts dict."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    doc.add_heading("Merged document", level=1)
    sub = doc.add_paragraph()
    sub.add_run(f"Pair: {pair_id}").italic = True
    if lhs_doc and rhs_doc:
        sub.add_run(
            f"\nLHS: {lhs_doc.get('filename') or lhs_doc.get('doc_id') or '?'}"
            f"\nRHS: {rhs_doc.get('filename') or rhs_doc.get('doc_id') or '?'}"
        )

    counts: dict[str, int] = {
        "confirmed": 0,
        "rejected": 0,
        "pending": 0,
        "skipped": 0,
        "ambiguous": 0,
    }

    sorted_events = sorted(events or [], key=_sort_key)
    rev_id = 1
    for ev in sorted_events:
        status = (ev.get("status") or "").lower()
        decision = _classify_decision(ev)
        counts[decision] = counts.get(decision, 0) + 1
        lhs_q = ((ev.get("lhs") or {}).get("quote") or "").strip()
        rhs_q = ((ev.get("rhs") or {}).get("quote") or "").strip()

        # status=same — write a single line
        if status == "same":
            if lhs_q or rhs_q:
                doc.add_paragraph(rhs_q or lhs_q)
            continue

        if status in {"contradicts", "manual_review", "not_comparable"}:
            counts["ambiguous"] = counts.get("ambiguous", 0) + 1
            p = doc.add_paragraph()
            warn = p.add_run("⚠ ")
            warn.bold = True
            warn.font.color.rgb = RGBColor(0xE5, 0x48, 0x4D)
            if lhs_q:
                p.add_run(f"LHS: {lhs_q}\n")
            if rhs_q:
                p.add_run(f"RHS: {rhs_q}")
            continue

        # decision-driven branches
        if decision == "confirmed":
            if status == "added":
                if rhs_q:
                    doc.add_paragraph(rhs_q)
            elif status == "deleted":
                counts["skipped"] += 1  # text dropped
            elif status in {"modified", "partial"}:
                if rhs_q:
                    doc.add_paragraph(rhs_q)
        elif decision == "rejected":
            if status == "added":
                counts["skipped"] += 1  # rejected addition → nothing
            elif status == "deleted":
                if lhs_q:
                    doc.add_paragraph(lhs_q)
            elif status in {"modified", "partial"}:
                if lhs_q:
                    doc.add_paragraph(lhs_q)
        else:
            # pending — keep as track-change for further review
            p = doc.add_paragraph()
            if status == "added" and rhs_q:
                _ins(p, rhs_q, rev_id=rev_id)
                rev_id += 1
            elif status == "deleted" and lhs_q:
                _del(p, lhs_q, rev_id=rev_id)
                rev_id += 1
            elif status in {"modified", "partial"}:
                if lhs_q:
                    _del(p, lhs_q, rev_id=rev_id)
                    rev_id += 1
                if rhs_q:
                    _ins(p, rhs_q, rev_id=rev_id)
                    rev_id += 1

    # Audit footer
    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.add_run("—" * 40).italic = True
    audit = doc.add_paragraph()
    a1 = audit.add_run("Applied changes: ")
    a1.bold = True
    audit.add_run(
        f"{counts['confirmed']} confirmed · "
        f"{counts['rejected']} rejected · "
        f"{counts['pending']} pending (kept as track-changes) · "
        f"{counts.get('ambiguous', 0)} ambiguous"
    )

    doc.save(str(out_path))
    return counts
