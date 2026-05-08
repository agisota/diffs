from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _set_space(el):
    el.set(qn("xml:space"), "preserve")


def add_ins_run(p, text: str, author: str = "DocDiffOps", rid: int = 1):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rid))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), datetime.now(timezone.utc).isoformat())
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    _set_space(t)
    t.text = text or ""
    r.append(t)
    ins.append(r)
    p._p.append(ins)


def add_del_run(p, text: str, author: str = "DocDiffOps", rid: int = 1):
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), str(rid))
    dele.set(qn("w:author"), author)
    dele.set(qn("w:date"), datetime.now(timezone.utc).isoformat())
    r = OxmlElement("w:r")
    t = OxmlElement("w:delText")
    _set_space(t)
    t.text = text or ""
    r.append(t)
    dele.append(r)
    p._p.append(dele)


def render_track_changes_docx(out_path: Path, pair_summary: dict[str, Any], events: list[dict[str, Any]], limit: int = 200) -> Path:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    doc.add_heading("DocDiffOps — Track Changes Diff", level=1)
    doc.add_paragraph(f"Pair: {pair_summary.get('lhs_doc_id')} ↔ {pair_summary.get('rhs_doc_id')}")
    doc.add_paragraph(f"Events: {len(events)}. High risk: {pair_summary.get('high_count', 0)}. Review required: {pair_summary.get('review_required_count', 0)}.")

    doc.add_heading("Redline events", level=2)
    rid = 1
    for e in events[:limit]:
        status = e.get("status")
        lhs = e.get("lhs") or {}
        rhs = e.get("rhs") or {}
        doc.add_paragraph(f"{e.get('event_id')} | {status} | severity={e.get('severity')} | score={e.get('score')}", style=None)
        p = doc.add_paragraph()
        if status == "deleted":
            p.add_run("DEL: ").bold = True
            add_del_run(p, lhs.get("quote", ""), rid=rid)
        elif status == "added":
            p.add_run("INS: ").bold = True
            add_ins_run(p, rhs.get("quote", ""), rid=rid)
        elif status in {"partial", "modified", "contradicts"}:
            p.add_run("OLD: ").bold = True
            add_del_run(p, lhs.get("quote", ""), rid=rid)
            p.add_run("\nNEW: ").bold = True
            add_ins_run(p, rhs.get("quote", ""), rid=rid + 1)
        else:
            p.add_run(lhs.get("quote") or rhs.get("quote") or "")
        doc.add_paragraph(e.get("explanation_short", ""))
        rid += 2

    if len(events) > limit:
        doc.add_paragraph(f"Обрезано до {limit} событий. Полная матрица в evidence_matrix.xlsx.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
