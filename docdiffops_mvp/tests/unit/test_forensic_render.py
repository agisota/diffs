"""Smoke tests for forensic v8 renderers (XLSX / DOCX / PDF).

These tests verify the renderer produces a non-empty file whose binary
header matches the expected format and which contains Cyrillic glyphs
extractable by ``pdftotext`` (PDF) or python-docx / openpyxl (DOCX/XLSX).
"""
from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

import pytest

from docdiffops.forensic import build_forensic_bundle


def _sample_bundle():
    docs = [
        {"id": "D01", "code": "FZ_115", "rank": 1, "title": "115-ФЗ",
         "type": "law", "url": "https://example/fz115"},
        {"id": "D02", "code": "FZ_109", "rank": 1, "title": "109-ФЗ",
         "type": "law", "url": "https://example/fz109"},
        {"id": "D03", "code": "ANALYTIC", "rank": 3, "title": "ВЦИОМ-аналитика",
         "type": "analytic", "url": ""},
    ]
    pairs = [
        {"id": "P1", "left": "D01", "right": "D02",
         "events": [{"status": "partial", "topic": "Цифровой профиль"}]},
        {"id": "P2", "left": "D01", "right": "D03",
         "events": [{"status": "same"}]},  # rank3↔rank1 → manual_review
        {"id": "P3", "left": "D02", "right": "D03", "events": []},
    ]
    return build_forensic_bundle(documents=docs, pairs=pairs, events=[],
                                 amendment_graph={"D01": ["D02"]})


def test_render_v8_xlsx_creates_non_empty_workbook(tmp_path: Path):
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    out = tmp_path / "v8.xlsx"
    render_v8_xlsx(_sample_bundle(), out)
    assert out.exists() and out.stat().st_size > 5000

    wb = load_workbook(out)
    assert any("Обложка" in n for n in wb.sheetnames)
    assert any("Реестр" in n for n in wb.sheetnames)
    assert any("Документ" in n for n in wb.sheetnames)
    assert any("Пары" in n for n in wb.sheetnames)


def test_render_v8_docx_explanatory_includes_cyrillic(tmp_path: Path):
    from docx import Document
    from docdiffops.forensic_render import render_v8_docx_explanatory

    out = tmp_path / "explanatory.docx"
    render_v8_docx_explanatory(_sample_bundle(), out)
    assert out.exists() and out.stat().st_size > 3000

    d = Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "v8" in text or "Интегральное" in text or "методика" in text.lower()


def test_render_v8_docx_redgreen_marks_status(tmp_path: Path):
    from docx import Document
    from docdiffops.forensic_render import render_v8_docx_redgreen

    out = tmp_path / "redgreen.docx"
    render_v8_docx_redgreen(_sample_bundle(), out)
    assert out.exists() and out.stat().st_size > 3000

    d = Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    # Russian status labels are now used in the redgreen body
    assert any(t in text for t in (
        "Совпадение", "Частичное совпадение", "Противоречие",
        "Устаревшее", "Ручная проверка", "Несопоставимо",
    ))


@pytest.mark.skipif(shutil.which("pdftotext") is None,
                    reason="pdftotext not installed")
def test_render_v8_pdf_summary_has_cyrillic(tmp_path: Path):
    from docdiffops.forensic_render import render_v8_pdf_summary

    out = tmp_path / "summary.pdf"
    render_v8_pdf_summary(_sample_bundle(), out)
    assert out.exists() and out.stat().st_size > 5000

    r = subprocess.run(["pdftotext", str(out), "-"], capture_output=True)
    txt = r.stdout.decode(errors="replace")
    cyr = sum(1 for c in txt if "Ѐ" <= c <= "ӿ")
    assert cyr > 30, f"Expected Cyrillic chars in PDF, got {cyr}"


# ---------------------------------------------------------------------------
# Step 1.4 — explanations surface in rendered artifacts
# ---------------------------------------------------------------------------


def _bundle_with_explanations():
    from docdiffops.forensic import build_forensic_bundle
    docs = [
        {"id": "D01", "code": "C1", "rank": 1, "title": "t1", "type": "law"},
        {"id": "D02", "code": "C2", "rank": 1, "title": "t2", "type": "law"},
    ]
    pairs = [{"id": "P1", "left": "D01", "right": "D02",
              "events": [{"status": "same", "explanation_short": "Тест обоснования"}]}]
    return build_forensic_bundle(documents=docs, pairs=pairs, events=[], amendment_graph={})


def test_xlsx_explanations_column_present(tmp_path: Path):
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    b = _bundle_with_explanations()
    out = tmp_path / "v8.xlsx"
    render_v8_xlsx(b, out)

    wb = load_workbook(out)
    ws = wb["04 Пары v8"]
    headers = [cell.value for cell in ws[1]]
    assert "Обоснование" in headers

    col_idx = headers.index("Обоснование")
    data_val = list(ws.iter_rows(min_row=2, max_row=2, values_only=True))[0][col_idx]
    assert "Тест обоснования" in (data_val or "")


def test_redgreen_docx_has_explanation_text(tmp_path: Path):
    from docx import Document
    from docdiffops.forensic_render import render_v8_docx_redgreen

    b = _bundle_with_explanations()
    out = tmp_path / "redgreen.docx"
    render_v8_docx_redgreen(b, out)

    d = Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Тест обоснования" in text


# ---------------------------------------------------------------------------
# Action / RACI / corpus-supplementary rendering
# ---------------------------------------------------------------------------


def _bundle_with_actions(corpus: str | None = None):
    from docdiffops.forensic import build_forensic_bundle
    from docdiffops.forensic_actions import apply_actions_to_bundle
    docs = [
        {"id": "D18", "code": "MINEK_BROCHURE", "rank": 2, "title": "брошюра", "type": "brochure"},
        {"id": "D20", "code": "PP_2573", "rank": 1, "title": "ПП 2573", "type": "law"},
    ]
    pairs = [{"id": "P1", "left": "D18", "right": "D20",
              "events": [{"status": "partial"}]}]
    b = build_forensic_bundle(documents=docs, pairs=pairs, events=[],
                              amendment_graph={})
    return apply_actions_to_bundle(b, corpus=corpus)


def test_xlsx_actions_column_in_pairs_sheet(tmp_path: Path):
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    b = _bundle_with_actions()
    out = tmp_path / "v8.xlsx"
    render_v8_xlsx(b, out)
    ws = load_workbook(out)["04 Пары v8"]
    headers = [cell.value for cell in ws[1]]
    assert "actions" in headers


def test_xlsx_actions_sheet_present_when_catalogue_attached(tmp_path: Path):
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    b = _bundle_with_actions()
    out = tmp_path / "v8.xlsx"
    render_v8_xlsx(b, out)
    wb = load_workbook(out)
    assert any("Действия" in n for n in wb.sheetnames)
    # RACI columns exist on the action sheet
    ws = next(wb[n] for n in wb.sheetnames if "Действия" in n)
    headers = [c.value for c in ws[1]]
    assert "R" in headers and "A" in headers and "C" in headers and "I" in headers


def test_xlsx_corpus_supplementaries_only_when_migration_v8(tmp_path: Path):
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    # Generic batch — no corpus content
    generic = _bundle_with_actions(corpus=None)
    out_g = tmp_path / "generic.xlsx"
    render_v8_xlsx(generic, out_g)
    wb_g = load_workbook(out_g)
    assert not any("Брошюра" in n for n in wb_g.sheetnames)
    assert not any("ЕАЭС" in n for n in wb_g.sheetnames)

    # Migration corpus — sheets appear
    mig = _bundle_with_actions(corpus="migration_v8")
    out_m = tmp_path / "migration.xlsx"
    render_v8_xlsx(mig, out_m)
    wb_m = load_workbook(out_m)
    assert any("Брошюра" in n for n in wb_m.sheetnames)
    assert any("Klerk" in n for n in wb_m.sheetnames)
    assert any("ЕАЭС" in n for n in wb_m.sheetnames)
    assert any("Цепочка" in n for n in wb_m.sheetnames)


def test_explanatory_docx_has_actions_section_when_catalogue_attached(tmp_path: Path):
    from docx import Document
    from docdiffops.forensic_render import render_v8_docx_explanatory

    b = _bundle_with_actions()
    out = tmp_path / "explanatory.docx"
    render_v8_docx_explanatory(b, out)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Действия" in text


def test_redgreen_docx_brochure_section_only_with_migration_corpus(tmp_path: Path):
    from docx import Document
    from docdiffops.forensic_render import render_v8_docx_redgreen

    out_g = tmp_path / "generic_rg.docx"
    render_v8_docx_redgreen(_bundle_with_actions(corpus=None), out_g)
    text_g = "\n".join(p.text for p in Document(str(out_g)).paragraphs)
    assert "Брошюра" not in text_g

    out_m = tmp_path / "migration_rg.docx"
    render_v8_docx_redgreen(_bundle_with_actions(corpus="migration_v8"), out_m)
    text_m = "\n".join(p.text for p in Document(str(out_m)).paragraphs)
    assert "Брошюра" in text_m


# ---------------------------------------------------------------------------
# ULTRA-HQ design — cover pages, KPI tiles, Russian language coverage
# ---------------------------------------------------------------------------


def test_explanatory_docx_has_cover_title_and_russian_sections(tmp_path: Path):
    from docx import Document
    from docdiffops.forensic_render import render_v8_docx_explanatory

    out = tmp_path / "exp.docx"
    render_v8_docx_explanatory(_sample_bundle(), out)
    d = Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    # Cover title (Russian)
    assert "Криминалистический сравнительный анализ" in text
    # Numbered Russian sections (ULTRA-HQ "Раздел N." headings)
    assert "Раздел 1." in text
    assert "Раздел 2." in text
    assert "Раздел 3." in text
    assert "Раздел 4." in text
    # TOC marker
    assert "Оглавление" in text


def test_explanatory_docx_sets_russian_language_property(tmp_path: Path):
    from docx import Document
    from docdiffops.forensic_render import render_v8_docx_explanatory

    out = tmp_path / "exp.docx"
    render_v8_docx_explanatory(_sample_bundle(), out)
    d = Document(str(out))
    assert d.core_properties.language == "ru-RU"
    assert "DocDiffOps" in (d.core_properties.title or "")


def test_redgreen_docx_has_cover_band_and_legend(tmp_path: Path):
    from docx import Document
    from docdiffops.forensic_render import render_v8_docx_redgreen

    out = tmp_path / "rg.docx"
    render_v8_docx_redgreen(_sample_bundle(), out)
    d = Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "DOCDIFFOPS · FORENSIC v8 · РЕДАКЦИОННЫЙ DIFF" in text
    # Status legend cells contain Russian names
    table_text = " ".join(
        cell.text for tbl in d.tables for row in tbl.rows for cell in row.cells
    )
    assert "Совпадение" in table_text


def test_xlsx_cover_has_kpi_tiles_and_russian_legend(tmp_path: Path):
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    out = tmp_path / "cover.xlsx"
    render_v8_xlsx(_sample_bundle(), out)
    wb = load_workbook(out)
    assert any("Обложка" in n for n in wb.sheetnames)
    ws = next(wb[n] for n in wb.sheetnames if "Обложка" in n)
    # Cover title is in A1
    assert "Криминалистический" in (ws["A1"].value or "")
    # Russian status names appear in legend
    all_cells = [c.value for row in ws.iter_rows() for c in row]
    text = " ".join(str(v) for v in all_cells if v)
    assert "Совпадение" in text
    assert "Противоречие" in text


def test_pdf_summary_has_russian_title_and_kpi(tmp_path: Path):
    import shutil
    import subprocess
    if shutil.which("pdftotext") is None:
        pytest.skip("pdftotext not installed")
    from docdiffops.forensic_render import render_v8_pdf_summary

    out = tmp_path / "summary.pdf"
    render_v8_pdf_summary(_sample_bundle(), out)
    txt = subprocess.run(["pdftotext", str(out), "-"],
                         capture_output=True).stdout.decode(errors="replace")
    assert "Криминалистический" in txt
    assert "Документов" in txt
    assert "Раздел 1." in txt
    assert "Легенда статусов" in txt


# ---------------------------------------------------------------------------
# PR-6.4 — render_integral_matrix_pdf tests
# ---------------------------------------------------------------------------


def _make_integral_bundle(n_docs: int) -> dict:
    """Build a synthetic forensic bundle with *n_docs* documents."""
    from docdiffops.forensic import build_forensic_bundle

    docs = [
        {"id": f"D{i:02d}", "code": f"DOC{i:02d}", "rank": 1 + (i % 3),
         "title": f"Документ {i:02d}", "type": "law"}
        for i in range(1, n_docs + 1)
    ]
    # Create a handful of pairs with varied statuses to make the matrix interesting
    statuses = ["same", "partial", "contradicts", "same", "manual_review"]
    pairs = []
    pid = 1
    for i in range(1, min(n_docs, 6)):
        for j in range(i + 1, min(n_docs, i + 4)):
            st = statuses[pid % len(statuses)]
            pairs.append({
                "id": f"P{pid}",
                "left": f"D{i:02d}",
                "right": f"D{j:02d}",
                "events": [{"status": st}],
            })
            pid += 1

    # Create synthetic events for top-N page
    events_raw = [
        {"event_id": f"E{k:03d}", "status": "contradiction",
         "severity": "high" if k % 3 == 0 else ("medium" if k % 3 == 1 else "low"),
         "risk": "высокий" if k % 3 == 0 else "средний",
         "confidence": 0.9 - k * 0.01,
         "conclusion": f"Вывод события {k:03d}: обнаружено расхождение в положениях."}
        for k in range(1, 16)
    ]

    bundle = build_forensic_bundle(
        documents=docs, pairs=pairs,
        events=events_raw, amendment_graph={}
    )
    # Embed raw events for top-N page (build_forensic_bundle stores control_numbers
    # but not the events list itself; add them manually for the renderer)
    bundle["events"] = events_raw
    return bundle


@pytest.mark.skipif(shutil.which("pdfinfo") is None,
                    reason="pdfinfo not installed")
def test_integral_matrix_27x27_a3_auto(tmp_path: Path):
    """27 docs with page_size='auto' must select A3 landscape."""
    from docdiffops.forensic_render import render_integral_matrix_pdf

    bundle = _make_integral_bundle(27)
    out = tmp_path / "integral_27.pdf"
    render_integral_matrix_pdf(bundle, out, page_size="auto")
    assert out.exists() and out.stat().st_size > 5000

    info = subprocess.run(["pdfinfo", str(out)], capture_output=True)
    txt = info.stdout.decode(errors="replace")
    # A3 landscape: ~841 × 1190 mm → pdfinfo reports in pts ≈ 2383.94 x 1683.78
    # pdfinfo "Page size:" line e.g. "Page size:      1190.55 x 841.89 pts (A3)"
    # or "1683.78 x 1190.55 pts" depending on orientation reporting
    # We just verify the page is wider than tall (landscape) and larger than A4
    import re
    m = re.search(r"Page size:\s*([\d.]+)\s*x\s*([\d.]+)", txt)
    assert m, f"pdfinfo did not report page size: {txt}"
    w, h = float(m.group(1)), float(m.group(2))
    # A3 landscape: width > height; width should be ~1190 pts
    assert w > h, f"Expected landscape (w>h), got w={w} h={h}"
    assert w > 900, f"Expected A3-width > 900 pts, got {w}"


@pytest.mark.skipif(shutil.which("pdfinfo") is None,
                    reason="pdfinfo not installed")
def test_integral_matrix_8x8_a4_auto(tmp_path: Path):
    """8 docs with page_size='auto' must select A4 portrait."""
    from docdiffops.forensic_render import render_integral_matrix_pdf

    bundle = _make_integral_bundle(8)
    out = tmp_path / "integral_8.pdf"
    render_integral_matrix_pdf(bundle, out, page_size="auto")
    assert out.exists() and out.stat().st_size > 3000

    info = subprocess.run(["pdfinfo", str(out)], capture_output=True)
    txt = info.stdout.decode(errors="replace")
    import re
    m = re.search(r"Page size:\s*([\d.]+)\s*x\s*([\d.]+)", txt)
    assert m, f"pdfinfo did not report page size: {txt}"
    w, h = float(m.group(1)), float(m.group(2))
    # A4 portrait: height > width; width ~595 pts
    assert h > w, f"Expected portrait (h>w), got w={w} h={h}"
    assert w < 700, f"Expected A4-width < 700 pts, got {w}"


@pytest.mark.skipif(shutil.which("pdftotext") is None,
                    reason="pdftotext not installed")
def test_integral_matrix_status_legend_present(tmp_path: Path):
    """Page 2 (legend) must contain at least 5 of the 7 Russian status names."""
    from docdiffops.forensic_render import render_integral_matrix_pdf

    bundle = _make_integral_bundle(8)
    out = tmp_path / "integral_legend.pdf"
    render_integral_matrix_pdf(bundle, out, page_size="A4-portrait")
    assert out.exists()

    # Extract page 2 only using pdftotext -f 2 -l 2
    r = subprocess.run(
        ["pdftotext", "-f", "2", "-l", "2", str(out), "-"],
        capture_output=True,
    )
    txt = r.stdout.decode(errors="replace")

    expected_names = [
        "Совпадение", "Частичное", "Противоречие",
        "Устаревшее", "Пробел", "Ручная проверка", "Несопоставимо",
    ]
    found = sum(1 for name in expected_names if name in txt)
    assert found >= 5, (
        f"Expected at least 5 status names on legend page, found {found}. "
        f"Page 2 text snippet: {txt[:500]}"
    )


@pytest.mark.skipif(shutil.which("pdftotext") is None,
                    reason="pdftotext not installed")
def test_integral_matrix_top_n_events(tmp_path: Path):
    """top_n_events=5 → page 4 must contain exactly 5 event rows (E001..E005 IDs)."""
    from docdiffops.forensic_render import render_integral_matrix_pdf

    bundle = _make_integral_bundle(8)
    out = tmp_path / "integral_topn.pdf"
    render_integral_matrix_pdf(bundle, out, page_size="A4-portrait", top_n_events=5)
    assert out.exists()

    # Extract page 4 only
    r = subprocess.run(
        ["pdftotext", "-f", "4", "-l", "4", str(out), "-"],
        capture_output=True,
    )
    txt = r.stdout.decode(errors="replace")

    # Our synthetic events have IDs E001..E015; top-5 by severity should be
    # E001,E004,E007,E010,E013 (high severity every 3rd, k%3==0 from k=1 means k=3,6,9,12,15)
    # Actually k%3==0: k=3,6,9,12,15 are high. k%3==1: k=1,4,7,10,13 are medium.
    # Sorted by severity: high first, then medium.
    # Simplest check: count event IDs present in the text
    import re
    event_ids_found = set(re.findall(r"E\d{3}", txt))
    assert len(event_ids_found) == 5, (
        f"Expected exactly 5 event IDs on page 4, found {len(event_ids_found)}: "
        f"{event_ids_found}. Page 4 snippet: {txt[:500]}"
    )


@pytest.mark.skipif(shutil.which("pdftotext") is None,
                    reason="pdftotext not installed")
def test_integral_matrix_cyrillic_renders(tmp_path: Path):
    """Full PDF must contain the Cyrillic word 'Интегральное' (validates font)."""
    from docdiffops.forensic_render import render_integral_matrix_pdf

    bundle = _make_integral_bundle(5)
    out = tmp_path / "integral_cyr.pdf"
    render_integral_matrix_pdf(bundle, out, page_size="A4-portrait")
    assert out.exists()

    r = subprocess.run(["pdftotext", str(out), "-"], capture_output=True)
    txt = r.stdout.decode(errors="replace")
    assert "Интегральное" in txt, (
        f"Expected 'Интегральное' in PDF text. Got snippet: {txt[:300]}"
    )


# ---------------------------------------------------------------------------
# PR-6.2 — XLSX 10 → 14 sheets (correlation extension)
# ---------------------------------------------------------------------------


def _make_correlations_bundle():
    """Return (bundle, correlations) with ≥3 themes, ≥4 docs, ≥3 events."""
    from docdiffops.forensic import build_forensic_bundle
    from docdiffops.forensic_correlations import (
        compute_correlation_matrix,
        compute_claim_provenance,
        compute_dependency_graph,
        compute_coverage_heatmap,
    )

    docs = [
        {"id": "D01", "code": "LAW1", "rank": 1, "title": "Закон 1", "type": "law"},
        {"id": "D02", "code": "LAW2", "rank": 1, "title": "Закон 2", "type": "law"},
        {"id": "D03", "code": "PP1",  "rank": 2, "title": "ПП 1",   "type": "regulation"},
        {"id": "D04", "code": "ANA1", "rank": 3, "title": "Аналитика", "type": "analytic"},
    ]
    pairs = [
        {"id": "P1", "left": "D01", "right": "D02",
         "events": [{"status": "same"}]},
        {"id": "P2", "left": "D01", "right": "D03",
         "events": [{"status": "partial"}]},
        {"id": "P3", "left": "D02", "right": "D04",
         "events": [{"status": "contradicts"}]},
    ]
    bundle = build_forensic_bundle(
        documents=docs, pairs=pairs, events=[], amendment_graph={}
    )

    # topic_clusters acts as themes in the bundle
    themes = [
        {"id": "T01", "name": "Тема A", "label": "Тема A", "needles": ["a"]},
        {"id": "T02", "name": "Тема B", "label": "Тема B", "needles": ["b"]},
        {"id": "T03", "name": "Тема C", "label": "Тема C", "needles": ["c"]},
    ]
    theme_doc_links = [
        {"theme_id": "T01", "doc_id": "D01", "status": "match"},
        {"theme_id": "T01", "doc_id": "D02", "status": "partial_overlap"},
        {"theme_id": "T02", "doc_id": "D03", "status": "match"},
        {"theme_id": "T03", "doc_id": "D01", "status": "match"},
        {"theme_id": "T03", "doc_id": "D04", "status": "match"},
    ]
    theses = [
        {"thesis_id": "TH1", "thesis": "Тезис первый",
         "theme": "T01", "coordinate": "D01 стр. 1"},
        {"thesis_id": "TH2", "thesis": "Тезис второй",
         "theme": "T02", "coordinate": "D03 стр. 5"},
        {"thesis_id": "TH3", "thesis": "Тезис третий",
         "theme": "T03", "coordinate": ""},
    ]
    events_raw = [
        {"event_id": "E001", "theme_id": "T01", "theme": "T01",
         "left_id": "D01", "right_id": "D02",
         "status": "match", "source_rank_left": "1", "source_rank_right": "1"},
        {"event_id": "E002", "theme_id": "T02", "theme": "T02",
         "left_id": "D03", "right_id": "D04",
         "status": "contradiction", "source_rank_left": "2", "source_rank_right": "3"},
        {"event_id": "E003", "theme_id": "T03", "theme": "T03",
         "left_id": "D01", "right_id": "D04",
         "status": "partial_overlap", "source_rank_left": "1", "source_rank_right": "3"},
    ]
    pair_relations = [
        {"left_id": "D01", "right_id": "D02", "comparison_type": "тематическое сопоставление",
         "status": "match", "relevance": "высокая", "left": "LAW1", "right": "LAW2"},
        {"left_id": "D01", "right_id": "D03", "comparison_type": "provenance/архивное сопоставление",
         "status": "partial_overlap", "relevance": "средняя", "left": "LAW1", "right": "PP1"},
        {"left_id": "D02", "right_id": "D04", "comparison_type": "методический/forensic контекст",
         "status": "contradiction", "relevance": "высокая", "left": "LAW2", "right": "ANA1"},
    ]

    corr_matrix = compute_correlation_matrix(themes, docs, theme_doc_links)
    claim_prov = compute_claim_provenance(theses, events_raw, docs)
    dep_graph = compute_dependency_graph(pair_relations, docs)
    coverage_hm = compute_coverage_heatmap(corr_matrix, docs)

    correlations = {
        "correlation_matrix": corr_matrix,
        "claim_provenance": claim_prov,
        "dependency_graph": dep_graph,
        "coverage_heatmap": coverage_hm,
    }
    return bundle, correlations


# ---- Test 1: backwards-compat — no correlations → same base sheet count ----


def test_xlsx_default_10_sheets_unchanged(tmp_path: Path):
    """render_v8_xlsx without correlations kwarg must produce the same
    base sheets as before PR-6.2 (backwards-compat guarantee)."""
    from openpyxl import load_workbook
    from docdiffops.forensic import build_forensic_bundle
    from docdiffops.forensic_actions import apply_actions_to_bundle
    from docdiffops.forensic_render import render_v8_xlsx

    docs = [
        {"id": "D18", "code": "MINEK", "rank": 2, "title": "брошюра", "type": "brochure"},
        {"id": "D20", "code": "PP2573", "rank": 1, "title": "ПП 2573", "type": "law"},
    ]
    pairs = [{"id": "P1", "left": "D18", "right": "D20",
              "events": [{"status": "partial"}]}]
    bundle = build_forensic_bundle(documents=docs, pairs=pairs, events=[],
                                   amendment_graph={})
    bundle = apply_actions_to_bundle(bundle)  # adds 08 Действия → 10 sheets

    out = tmp_path / "base.xlsx"
    render_v8_xlsx(bundle, out)  # no correlations kwarg

    wb = load_workbook(out)
    expected_names = [
        "00 Обложка",
        "01 Реестр источников",
        "02 Документ × Документ",
        "03 Тема × Документ",
        "04 Пары v8",
        "05 Manual review",
        "06 Outdated (изменения)",
        "07 Topics catalogue",
        "08 Действия",
        "13 QA",
    ]
    assert wb.sheetnames == expected_names, (
        f"Expected exactly {expected_names}, got {wb.sheetnames}"
    )
    # No correlation sheets present
    assert "correlation_matrix" not in wb.sheetnames
    assert "dependency_graph" not in wb.sheetnames


# ---- Test 2: correlations kwarg → ≥14 sheets --------------------------------


def test_xlsx_with_correlations_14_sheets(tmp_path: Path):
    """render_v8_xlsx with correlations kwarg must produce ≥13 sheets
    (9 base + 4 correlation) and include all 4 new named sheets."""
    from openpyxl import load_workbook
    from docdiffops.forensic_actions import apply_actions_to_bundle
    from docdiffops.forensic_render import render_v8_xlsx

    bundle, correlations = _make_correlations_bundle()
    # Add actions catalogue so base reaches 10 sheets → total ≥14 with correlations
    bundle = apply_actions_to_bundle(bundle)

    out = tmp_path / "corr.xlsx"
    result_path = render_v8_xlsx(bundle, out, correlations=correlations)

    # Return value must be the same path
    assert result_path == out

    wb = load_workbook(out)
    assert len(wb.sheetnames) >= 14, (
        f"Expected ≥14 sheets, got {len(wb.sheetnames)}: {wb.sheetnames}"
    )
    for sheet_name in ("correlation_matrix", "dependency_graph",
                       "claim_provenance", "coverage_heatmap"):
        assert sheet_name in wb.sheetnames, (
            f"Sheet '{sheet_name}' missing from {wb.sheetnames}"
        )


# ---- Test 3: correlation_matrix has ColorScaleRule -------------------------


def test_xlsx_correlation_heatmap_color_scale(tmp_path: Path):
    """Sheet 'correlation_matrix' must have ≥1 colorScale rule in its
    conditional_formatting ranges (ColorScaleRule is a factory in openpyxl
    that returns Rule objects with type='colorScale')."""
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    bundle, correlations = _make_correlations_bundle()
    out = tmp_path / "corr_cf.xlsx"
    render_v8_xlsx(bundle, out, correlations=correlations)

    wb = load_workbook(out)
    ws = wb["correlation_matrix"]

    # Collect all conditional formatting rules on this sheet.
    # ColorScaleRule is a factory function in openpyxl — it returns a Rule
    # object with type == "colorScale". Check via attribute, not isinstance.
    all_rules = []
    for _range, rules in ws.conditional_formatting._cf_rules.items():
        all_rules.extend(rules)

    color_scale_rules = [r for r in all_rules if getattr(r, "type", None) == "colorScale"]
    assert len(color_scale_rules) >= 1, (
        f"Expected ≥1 colorScale rule on correlation_matrix, "
        f"found {len(color_scale_rules)}. "
        f"All rule types: {[getattr(r, 'type', None) for r in all_rules]}"
    )


# ---- Test 4: dependency_graph rows sorted by relation_type, weight desc ----


def test_xlsx_dependency_graph_sorted(tmp_path: Path):
    """Sheet 'dependency_graph' must be sorted by relation_type ASC,
    then weight DESC."""
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    bundle, correlations = _make_correlations_bundle()
    out = tmp_path / "corr_sort.xlsx"
    render_v8_xlsx(bundle, out, correlations=correlations)

    wb = load_workbook(out)
    ws = wb["dependency_graph"]

    headers = [cell.value for cell in ws[1]]
    rt_idx = headers.index("relation_type")
    wt_idx = headers.index("weight")

    rows = list(ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=True))
    if len(rows) < 2:
        # Not enough rows to check ordering — vacuously true
        return

    for i in range(len(rows) - 1):
        rt_a = str(rows[i][rt_idx] or "")
        rt_b = str(rows[i + 1][rt_idx] or "")
        wt_a = float(rows[i][wt_idx] or 0)
        wt_b = float(rows[i + 1][wt_idx] or 0)

        # relation_type must be non-decreasing
        assert rt_a <= rt_b, (
            f"Row {i + 2}: relation_type out of order: "
            f"'{rt_a}' > '{rt_b}'"
        )
        # When relation_type equal, weight must be non-increasing
        if rt_a == rt_b:
            assert wt_a >= wt_b, (
                f"Row {i + 2}: weight out of order within same relation_type "
                f"'{rt_a}': {wt_a} < {wt_b}"
            )


# ---- Test 5: frozen panes on all 4 new sheets --------------------------------


def test_xlsx_frozen_panes_on_new_sheets(tmp_path: Path):
    """All 4 new correlation sheets must have freeze_panes set (non-None)."""
    from openpyxl import load_workbook
    from docdiffops.forensic_render import render_v8_xlsx

    bundle, correlations = _make_correlations_bundle()
    out = tmp_path / "corr_freeze.xlsx"
    render_v8_xlsx(bundle, out, correlations=correlations)

    wb = load_workbook(out)
    new_sheets = ("correlation_matrix", "dependency_graph",
                  "claim_provenance", "coverage_heatmap")
    for sheet_name in new_sheets:
        ws = wb[sheet_name]
        assert ws.freeze_panes is not None, (
            f"Sheet '{sheet_name}' has no frozen panes (freeze_panes is None)"
        )
