"""Tests for forensic_delta_render — XLSX/DOCX/PDF delta renderers."""
from __future__ import annotations

from pathlib import Path

import pytest

from docdiffops.forensic import build_forensic_bundle
from docdiffops.forensic_delta import compare_bundles
from docdiffops.forensic_delta_render import (
    DELTA_TITLE_RU,
    DIRECTION_RU,
    render_delta_docx,
    render_delta_pdf,
    render_delta_xlsx,
)


def _delta(old_status: str = "contradicts",
           new_status: str = "partial") -> dict:
    docs = [
        {"id": "D1", "code": "C1", "rank": 1, "title": "t1", "type": "law"},
        {"id": "D2", "code": "C2", "rank": 1, "title": "t2", "type": "law"},
    ]
    old_b = build_forensic_bundle(
        documents=docs,
        pairs=[{"id": "P1", "left": "D1", "right": "D2",
                "events": [{"status": old_status}]}],
        events=[], amendment_graph={},
    )
    new_b = build_forensic_bundle(
        documents=docs,
        pairs=[{"id": "P1", "left": "D1", "right": "D2",
                "events": [{"status": new_status}]},
               {"id": "P2", "left": "D1", "right": "D2",
                "events": [{"status": "same"}]}],  # P2 is new
        events=[], amendment_graph={},
    )
    return compare_bundles(old_b, new_b)


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def test_render_delta_xlsx_creates_workbook_with_cover_and_changes(tmp_path: Path):
    from openpyxl import load_workbook
    out = tmp_path / "delta.xlsx"
    render_delta_xlsx(_delta(), out)
    assert out.exists() and out.stat().st_size > 5000

    wb = load_workbook(out)
    assert "00 Обложка" in wb.sheetnames
    assert "01 Изменения статусов" in wb.sheetnames
    assert "02 Распределение (Δ)" in wb.sheetnames
    assert "03 Новые пары" in wb.sheetnames

    # Cover sheet has Russian title and KPI tile values
    ws_cover = wb["00 Обложка"]
    assert "Дельта-отчёт" in (ws_cover["A1"].value or "")
    all_text = " ".join(str(c.value) for row in ws_cover.iter_rows() for c in row if c.value)
    assert "Изменено" in all_text
    assert "Закрыто" in all_text


def test_render_delta_xlsx_status_changes_table_has_russian_columns(tmp_path: Path):
    from openpyxl import load_workbook
    out = tmp_path / "delta.xlsx"
    render_delta_xlsx(_delta(), out)
    ws = load_workbook(out)["01 Изменения статусов"]
    headers = [c.value for c in ws[1]]
    assert "old_status_ru" in headers
    assert "new_status_ru" in headers
    assert "direction_ru" in headers


def test_render_delta_xlsx_omits_new_pairs_sheet_when_no_new_pairs(tmp_path: Path):
    from openpyxl import load_workbook
    docs = [{"id": "D1", "code": "C1", "rank": 1, "title": "t1", "type": "law"},
            {"id": "D2", "code": "C2", "rank": 1, "title": "t2", "type": "law"}]
    b = build_forensic_bundle(documents=docs,
                              pairs=[{"id": "P1", "left": "D1", "right": "D2",
                                      "events": [{"status": "same"}]}],
                              events=[], amendment_graph={})
    delta = compare_bundles(b, b)  # identical → no new/removed pairs
    out = tmp_path / "delta.xlsx"
    render_delta_xlsx(delta, out)
    wb = load_workbook(out)
    assert "03 Новые пары" not in wb.sheetnames
    assert "04 Удалённые пары" not in wb.sheetnames


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_render_delta_docx_has_cover_and_russian_sections(tmp_path: Path):
    from docx import Document
    out = tmp_path / "delta.docx"
    render_delta_docx(_delta(), out)
    assert out.exists() and out.stat().st_size > 3000
    d = Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    assert DELTA_TITLE_RU in text
    assert "Раздел 1. Изменения статусов" in text
    assert "Краткая сводка" in text
    # Russian status labels in tables
    table_text = " ".join(c.text for tbl in d.tables for r in tbl.rows for c in r.cells)
    assert "Противоречие" in table_text or "Частичное совпадение" in table_text


def test_render_delta_docx_sets_russian_language_property(tmp_path: Path):
    from docx import Document
    out = tmp_path / "delta.docx"
    render_delta_docx(_delta(), out)
    d = Document(str(out))
    assert d.core_properties.language == "ru-RU"


def test_render_delta_docx_handles_empty_changes(tmp_path: Path):
    from docx import Document
    docs = [{"id": "D1", "code": "C1", "rank": 1, "title": "t1", "type": "law"},
            {"id": "D2", "code": "C2", "rank": 1, "title": "t2", "type": "law"}]
    b = build_forensic_bundle(documents=docs,
                              pairs=[{"id": "P1", "left": "D1", "right": "D2",
                                      "events": [{"status": "same"}]}],
                              events=[], amendment_graph={})
    delta = compare_bundles(b, b)
    out = tmp_path / "delta_empty.docx"
    render_delta_docx(delta, out)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "не зафиксировано" in text


def test_render_delta_docx_direction_translation_uses_russian_labels():
    # Direction labels should map to Russian
    assert DIRECTION_RU["improved"] == "Улучшилось"
    assert DIRECTION_RU["degraded"] == "Ухудшилось"
    assert DIRECTION_RU["unchanged"] == "Без изменений"


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def test_render_delta_pdf_creates_non_empty_file(tmp_path: Path):
    out = tmp_path / "delta.pdf"
    render_delta_pdf(_delta(), out)
    assert out.exists() and out.stat().st_size > 4000
    # PDF magic bytes
    assert out.read_bytes().startswith(b"%PDF")


@pytest.mark.skipif(__import__("shutil").which("pdftotext") is None,
                    reason="pdftotext not installed")
def test_render_delta_pdf_contains_russian_kpi_labels(tmp_path: Path):
    import subprocess
    out = tmp_path / "delta.pdf"
    render_delta_pdf(_delta(), out)
    txt = subprocess.run(["pdftotext", str(out), "-"],
                         capture_output=True).stdout.decode(errors="replace")
    assert "Дельта-отчёт" in txt
    assert "Изменено" in txt
    assert "Раздел 1" in txt
