"""Tests for docdiffops.forensic_note — PR-6.3 (Sprint 6).

6+ tests covering:
  1. test_docx_has_10_chapter_headings
  2. test_pdf_has_at_least_10_pages
  3. test_pdf_renders_cyrillic
  4. test_docx_table_of_contents_present
  5. test_font_fallback_chain
  6. test_empty_correlations_doesnt_crash
  7. test_docx_kpi_block_present
  8. test_pdf_renders_with_full_bundle
"""
from __future__ import annotations

import subprocess
import sys
import warnings
from pathlib import Path
from typing import Any
from unittest.mock import patch

import pytest

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

MINIMAL_DOCS = [
    {"id": "D01", "code": "FZ_115", "rank": 1, "title": "115-ФЗ", "type": "law",
     "short": "115-ФЗ О гражданстве", "date": "2025"},
    {"id": "D02", "code": "FZ_109", "rank": 1, "title": "109-ФЗ", "type": "law",
     "short": "109-ФЗ О миграционном учёте", "date": "2025"},
    {"id": "D03", "code": "PP_1510", "rank": 2, "title": "ПП-1510", "type": "regulation",
     "short": "ПП 1510 о порядке", "date": "2025"},
    {"id": "D04", "code": "VCIOM_01", "rank": 3, "title": "ВЦИОМ", "type": "analytics",
     "short": "ВЦИОМ опрос 2025", "date": "2025"},
    {"id": "D05", "code": "CONC_2030", "rank": 1, "title": "Концепция 2026-2030",
     "type": "strategy", "short": "Концепция миграционной политики", "date": "2026"},
]

MINIMAL_EVENTS = [
    {"id": "E001", "event_id": "E001", "status": "contradiction", "risk": "высокий",
     "theme": "Регистрация", "confidence": 0.91,
     "left_doc": "115-ФЗ", "right_doc": "ВЦИОМ",
     "source_rank_left": 1, "source_rank_right": 3,
     "claim_left": "Регистрация обязательна", "conclusion": "Уточнить источник"},
    {"id": "E002", "event_id": "E002", "status": "partial_overlap", "risk": "средний",
     "theme": "Трудоустройство", "confidence": 0.78,
     "left_doc": "109-ФЗ", "right_doc": "ПП-1510"},
    {"id": "E003", "event_id": "E003", "status": "manual_review", "risk": "низкий",
     "theme": "Квоты", "confidence": 0.65,
     "left_doc": "ПП-1510", "right_doc": "Концепция 2026-2030"},
]

MINIMAL_PAIRS = [
    {"id": "P1", "left": "D01", "right": "D02",
     "events": [{"status": "partial_overlap"}]},
    {"id": "P2", "left": "D01", "right": "D04",
     "events": [{"status": "contradiction"}]},
    {"id": "P3", "left": "D02", "right": "D05",
     "events": [{"status": "manual_review"}]},
]


def _make_bundle() -> dict[str, Any]:
    from docdiffops.forensic import build_forensic_bundle
    bundle: dict[str, Any] = build_forensic_bundle(
        documents=MINIMAL_DOCS,
        pairs=MINIMAL_PAIRS,
        events=MINIMAL_EVENTS,
        amendment_graph={},
    )
    return bundle


def _make_correlations() -> dict[str, Any]:
    """Synthetic correlations dict matching the expected schema."""
    return {
        "correlation_matrix": {
            "T01": {"D01": 3, "D02": 2, "D04": 1},
            "T02": {"D02": 2, "D05": 4},
            "T03": {"D03": 1},
        },
        "dependency_graph": [
            {"from_doc_id": "D01", "to_doc_id": "D04", "relation_type": "references", "weight": 2},
            {"from_doc_id": "D02", "to_doc_id": "D03", "relation_type": "amends", "weight": 3},
        ],
        "claim_provenance": [
            {
                "thesis_id": "TH001",
                "thesis_text": "Регистрация иностранных граждан обязательна",
                "primary_doc_id": "D01",
                "primary_rank": 1,
                "confirming_docs": "D02;D05",
                "refuting_docs": "",
                "evidence_event_ids": "E001",
            },
            {
                "thesis_id": "TH002",
                "thesis_text": "Трудовая квота устанавливается ежегодно",
                "primary_doc_id": "D03",
                "primary_rank": 2,
                "confirming_docs": "D05",
                "refuting_docs": "D04",
                "evidence_event_ids": "E002",
            },
        ],
        "coverage_heatmap": [
            {"theme_id": "T01", "theme_name": "Регистрация",
             "rank_1_count": 2, "rank_2_count": 1, "rank_3_count": 1},
            {"theme_id": "T02", "theme_name": "Трудоустройство",
             "rank_1_count": 0, "rank_2_count": 1, "rank_3_count": 3},
        ],
    }


# ---------------------------------------------------------------------------
# Test 1: DOCX has exactly 10+ Heading 1 paragraphs (one per chapter)
# ---------------------------------------------------------------------------

def test_docx_has_10_chapter_headings(tmp_path: Path) -> None:
    """DOCX must have at least 10 Heading 1 paragraphs (one per chapter)."""
    from docx import Document
    from docdiffops.forensic_note import render_explanatory_note_docx

    out = tmp_path / "note.docx"
    render_explanatory_note_docx(_make_bundle(), _make_correlations(), out)

    doc = Document(str(out))
    heading1_paras = [
        p for p in doc.paragraphs
        if p.style.name.startswith("Heading 1")
    ]
    assert len(heading1_paras) >= 10, (
        f"Expected ≥10 Heading 1 paragraphs, found {len(heading1_paras)}: "
        + str([p.text for p in heading1_paras])
    )


# ---------------------------------------------------------------------------
# Test 2: PDF has at least 10 pages
# ---------------------------------------------------------------------------

def test_pdf_has_at_least_10_pages(tmp_path: Path) -> None:
    """PDF must have at least 10 pages (one per chapter)."""
    from docdiffops.forensic_note import render_explanatory_note_pdf

    out = tmp_path / "note.pdf"
    render_explanatory_note_pdf(_make_bundle(), _make_correlations(), out)

    assert out.exists(), "PDF file was not created"
    assert out.stat().st_size > 0, "PDF file is empty"

    # Try pdfinfo first, fall back to pypdf/PyPDF2
    page_count = _get_pdf_page_count(out)
    assert page_count >= 10, f"Expected ≥10 pages, got {page_count}"


def _get_pdf_page_count(pdf_path: Path) -> int:
    """Return page count via pdfinfo (preferred) or pypdf fallback."""
    result = subprocess.run(
        ["pdfinfo", str(pdf_path)],
        capture_output=True,
        text=True,
    )
    if result.returncode == 0:
        for line in result.stdout.splitlines():
            if line.startswith("Pages:"):
                return int(line.split(":")[1].strip())

    # Fallback: pypdf
    try:
        import pypdf  # type: ignore[import]
        with open(pdf_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            return len(reader.pages)
    except ImportError:
        pass

    # Fallback 2: count %%Page markers
    data = pdf_path.read_bytes()
    count = data.count(b"/Page ")
    return max(count, 1)


# ---------------------------------------------------------------------------
# Test 3: PDF renders Cyrillic text ("Введение")
# ---------------------------------------------------------------------------

def test_pdf_renders_cyrillic(tmp_path: Path) -> None:
    """pdftotext output must contain the word 'Введение'."""
    from docdiffops.forensic_note import render_explanatory_note_pdf

    out = tmp_path / "note.pdf"
    render_explanatory_note_pdf(_make_bundle(), _make_correlations(), out)

    result = subprocess.run(
        ["pdftotext", str(out), "-"],
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    if result.returncode != 0:
        pytest.skip("pdftotext not available — skipping Cyrillic check")

    text = result.stdout
    assert "Введение" in text or "введение" in text.lower(), (
        "Cyrillic 'Введение' not found in PDF text — possible mojibake. "
        f"First 500 chars: {text[:500]!r}"
    )


# ---------------------------------------------------------------------------
# Test 4: DOCX has a table of contents ("Содержание") with chapter entries
# ---------------------------------------------------------------------------

def test_docx_table_of_contents_present(tmp_path: Path) -> None:
    """DOCX must contain a 'Содержание' heading and TOC entries."""
    from docx import Document
    from docdiffops.forensic_note import render_explanatory_note_docx

    out = tmp_path / "note.docx"
    render_explanatory_note_docx(_make_bundle(), _make_correlations(), out)

    doc = Document(str(out))
    all_text = " ".join(p.text for p in doc.paragraphs)
    # Must have the TOC heading
    assert "Содержание" in all_text, (
        "'Содержание' heading not found in DOCX paragraphs"
    )
    # Must have at least the first chapter mentioned in the TOC body
    assert "Введение" in all_text, (
        "'Введение' not found in DOCX — TOC or chapter missing"
    )
    # Must have the 10th chapter
    assert "Приложения" in all_text, (
        "'Приложения' not found — chapter 10 or TOC entry missing"
    )


# ---------------------------------------------------------------------------
# Test 5: Font fallback chain — when NotoSans is absent, fallback fires
# ---------------------------------------------------------------------------

def test_font_fallback_chain(tmp_path: Path) -> None:
    """When NotoSans TTF doesn't exist, DejaVu or Helvetica fallback is used."""
    from docdiffops.forensic_note import _register_cyrillic_font_for_pdf

    # Patch os.path.exists to make NotoSans appear absent
    original_exists = __import__("os").path.exists

    def _fake_exists(path: str) -> bool:
        from docdiffops.forensic_note import _NOTO_TTF
        if path == _NOTO_TTF:
            return False
        return original_exists(path)

    with patch("os.path.exists", side_effect=_fake_exists):
        # Re-import so it picks up the patch in the module's os.path.exists
        with patch("docdiffops.forensic_note._NOTO_TTF", "/nonexistent/NotoSans.ttf"):
            # Either DejaVu found → returns "DejaVuSans"
            # or Helvetica fallback → emits RuntimeWarning and returns "Helvetica"
            with warnings.catch_warnings(record=True) as caught:
                warnings.simplefilter("always")
                font_used = _register_cyrillic_font_for_pdf()

            from docdiffops.forensic_note import _DEJAVU_TTF
            import os
            if os.path.exists(_DEJAVU_TTF):
                assert font_used == "DejaVuSans", (
                    f"Expected DejaVuSans fallback, got {font_used!r}"
                )
            else:
                assert font_used == "Helvetica", (
                    f"Expected Helvetica last-resort fallback, got {font_used!r}"
                )
                rw = [w for w in caught if issubclass(w.category, RuntimeWarning)]
                assert rw, "Expected RuntimeWarning when falling back to Helvetica"


# ---------------------------------------------------------------------------
# Test 6: Empty correlations dict does not crash
# ---------------------------------------------------------------------------

def test_empty_correlations_doesnt_crash(tmp_path: Path) -> None:
    """Both render functions must succeed with an empty correlations dict."""
    from docdiffops.forensic_note import (
        render_explanatory_note_docx,
        render_explanatory_note_pdf,
    )
    from docx import Document

    bundle = _make_bundle()
    empty_corr: dict[str, Any] = {}

    docx_out = tmp_path / "empty_corr.docx"
    pdf_out = tmp_path / "empty_corr.pdf"

    # Must not raise
    render_explanatory_note_docx(bundle, empty_corr, docx_out)
    render_explanatory_note_pdf(bundle, empty_corr, pdf_out)

    assert docx_out.exists() and docx_out.stat().st_size > 0
    assert pdf_out.exists() and pdf_out.stat().st_size > 0

    # DOCX should still have 10 headings
    doc = Document(str(docx_out))
    heading1_paras = [p for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    assert len(heading1_paras) >= 10


# ---------------------------------------------------------------------------
# Test 7: DOCX KPI block is present
# ---------------------------------------------------------------------------

def test_docx_kpi_block_present(tmp_path: Path) -> None:
    """DOCX must contain a KPI summary table with 'Документов' label."""
    from docx import Document
    from docdiffops.forensic_note import render_explanatory_note_docx

    out = tmp_path / "note_kpi.docx"
    render_explanatory_note_docx(_make_bundle(), _make_correlations(), out)

    doc = Document(str(out))
    # Collect all cell text from tables
    all_table_text = " ".join(
        cell.text
        for tbl in doc.tables
        for row in tbl.rows
        for cell in row.cells
    )
    assert "Документов" in all_table_text, (
        "KPI label 'Документов' not found in any DOCX table"
    )


# ---------------------------------------------------------------------------
# Test 8: PDF renders with full bundle (no crash, non-trivial size)
# ---------------------------------------------------------------------------

def test_pdf_renders_with_full_bundle(tmp_path: Path) -> None:
    """PDF render with populated bundle + correlations produces non-trivial file."""
    from docdiffops.forensic_note import render_explanatory_note_pdf

    out = tmp_path / "full.pdf"
    render_explanatory_note_pdf(
        _make_bundle(),
        _make_correlations(),
        out,
        pipeline_version="10.0.0",
        generated_at="2026-05-09T10:00:00Z",
    )
    assert out.exists()
    size = out.stat().st_size
    # A minimal 10-chapter PDF should be at least 20 KB
    assert size > 20_000, f"PDF suspiciously small: {size} bytes"
